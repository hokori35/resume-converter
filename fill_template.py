import json
import re
import subprocess
import tempfile
import argparse
from datetime import datetime
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree
try:
    import win32com.client  # type: ignore
except Exception:
    win32com = None

try:
    from batch_extract_json import (
        find_valid_issn_in_text,
        normalize_experience_date_jin,
        parse_pub_info as _parse_pub_info_impl,
    )
except ImportError:
    _parse_pub_info_impl = None
    find_valid_issn_in_text = None

    def normalize_experience_date_jin(text):
        import re as _re

        s = _re.sub(r"\s+", " ", (text or "").strip())
        s = _re.sub(
            r"(\d{4}\.\d{1,2})\s*[-～–—－]\s*今(?!至)",
            r"\1-至今",
            s,
        )
        s = _re.sub(r"(\d{4}\.\d{1,2})\s+今(?!至)", r"\1-至今", s)
        s = _re.sub(r"(\d{4}\.\d{1,2})\s+至今", r"\1-至今", s)
        return s


def set_font(
    run,
    font_cn="宋体",
    font_en="Times New Roman",
    size_pt=10.5,
    bold=False,
    underline=False,
    italic=False,
):
    """同时设置中英文字体、字号及样式。"""
    run.font.name = font_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.underline = underline
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)


def set_paragraph_format(paragraph, first_line_chars=0, line_spacing=1.0):
    """设置段落格式：首行缩进(按中文字符宽度近似)与行距。"""
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(first_line_chars * 10.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = line_spacing


def _is_other_section_subtitle_line(line):
    """
    「其他成果」小节标题：顶格（无首行缩进）、正文加粗。
    如「其他论文：」「1.会议论文：」「2.发明专利：」「1.授权发明专利：」「2.获奖情况：」或单独成行短标题「发明专利：」；
    另支持「1.实践获奖」「2.国际赛事获奖」类编号小标题（可无句末冒号；与条目「（1）…」区分）。
    「1.单位,2018-2022. 助教」「1. IEDA…」等教学/课程枚举不算标题（首行缩进 2、不加粗）。
    枚举条「（1）…」不算标题。
    """
    t = (line or "").strip()
    if not t:
        return False
    if t[0] in "（(":
        return False
    # 「一、其他论文」「二、科研奖励」等：顶格加粗（勿把「一、多目标…」长项目句当标题）
    if re.match(r"^[一二三四五六七八九十]+、", t):
        if len(t) <= 14:
            return True
        if len(t) <= 48 and any(
            k in t
            for k in (
                "论文",
                "奖励",
                "专利",
                "成果",
                "业绩",
                "著作",
                "会议",
                "标准",
                "软件",
                "竞赛",
            )
        ):
            return True
    # 「1.单位,2018-2022. 助教」「1. IEDA…」等：正文枚举，非小节标题（否则顶格加粗、无首行缩进）
    if re.match(r"^\d+[.。、]\s*", t):
        if re.search(r",\s*20\d{2}|，\s*20\d{2}", t):
            return False
        if "助教" in t or "讲师" in t or "副教授" in t:
            return False
        if re.search(r"^\d+[.。、]\s*IEDA", t, re.I):
            return False
    if re.match(r"^\d+[.。、]\s*.+?[：:]\s*$", t):
        return True
    # 行末为冒号的标题行（放宽长度，避免长标签被当成正文而缩进）
    if re.search(r"[：:]\s*$", t) and len(t) <= 120:
        return True
    # 「1.2023.07获得…」「2.2024年…」等：序号后紧跟年份，是条目非小节标题
    if re.match(r"^\d+[.。、]\s*\d{4}[./年]", t):
        return False
    # 「1.实践获奖」：行末可无冒号；含「（」一般为（1）条目或署名说明，不作为小节标题
    if (
        len(t) <= 36
        and re.match(r"^\d+[.。、]\s*\S", t)
        and "（" not in t
        and "(" not in t
    ):
        return True
    return False


def _apply_other_paragraph_layout(paragraph, *, is_subtitle):
    """其他栏：小标题顶格（左缩进 0、首行缩进 0）；正文首行缩进 2 字符。"""
    _tight_paragraph(paragraph)
    pf = paragraph.paragraph_format
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    set_paragraph_format(paragraph, first_line_chars=0 if is_subtitle else 2, line_spacing=1.0)


def _normalize_other_raw_cn_sections(s):
    """
    将「一、…二、…三、…」挤在同一行时的「二、」等前插入换行，便于按段加粗小标题。
    """
    t = (s or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not t or t.count("\n") >= 4:
        return t
    if "一、" in t[:30] and re.search(r"(?:二|三|四|五|六|七|八|九|十、)", t):
        t = re.sub(
            r"(?=二、|三、|四、|五、|六、|七、|八、|九、|十、)",
            "\n",
            t,
        )
    if re.search(r"[（(][一二三四五六七八九十]+[）)]", t):
        t = re.sub(
            r"(?=[（(](?:二|三|四|五|六|七|八|九|十)[）)])",
            "\n",
            t,
        )
    if re.search(r"^[一二三四五六七八九十]+、", t, re.M):
        t = re.sub(
            r"(?=(?<=\n)(?:二|三|四|五|六|七|八|九|十)、)|(?=^(?:二|三|四|五|六|七|八|九|十)、)",
            "\n",
            t,
            flags=re.M,
        )
    return t.strip()


_OTHER_CN_SECTION_HEAD = re.compile(
    r"^[（(]([一二三四五六七八九十]+)[）)]\s*(.+?)\s*$"
)
_OTHER_BRACKET_ITEM_SPLIT = re.compile(r"(?=\[\d+\]\s*)")


def _other_strip_item_number_prefix(s):
    t = (s or "").strip()
    t = re.sub(r"^\[\d+\]\s*", "", t)
    t = re.sub(r"^[（(]\s*\d+\s*[）)]\s*", "", t)
    # 去掉行首 1. / 2. / 1. IEDA 等，统一由（1）（2）重编
    if re.match(r"^\d+[.。、]\s*", t):
        t = re.sub(r"^\d+[.。、]\s*", "", t)
    return t


def _other_is_soft_copyright_item(text):
    plain = _other_plain_text_for_subtitle_detect(text)
    if re.search(r"\d{4}SR\d+", plain, re.I):
        return True
    if "软著" in plain:
        return True
    if re.search(r"V\d+(?:\.\d+)?", plain, re.I) and not re.search(
        r"\bZL\s*\d", plain, re.I
    ):
        return "场景监视" in plain or "系统V" in plain
    return False


def _other_is_invention_patent_item(text):
    plain = _other_plain_text_for_subtitle_detect(text)
    if re.search(r"\bZL\s*\d", plain, re.I):
        return True
    if "授权日期" in plain:
        return True
    if "发明专利" in plain and "软著" not in plain:
        return True
    return False


def _other_canonical_section_title(title):
    t = (title or "").strip().rstrip("：:")
    if "出版专著" in t or "专著" in t and "专利" not in t:
        return "出版专著"
    if "荣誉" in t or "奖励" in t or "获奖" in t:
        return "荣誉与奖励"
    if "软著" in t and "专利" not in t:
        return "软著"
    if "发明专利" in t or "授权发明" in t:
        return "发明专利"
    return t


def _other_split_line_items(ln):
    line = (ln or "").strip()
    if not line:
        return []
    if not _OTHER_BRACKET_ITEM_SPLIT.search(line):
        return [line]
    parts = [p.strip() for p in _OTHER_BRACKET_ITEM_SPLIT.split(line) if p.strip()]
    return parts or [line]


_OTHER_BLOCK_SECTION_KEYWORDS = (
    "演讲",
    "研讨",
    "教学",
    "竞赛",
    "荣誉",
    "奖励",
    "专利",
    "软著",
    "专著",
    "著作",
    "论文",
    "报告",
    "兼职",
    "发明",
    "项目",
    "标准",
    "任职",
    "经历",
    "成果",
)


def _other_line_is_standalone_keyword_title(plain):
    """单独成行的小节名（无冒号）：专利、获奖、软著 等。"""
    t = (plain or "").strip().rstrip("：:，,")
    if not t or len(t) > 20:
        return False
    if re.match(r"^[（(]", t) or re.match(r"^\d", t):
        return False
    if re.search(r"[：:]", t):
        return False
    if t in _OTHER_BLOCK_SECTION_KEYWORDS:
        return True
    if len(t) <= 12 and any(
        kw in t and len(t) <= len(kw) + 6 for kw in _OTHER_BLOCK_SECTION_KEYWORDS
    ):
        return True
    return False


def _other_line_is_short_colon_title(plain):
    """短行冒号结尾的小标题：其他论文：、专利：、软件著作权:（非 1. 2018 类条目）。"""
    t = (plain or "").strip()
    if len(t) < 2 or len(t) > 40:
        return False
    if not re.search(r"[：:]\s*$", t):
        return False
    if re.match(r"^[（(]", t):
        return False
    if re.match(r"^\d+[.。、]\s*\d{4}[./年\-]", t):
        return False
    if re.match(r"^\d+[.。、]\s+\d", t) and not re.match(
        r"^\d+[.。、]\s*.{1,20}[：:]\s*$", t
    ):
        return False
    head = t.rstrip("：:")
    if len(head) > 28:
        return False
    if head.count("，") >= 2 or head.count(",") >= 3:
        return False
    if re.match(r"^\d+[.。、]\s*.+?[：:]\s*$", t) and len(t) <= 48:
        return True
    if not re.match(r"^\d", t):
        return True
    return False


def _other_line_is_block_section_header(plain, raw=""):
    """块级小节标题：[**演讲：**]、（一）…、一、其他论文、研讨会：、1.软著： 等。"""
    t = (plain or "").strip()
    r = (raw or "").strip()
    if not t:
        return False
    if re.match(r"^\[\*\*.+?\*\*\]\s*[：:]?\s*$", r):
        return True
    if _OTHER_CN_SECTION_HEAD.match(t):
        return True
    if re.match(r"^[一二三四五六七八九十]+、", t) and len(t) <= 36:
        return True
    if re.match(r"^\d+[.。、]\s*.+?[：:]\s*$", t) and len(t) <= 48:
        return True
    if _other_line_is_short_colon_title(t):
        return True
    if _other_line_is_standalone_keyword_title(t):
        return True
    return False


def _other_promote_leading_keyword_title(title, items):
    """首行实为「专利」等无冒号小标题时，提升为 section 标题并从条目中剔除。"""
    if title:
        return title, items
    if not items:
        return title, items
    head_plain = _other_plain_text_for_subtitle_detect(items[0]).strip()
    if _other_line_is_standalone_keyword_title(head_plain):
        return _other_canonical_section_title(head_plain), items[1:]
    return title, items


def _other_infer_section_title_from_items(items):
    """无标题小节：根据条目内容推断（如均为中国专利/美国专利）。"""
    joined = " ".join(_other_plain_text_for_subtitle_detect(x) for x in (items or [])[:6])
    if "专利" in joined or re.search(r"\b(?:CN|US|TWI)\d", joined, re.I):
        return "专利"
    if _other_is_soft_copyright_item(joined):
        return "软著"
    if any(k in joined for k in ("奖学金", "获奖", "奖励", "优秀", "称号")):
        return "荣誉与奖励"
    return "其他成果"


def _other_section_title_from_header(plain, raw=""):
    t = (plain or "").strip()
    m = _OTHER_CN_SECTION_HEAD.match(t)
    if m:
        return m.group(2).strip().rstrip("：:")
    m = re.match(r"^([一二三四五六七八九十]+)、\s*(.+)$", t)
    if m:
        return m.group(2).strip().rstrip("：:") or m.group(0).rstrip("：:")
    m = re.match(r"^\d+[.。、]\s*(.+?)[：:]\s*$", t)
    if m:
        return m.group(1).strip()
    m = re.match(r"^\[\*\*(.+?)\*\*\]\s*[：:]?\s*$", (raw or "").strip())
    if m:
        return m.group(1).strip().rstrip("：:")
    if _other_line_is_standalone_keyword_title(t):
        return _other_canonical_section_title(t)
    return t.rstrip("：:")


def _other_line_is_subsection_header(plain):
    """已废弃路径保留：与块级标题判定一致。"""
    return _other_line_is_block_section_header(plain)


def _other_line_keep_original_numbering(plain):
    """仅保留 ①②③ 等特殊序号；阿拉伯数字序号一律改为（1）（2）。"""
    t = (plain or "").strip()
    return bool(re.match(r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫]\s*", t))


def _other_already_numbered_sections(t):
    """已是 1.演讲：/2.研讨会： 等多节编号则不再整体重排。"""
    if re.search(r"^\[\*\*", t or "", re.M):
        return False
    if re.search(r"^[（(][一二三四五六七八九十]+[）)]", t or "", re.M):
        return False
    heads = re.findall(r"^\d+[.。、]\S{1,24}[：:]", t or "", re.M)
    return len(heads) >= 2


def _parse_other_sections(t):
    """返回 [(title, items, from_markdown_header), ...]"""
    sections = []
    cur_title = None
    cur_items = []
    cur_md = False

    def _flush():
        nonlocal cur_title, cur_items, cur_md
        tit, its = _other_promote_leading_keyword_title(cur_title, list(cur_items))
        if tit is not None:
            sections.append((tit, its, cur_md))
        elif its:
            sections.append(("", its, False))
        cur_title = None
        cur_items = []
        cur_md = False

    for ln in (t or "").split("\n"):
        for chunk in _other_split_line_items(ln):
            plain = _other_plain_text_for_subtitle_detect(chunk).strip()
            raw = chunk.strip()
            if not plain:
                continue
            if _other_line_is_block_section_header(plain, raw):
                _flush()
                cur_title = _other_section_title_from_header(plain, raw)
                cur_md = bool(re.match(r"^\[\*\*", raw))
                continue
            cur_items.append(chunk)
    tit, its = _other_promote_leading_keyword_title(cur_title, list(cur_items))
    if tit is not None:
        sections.append((tit, its, cur_md))
    elif its:
        sections.append(("", its, False))
    return sections


def _other_is_plain_numbered_list(t):
    """全文仅为「1. 条目」枚举、无小节标题时勿整体重排。"""
    lines = [ln.strip() for ln in (t or "").split("\n") if ln.strip()]
    if not lines:
        return False
    for ln in lines:
        plain = _other_plain_text_for_subtitle_detect(ln).strip()
        if _other_line_is_block_section_header(plain, ln):
            return False
    return all(
        re.match(r"^\d+[.。、]\s+", _other_plain_text_for_subtitle_detect(ln).strip())
        for ln in lines
    )


def _other_should_structure_normalize(t):
    """凡能识别出小节标题，或含（一）/[1]/[** 等结构，均自动编号为 1.标题：。"""
    if not (t or "").strip():
        return False
    if _other_already_numbered_sections(t):
        return False
    if _other_is_plain_numbered_list(t):
        return False
    sections = _parse_other_sections(t)
    if any(tit for tit, _, _ in sections):
        return True
    if re.search(r"^\[\*\*", t, re.M):
        return True
    if re.search(r"^[（(][一二三四五六七八九十]+[）)]", t, re.M):
        return True
    if re.search(r"^\[\d+\]", t, re.M):
        return True
    if re.search(r"^[一二三四五六七八九十]+、", t, re.M):
        return True
    return False


def _other_renumber_items(items, *, keep_bracket_nums=False):
    """
    第二层条目：（1）（2）…，或保留 [1][2]（[**小节**] 来源时）；
    第三级 1.HE、1. IEDA、① 等保持原编号。
    """
    out = []
    item_no = 1
    for raw in items:
        plain = _other_plain_text_for_subtitle_detect(raw).strip()
        if not plain:
            continue
        if keep_bracket_nums and re.match(r"^\[\d+\]", plain):
            out.append(raw.strip())
            continue
        if _other_line_keep_original_numbering(plain):
            out.append(raw.strip())
            continue
        body = _other_strip_item_number_prefix(raw)
        if not body:
            continue
        out.append(f"（{item_no}）{body}")
        item_no += 1
    return out


def _other_format_numbered_section(sec_no, title, items, *, keep_bracket_nums=False):
    title = _other_canonical_section_title(title)
    if not title:
        title = _other_infer_section_title_from_items(items)
    lines = [f"{sec_no}.{title}："]
    lines.extend(_other_renumber_items(items, keep_bracket_nums=keep_bracket_nums))
    return lines


def _normalize_other_achievements_structure(s):
    """
    「其他」栏：将 [**演讲：**]/（一）/[1] 等转为 1.小节：（1）条目；
    其下 1.HE、1. IEDA 等更深层编号保持原样。
    """
    t = _normalize_other_raw_cn_sections(s)
    if not t or not _other_should_structure_normalize(t):
        return t

    sections = _parse_other_sections(t)
    if not sections:
        return t

    out_lines = []
    sec_no = 1
    for title, items, from_md in sections:
        title, items = _other_promote_leading_keyword_title(title or "", items)
        if "发明专利" in title and "软著" in title:
            soft, patent = [], []
            for it in items:
                if _other_is_soft_copyright_item(it):
                    soft.append(it)
                elif _other_is_invention_patent_item(it):
                    patent.append(it)
                elif re.search(r"\d{4}SR\d+", _other_plain_text_for_subtitle_detect(it), re.I):
                    soft.append(it)
                else:
                    patent.append(it)
            if soft:
                out_lines.extend(_other_format_numbered_section(sec_no, "软著", soft))
                sec_no += 1
            if patent:
                out_lines.extend(
                    _other_format_numbered_section(sec_no, "发明专利", patent)
                )
                sec_no += 1
            continue
        clean = _other_canonical_section_title(title)
        out_lines.extend(
            _other_format_numbered_section(
                sec_no, clean, items, keep_bracket_nums=from_md
            )
        )
        sec_no += 1

    return "\n".join(out_lines).strip() if out_lines else t


def clear_cell(cell):
    """清空单元格全部段落，仅保留一个空段落容器。"""
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph("")


def clear_paragraph_runs(paragraph):
    """只清空 run，保留段落本身的原有段落格式。"""
    p = paragraph._p
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)


def append_run(
    paragraph,
    text,
    *,
    font_cn,
    font_en,
    size_pt,
    bold=False,
    underline=False,
    italic=False,
    color_rgb=None,
):
    run = paragraph.add_run(text)
    set_font(
        run,
        font_cn=font_cn,
        font_en=font_en,
        size_pt=size_pt,
        bold=bold,
        underline=underline,
        italic=italic,
    )
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(color_rgb[0], color_rgb[1], color_rgb[2])
    return run


_STAT_DETAIL_DECIMAL_TOKEN = re.compile(r"\d+\.\d+")


def _append_stats_detail_text_runs(paragraph, text, *, font_cn, font_en, size_pt, bold=False):
    """
    论文统计明细段：真实小数（如 2.5、1.5）整段黑色正常显示；
    仅为与「半篇」等小数对齐而写成的「N.0」，将小数点及末尾的 0 设为白色，肉眼只见整数 N，占位仍在。
    """
    s = str(text)
    pos = 0
    white = (255, 255, 255)
    for m in _STAT_DETAIL_DECIMAL_TOKEN.finditer(s):
        if m.start() > pos:
            append_run(
                paragraph,
                s[pos : m.start()],
                font_cn=font_cn,
                font_en=font_en,
                size_pt=size_pt,
                bold=bold,
            )
        token = m.group(0)
        m_align_zero = re.fullmatch(r"(\d+)\.(0+)", token)
        if m_align_zero:
            append_run(
                paragraph,
                m_align_zero.group(1),
                font_cn=font_cn,
                font_en=font_en,
                size_pt=size_pt,
                bold=bold,
            )
            append_run(
                paragraph,
                "." + m_align_zero.group(2),
                font_cn=font_cn,
                font_en=font_en,
                size_pt=size_pt,
                bold=bold,
                color_rgb=white,
            )
        else:
            append_run(
                paragraph,
                token,
                font_cn=font_cn,
                font_en=font_en,
                size_pt=size_pt,
                bold=bold,
            )
        pos = m.end()
    if pos < len(s):
        append_run(
            paragraph,
            s[pos:],
            font_cn=font_cn,
            font_en=font_en,
            size_pt=size_pt,
            bold=bold,
        )


def append_run_bold_underline_digits(paragraph, text, *, font_cn, font_en, size_pt):
    """整段加粗；数字及其两侧空格一并加下划线（拉长数字下划线，贴近空白版效果）。"""
    s = str(text)
    pos = 0
    # 仅在“空格+数字+空格”场景下延展下划线，避免误伤日期等连续数字文本。
    for m in re.finditer(r" {2}\d+ {2}", s):
        if m.start() > pos:
            append_run(
                paragraph,
                s[pos : m.start()],
                font_cn=font_cn,
                font_en=font_en,
                size_pt=size_pt,
                bold=True,
                underline=False,
            )
        append_run(
            paragraph,
            m.group(),
            font_cn=font_cn,
            font_en=font_en,
            size_pt=size_pt,
            bold=True,
            underline=True,
        )
        pos = m.end()
    # 兜底：若仍存在未命中的数字（例如非双空格包裹），按原规则加下划线。
    tail = s[pos:]
    if tail:
        tpos = 0
        for m in re.finditer(r"\d+", tail):
            if m.start() > tpos:
                append_run(
                    paragraph,
                    tail[tpos : m.start()],
                    font_cn=font_cn,
                    font_en=font_en,
                    size_pt=size_pt,
                    bold=True,
                    underline=False,
                )
            append_run(
                paragraph,
                m.group(),
                font_cn=font_cn,
                font_en=font_en,
                size_pt=size_pt,
                bold=True,
                underline=True,
            )
            tpos = m.end()
        if tpos < len(tail):
            append_run(paragraph, tail[tpos:], font_cn=font_cn, font_en=font_en, size_pt=size_pt, bold=True, underline=False)


def add_page_break_before_paragraph(paragraph):
    """在段落前插入分页符段落。"""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    paragraph._p.addprevious(p)


def normalize_numbered_parens_to_fullwidth(s):
    """将 (1) 类半角编号括号转为全角 （1），正文中的英文圆括号保留。"""
    return re.sub(r"\(\s*(\d+)\s*\)", r"（\1）", str(s or ""))


def normalize_task_section_parens(s):
    """
    岗位任务 / 学院举措正文行：条目编号统一为全角圆括号（1）（2）。
    支持行首「1）」「2)」及半角「(1)」；避免把「1995）」等年份误改。
    """
    out_lines = []
    for line in str(s or "").splitlines():
        t = line.strip()
        m = re.match(r"^(\d+)[）)]\s*(.*)$", t)
        if m:
            num, rest = m.group(1), m.group(2)
            if len(num) == 4 and num.isdigit() and 1900 <= int(num) <= 2099:
                out_lines.append(line)
            else:
                out_lines.append(f"（{num}）{rest}".rstrip() if rest else f"（{num}）")
        else:
            out_lines.append(re.sub(r"\(\s*(\d+)\s*\)", r"（\1）", line))
    return "\n".join(out_lines)


def split_political_paragraphs(text):
    """思想政治段落分段：仅按显式换行分段；无换行时保持整段不拆分。"""
    raw = str(text or "").strip()
    if not raw:
        return []
    if "\n" in raw:
        return [p.strip() for p in re.split(r"\n+", raw) if p.strip()]
    return [raw]


def format_with_double_space(items):
    """经历条目格式化：时间后空两格、学校与专业间空两格。"""
    _DATE_PAT = r"(\d{4}\.\d{1,2}-\d{4}\.\d{1,2}|\d{4}\.\d{1,2}[-\s]?至今)"
    out = []
    for item in items:
        s = normalize_experience_date_jin(str(item).strip())
        m = re.match(rf"^(.*?:)\s*{_DATE_PAT}\s+(.+?)\s+(.+)$", s)
        if m:
            date = _norm_date(m.group(2))
            out.append(f"{m.group(1)}{date}  {m.group(3)}{' ' * 8}{m.group(4)}")
        else:
            out.append(s)
    return out


def normalize_birth_text(text):
    """将出生信息规范为两行：日期 + 年龄。"""
    s = str(text or "").strip()
    m = re.match(r"^(\d{4}\.\d{2})[（(]?\s*(\d+岁)\s*[)）]?$", s)
    if m:
        return f"{m.group(1)}\n（{m.group(2)}）"
    return s


def parse_marked_authors(text):
    """
    与 batch_extract_json.read_cell_marked 成对：从 JSON 还原作者单元格样式。
    [**xxx**]{.underline} -> 加粗+下划线
    [[xxx]]{.underline} -> 仅下划线
    [**xxx**] -> 仅加粗（后面不能紧跟 {.underline}，避免与第一种冲突）
    其余 -> 常规
    """
    token = re.compile(
        r"\[\*\*(.*?)\*\*\]\{\.underline\}|"
        r"\[\[(.*?)\]\]\{\.underline\}|"
        r"\[\*\*(.*?)\*\*\](?!\{\.underline\})"
    )
    pos = 0
    parts = []
    for m in token.finditer(text):
        if m.start() > pos:
            parts.append((text[pos : m.start()], False, False))
        if m.group(1) is not None:
            inner, bold, ul = m.group(1), True, True
        elif m.group(2) is not None:
            inner, bold, ul = m.group(2), False, True
        else:
            inner, bold, ul = m.group(3), True, False
        parts.append((inner, bold, ul))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False))
    return parts


def _other_line_has_style_markup(s):
    """是否为 read_cell_marked 编码的「其他」行（含加粗/下划线标记）。"""
    t = s or ""
    return "[**" in t or "[[" in t


def _other_marked_spans(text):
    """返回 read_cell_marked 已编码区间 [(start, end), …]，避免对标记内文字二次包裹。"""
    spans = []
    for m in re.finditer(
        r"\[\*\*(.*?)\*\*\](?:\{\.underline\})?|"
        r"\[\[(.*?)\]\](?:\{\.underline\})?|"
        r"\[\*\*(.*?)\*\*\](?!\{\.underline\})",
        text or "",
    ):
        spans.append((m.start(), m.end()))
    return spans


def _other_inject_applicant_bold_markup(text, applicant_name):
    """
    「其他」栏：未加 [**] 标记的本人姓名包为 [**姓名**]，供 parse_marked_authors 加粗。
    已标记或已在标记内的片段不重复包裹。
    """
    name = (applicant_name or "").strip()
    if not name or len(name) < 2:
        return text
    s = text or ""
    if name not in s:
        return s
    protected = _other_marked_spans(s)
    esc = re.escape(name)
    out = []
    pos = 0
    for m in re.finditer(esc, s):
        out.append(s[pos : m.start()])
        inside = any(a <= m.start() and m.end() <= b for a, b in protected)
        if inside:
            out.append(m.group(0))
        else:
            out.append(f"[**{m.group(0)}**]")
        pos = m.end()
    out.append(s[pos:])
    return "".join(out)


def _other_append_line_runs(paragraph, text, applicant_name, *, default_bold=False):
    """写入「其他」栏一行：保留原有标记，并为未标记的本人姓名加粗。"""
    prepared = _other_inject_applicant_bold_markup(text, applicant_name)
    if _other_line_has_style_markup(prepared):
        for seg, b, u in parse_marked_authors(prepared):
            append_run(
                paragraph,
                seg,
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=10.5,
                bold=b,
                underline=u,
            )
    else:
        append_run(
            paragraph,
            prepared.strip(),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            bold=default_bold,
        )


def _other_plain_text_for_subtitle_detect(line):
    """去掉样式标记后用于判断是否小节标题行。"""
    return "".join(x[0] for x in parse_marked_authors(line)).strip()


def _other_fragment_already_numbered(s):
    """片段已以「（1）」「(1)」等编号开头时，不再自动加「（n）」。"""
    t = (s or "").strip()
    return bool(re.match(r"^[（(]\s*\d+\s*[）)]\s*", t))


# 起止年月：闭区间、2023.03-至今、或开放式结束（如 2022.07- ）；「至今」可与前段无空格
_DATE_SEP = r"[-～–—－]"        # 横线 / 全角波浪线 / en-dash / em-dash / 全角连字号(U+FF0D)
_DATE_EST = r"(?:预计|约|暂定)?\s*"   # 结束日期前可选的估计前缀
_DATE_RANGE = (
    rf"[0-9]{{4}}\.[0-9]{{1,2}}\s*{_DATE_SEP}\s*{_DATE_EST}[0-9]{{4}}\.[0-9]{{1,2}}(?:毕业|出站|结束)?|"
    r"[0-9]{4}\.[0-9]{1,2}[-\s]?至今|"
    rf"[0-9]{{4}}\.[0-9]{{1,2}}\s*{_DATE_SEP}\s*"
)


def _norm_date(dr):
    """规范化日期段：统一分隔符为 -，去多余空格，保证「至今」前有空格。
    「预计2025.07毕业」→「2025.07（预计）」"""
    dr = normalize_experience_date_jin(dr)
    dr = re.sub(r"\s+", " ", dr.strip())
    dr = re.sub(rf"\s*{_DATE_SEP}\s*", "-", dr)   # 各种分隔符统一为 -
    # 「预计/约/暂定YYYY.MM毕业/出站/结束」→「YYYY.MM（预计）」
    dr = re.sub(
        r"-(预计|约|暂定)\s*([0-9]{4}\.[0-9]{1,2})(?:毕业|出站|结束)?$",
        r"-\2（预计）",
        dr,
    )
    dr = dr.replace("-至今", " 至今")
    dr = re.sub(r"([0-9])至今", r"\1 至今", dr)
    return dr


def split_experience_line(line):
    """
    将经历文本拆分为 4 列：前缀/时间、日期段、学校(或单位)、专业(或说明)。
    学习经历：`1.本科:2013.09-2017.06 中国民航大学 专业`（标签含冒号）。
    工作经历：`1.2023.03-至今 四川大学 博士后…`（序号后直接日期）；或
    `1.研究助理:2017.06-2018.06 香港科技大学深圳研究所`（职务在冒号前：第一格为序号+日期，第二格单位，第三格职务）。
    """
    s = normalize_experience_date_jin(str(line or "").strip())
    if s in ("无", "/"):
        return [s, "", "", ""]

    _DR_FULL = (
        rf"[0-9]{{4}}\.[0-9]{{1,2}}\s*{_DATE_SEP}\s*{_DATE_EST}[0-9]{{4}}\.[0-9]{{1,2}}(?:毕业|出站|结束)?"
        rf"|[0-9]{{4}}\.[0-9]{{1,2}}[-\s]?至今(?:[（(][^）)]*[）)][^\s]*)?"
    )
    m = re.match(rf"^(.*?:)\s*({_DR_FULL})\s+(.+?)\s+(.+)$", s)
    if m:
        return [m.group(1), _norm_date(m.group(2)), m.group(3), m.group(4)]

    # 学习经历「序号.学位 日期 学校 专业」（无冒号）：第一格序号+学位，后三格日期/学校/专业
    m_edu_nc = re.match(rf"^(\d+\.)\s*([^\d\s]{{1,6}})\s+({_DATE_RANGE})\s+(.+?)\s+(.+)$", s)
    if m_edu_nc:
        label = m_edu_nc.group(2).strip()
        dr = _norm_date(m_edu_nc.group(3))
        return [f"{m_edu_nc.group(1)}{label}:", dr, m_edu_nc.group(4).strip(), m_edu_nc.group(5).strip()]

    # 工作经历「序号.职务:起止年月 单位 …」：第一格「序号+日期」，第二格单位，第三格职务+补充
    m_job = re.match(rf"^(\d+\.)\s*(.+?):\s*({_DATE_RANGE})\s*(.+)$", s)
    if m_job:
        seq = m_job.group(1).rstrip()
        title = m_job.group(2).strip()
        dr = _norm_date(m_job.group(3))
        tail = m_job.group(4).strip()
        first_cell = f"{seq}{dr}"
        segs = tail.split(None, 1)
        org = segs[0]
        extra = segs[1] if len(segs) > 1 else ""
        desc = title if not extra else f"{title} {extra}".strip()
        return [first_cell, org, desc, ""]

    # 工作经历：第一格为「序号.+起止年月」，后两格为单位、职务/说明
    m_w = re.match(rf"^(\d+\.)\s*({_DATE_RANGE})\s*(.+)$", s)
    if m_w:
        first_cell = f"{m_w.group(1).rstrip()}{_norm_date(m_w.group(2))}"
        tail = m_w.group(3).strip()
        segs = tail.split(None, 1)
        if len(segs) == 1:
            return [first_cell, segs[0], "", ""]
        return [first_cell, segs[0], segs[1], ""]

    # 序号后无点号：「2 2023.10-2024.10 单位 职务」→ 补全点号后同上处理
    m_w_nodot = re.match(rf"^(\d+)\s+({_DATE_RANGE})\s*(.+)$", s)
    if m_w_nodot:
        first_cell = f"{m_w_nodot.group(1)}.{_norm_date(m_w_nodot.group(2))}"
        tail = m_w_nodot.group(3).strip()
        segs = tail.split(None, 1)
        if len(segs) == 1:
            return [first_cell, segs[0], "", ""]
        return [first_cell, segs[0], segs[1], ""]

    # 无标号续行（双学位/辅修/双校等）：裸「日期 学校 专业」，第一格留空
    m_cont = re.match(rf"^({_DR_FULL})\s+(.+?)\s+(.+)$", s)
    if m_cont:
        return ["", _norm_date(m_cont.group(1)), m_cont.group(2).strip(), m_cont.group(3).strip()]

    return [s, "", "", ""]


def split_overseas_line(line):
    """
    海外经历拆 4 列（版式四格与学习经历表对齐）：
    第 1 格：序号与起止年月同一格，如 1.2023.11-2024.12；
    第 2、3 格：日期后正文按第一个空白拆成两段（或「序号.博士:日期 学校 专业」：第二格学校，第三格类型+说明）；
    第 4 格：空。
    日期与后文可无空格（如 2022.11英国华威大学）；日期段允许 -至今、开放式 2022.07-。
    """
    s = normalize_experience_date_jin(str(line or "").strip())
    if s in ("无", "/"):
        return [s, "", "", ""]

    m_lab = re.match(rf"^(\d+\.)\s*(.+?):\s*({_DATE_RANGE})\s*(.*)$", s)
    if m_lab:
        seq = m_lab.group(1).rstrip()
        label = m_lab.group(2).strip()
        dr = _norm_date(m_lab.group(3))
        tail = (m_lab.group(4) or "").strip()
        first_cell = f"{seq}{dr}"
        if not tail:
            return [first_cell, "", label, ""]
        segs = tail.split(None, 1)
        org = segs[0]
        extra = segs[1] if len(segs) > 1 else ""
        col3 = label if not extra else f"{label} {extra}".strip()
        return [first_cell, org, col3, ""]

    # 「序号.标签 日期 机构 说明」（标签无冒号）：补全冒号后按学历行处理
    m_lab_nc = re.match(rf"^(\d+\.)\s*([^\d\s]{{1,8}})\s+({_DATE_RANGE})\s*(.*)$", s)
    if m_lab_nc:
        seq = m_lab_nc.group(1).rstrip()
        label = m_lab_nc.group(2).strip()
        dr = _norm_date(m_lab_nc.group(3))
        first_cell = f"{seq}{dr}"
        tail = (m_lab_nc.group(4) or "").strip()
        if not tail:
            return [first_cell, "", label, ""]
        segs = tail.split(None, 1)
        org = segs[0]
        extra = segs[1] if len(segs) > 1 else ""
        col3 = label if not extra else f"{label} {extra}".strip()
        return [first_cell, org, col3, ""]

    m = re.match(rf"^(\d+\.)\s*({_DATE_RANGE})(.*)$", s)
    if not m:
        # 序号后无点号：「2 2023.10-2024.10 单位 …」→ 补全点号
        m_nodot = re.match(rf"^(\d+)\s+({_DATE_RANGE})(.*)$", s)
        if m_nodot:
            first_cell = f"{m_nodot.group(1)}.{_norm_date(m_nodot.group(2))}"
            tail = (m_nodot.group(3) or "").strip()
            if not tail:
                return [first_cell, "", "", ""]
            segs = tail.split(None, 1)
            if len(segs) == 1:
                return [first_cell, segs[0], "", ""]
            return [first_cell, segs[0], segs[1], ""]
        # 裸日期（无序号）：「date 机构 说明」，序号格留空
        m_bare = re.match(rf"^({_DATE_RANGE})[:\s]\s*(.+)$", s)
        if m_bare:
            dr = _norm_date(m_bare.group(1))
            tail = m_bare.group(2).strip()
            segs = tail.split(None, 1)
            if len(segs) == 1:
                return [dr, segs[0], "", ""]
            return [dr, segs[0], segs[1], ""]
        return [s, "", "", ""]

    first_cell = f"{m.group(1).rstrip()}{_norm_date(m.group(2))}"
    tail = (m.group(3) or "").strip()
    if not tail:
        return [first_cell, "", "", ""]
    segs = tail.split(None, 1)
    if len(segs) == 1:
        return [first_cell, segs[0], "", ""]
    return [first_cell, segs[0], segs[1], ""]


_EDU_LINE_FULL = re.compile(
    rf"^(.+?:)\s*([0-9]{{4}}\.[0-9]{{1,2}}\s*{_DATE_SEP}\s*[0-9]{{4}}\.[0-9]{{1,2}}|[0-9]{{4}}\.[0-9]{{1,2}}[-\s]?至今(?:[（(][^）)]*[）)][^\s]*)?)\s+(.+?)\s+(.+)$"
)


def _normalize_piped_table_text(s, joiner="、"):
    """将抽取时的 | 合并为可读分隔符；joiner 为空则直接拼接（拟加入科研团队常见）。"""
    s = (s or "").strip()
    if not s or "|" not in s:
        return s
    parts = [p.strip() for p in s.split("|") if p.strip()]
    return joiner.join(parts)


def _format_project_institution_display(text):
    """
    科研项目「依托单位」窄列排版（接近正确004/005）：纯汉字且去空格后总长为 4，
    则在第 2 个字后换行；总长为 6，则在第 3 个字后换行。含英文、数字或非汉字字符时不处理。
    """
    raw = (text or "").strip()
    if not raw:
        return ""
    s = raw.replace("|", "").strip()
    if not s:
        return ""
    if re.search(r"[0-9A-Za-z]", s):
        return raw
    compact = re.sub(r"[\s\u3000]+", "", s)
    if not compact:
        return raw
    for ch in compact:
        if not (
            "\u4e00" <= ch <= "\u9fff"
            or "\u3400" <= ch <= "\u4dbf"
            or ch == "\u3007"
        ):
            return raw
    n = len(compact)
    if n == 4:
        return compact[:2] + "\n" + compact[2:]
    if n == 6:
        return compact[:3] + "\n" + compact[3:]
    return s


def _fill_project_institution_cell(cell, raw_institution):
    """写入科研项目依托单位：必要时四字名第二字后、六字名第三字后换行，段居中。"""
    disp = _format_project_institution_display(raw_institution)
    clear_cell(cell)
    if not str(disp).strip():
        _set_cell_text(
            cell,
            "",
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(cell)
        return
    parts = [ln.strip() for ln in disp.split("\n")]
    parts = [x for x in parts if x]
    if not parts:
        _set_cell_text(
            cell,
            "",
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(cell)
        return
    for idx, line in enumerate(parts):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        clear_paragraph_runs(para)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _tight_paragraph(para)
        append_run(
            para,
            line,
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
        )
    _paper_cell_center(cell)


def _load_985211_name_set(base_dir: Path):
    """读取同目录下 985211名单.txt（一行一所，# 行为注释）；无文件或为空则返回 None。"""
    p = base_dir / "985211名单.txt"
    if not p.is_file():
        return None
    out = set()
    for raw in p.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line and not line.startswith("#"):
            out.add(line)
    return out if out else None


# 国别/地区 + 学校：校名中若含以下片段，视为境外院校，不做「非 985211」的加粗+倾斜+下划线
#（与 985211名单 并列的豁免条件；例：英国华威大学、美国斯坦福大学、香港大学）
_FOREIGN_SCHOOL_SUBSTRINGS = (
    "英国",
    "美国",
    "加拿大",
    "澳大利亚",
    "澳洲",
    "德国",
    "法国",
    "日本",
    "韩国",
    "新加坡",
    "俄罗斯",
    "意大利",
    "西班牙",
    "瑞士",
    "荷兰",
    "比利时",
    "新西兰",
    "爱尔兰",
    "瑞典",
    "挪威",
    "丹麦",
    "芬兰",
    "奥地利",
    "波兰",
    "捷克",
    "匈牙利",
    "香港",
    "澳门",
    "台湾",
)


def _is_foreign_school_no_warn(school: str) -> bool:
    """
    True：境外院校常见写法，不因不在 985211 名单而套预警样式。
    - 校名字符串中含国别/地区关键词（如「英国」「美国」「香港」…）；
    - 或整段无汉字、仅有外文校名（例 University of Oxford）。
    """
    t = (school or "").strip()
    if not t:
        return False
    for k in _FOREIGN_SCHOOL_SUBSTRINGS:
        if k in t:
            return True
    if not re.search(r"[\u4e00-\u9fff]", t) and re.search(r"[A-Za-z]", t):
        return True
    return False


_CAS_ACADEMY_KEY = "中国科学院"


def _is_cas_academy_institute(school: str) -> bool:
    """中国科学院系统（含各所）不因不在名单而套「非 985211」样式。"""
    return bool((school or "").strip()) and _CAS_ACADEMY_KEY in str(school)


def _append_education_line_runs(paragraph, line, name_set):
    """学习经历一行：学校列片段若不在 985211 名单则加粗、倾斜、下划线。"""
    s = str(line or "").strip()
    m = _EDU_LINE_FULL.match(s)
    if not m or not name_set:
        append_run(paragraph, s, font_cn="仿宋", font_en="Times New Roman", size_pt=12)
        return
    school = m.group(3).strip()
    warn = (
        bool(school)
        and school not in name_set
        and not _is_foreign_school_no_warn(school)
        and not _is_cas_academy_institute(school)
    )
    i0, i1 = m.start(3), m.end(3)
    pre, mid, post = s[:i0], s[i0:i1], s[i1:]
    append_run(paragraph, pre, font_cn="仿宋", font_en="Times New Roman", size_pt=12)
    append_run(
        paragraph,
        mid,
        font_cn="仿宋",
        font_en="Times New Roman",
        size_pt=12,
        bold=warn,
        italic=warn,
        underline=warn,
    )
    append_run(paragraph, post, font_cn="仿宋", font_en="Times New Roman", size_pt=12)


def write_cell_keep_style(cell, text, *, font_cn="仿宋", font_en="Times New Roman", size_pt=12, bold=False, underline=False, italic=False):
    """保留单元格段落/表格结构，仅替换首段 runs 文本。"""
    if not cell.paragraphs:
        cell.add_paragraph("")
    p = cell.paragraphs[0]
    clear_paragraph_runs(p)
    append_run(
        p,
        str(text),
        font_cn=font_cn,
        font_en=font_en,
        size_pt=size_pt,
        bold=bold,
        underline=underline,
        italic=italic,
    )


def write_cell_keep_style_decimal_hidden(cell, text, *, font_cn="仿宋", font_en="Times New Roman", size_pt=12, bold=False):
    """含论文统计数字时：仅「N.0」类对齐用小数将「.0」白色；2.5 等真实小数整段正常色。"""
    if not cell.paragraphs:
        cell.add_paragraph("")
    p = cell.paragraphs[0]
    clear_paragraph_runs(p)
    _append_stats_detail_text_runs(p, str(text), font_cn=font_cn, font_en=font_en, size_pt=size_pt, bold=bold)


def write_cell_school_with_985211_highlight(cell, school_text, name_set, *, font_cn, font_en, size_pt):
    """学习经历「学校」列：不在 985211名单 中的校名加粗、倾斜、下划线。"""
    if not name_set:
        write_cell_keep_style(cell, school_text, font_cn=font_cn, font_en=font_en, size_pt=size_pt)
        return
    t = (school_text or "").strip()
    warn = (
        bool(t)
        and t not in name_set
        and not _is_foreign_school_no_warn(t)
        and not _is_cas_academy_institute(t)
    )
    write_cell_keep_style(
        cell,
        school_text,
        font_cn=font_cn,
        font_en=font_en,
        size_pt=size_pt,
        bold=warn,
        italic=warn,
        underline=warn,
    )


def _is_overseas_supplementary_note(line):
    """
    海外经历附注：不入嵌套表，写在海外表下方同一大单元格内、表格外单独成段。
    - 「注:」「注：」等引导的说明（如 注:2022.12-2023.02 办理博士后入站手续）；
    - 「（相关经历不连续的请补充说明:…）」等补充说明类。
    """
    s = (line or "").strip()
    if not s:
        return False
    # 显式「注:」/「注：」行，与带序号的经历正文明区分
    if re.match(r"^[\s\u3000]*注\s*[：:]", s):
        return True
    if "补充说明" in s:
        return True
    # 全角括号引导的补充句（与序号正文「1.博士:…」区分）
    if s.startswith(("（", "(")) and ("补充" in s or "不连续" in s):
        return True
    if "须完整" in s or "须连续" in s:
        return True
    if re.search(r"学习及工作经历连续|经历连续", s):
        return True
    if re.search(r"临时工作", s):
        return True
    return False


def _first_paragraph_after_nested_table(cell, nested_table):
    """返回嵌套表之后第一个 w:p；若无则在表后新建一段。"""
    tc = cell._tc
    tbl_el = nested_table._tbl
    children = list(tc)
    for i, el in enumerate(children):
        if el is tbl_el:
            for j in range(i + 1, len(children)):
                if children[j].tag == qn("w:p"):
                    return Paragraph(children[j], cell)
            break
    new_el = OxmlElement("w:p")
    tbl_el.addnext(new_el)
    return Paragraph(new_el, cell)


def _append_overseas_notes_below_table(cell, overseas_tbl, note_lines):
    """在海外嵌套表下方的单元格段落中写入补充说明（不入表）；表与「注」等附注之间空一行。"""
    note_lines = [x.strip() for x in (note_lines or []) if x and str(x).strip()]
    if not note_lines:
        return
    first_p = _first_paragraph_after_nested_table(cell, overseas_tbl)
    blank_el = OxmlElement("w:p")
    first_p._p.addprevious(blank_el)
    p_blank = Paragraph(blank_el, cell)
    _tight_paragraph(p_blank)

    clear_paragraph_runs(first_p)
    _tight_paragraph(first_p)
    append_run(
        first_p,
        note_lines[0],
        font_cn="仿宋",
        font_en="Times New Roman",
        size_pt=12,
    )
    prev_el = first_p._p
    for line in note_lines[1:]:
        new_el = OxmlElement("w:p")
        prev_el.addnext(new_el)
        prev_el = new_el
        p2 = Paragraph(new_el, cell)
        _tight_paragraph(p2)
        append_run(
            p2,
            line,
            font_cn="仿宋",
            font_en="Times New Roman",
            size_pt=12,
        )


def _partition_experience_notes_from_lines(lines):
    """从学习/工作/海外行列表中拆出附注（兼容未写入 experience_notes 的旧 JSON）。"""
    rest, notes = [], []
    for x in lines or []:
        s = str(x).strip()
        if not s:
            continue
        if _is_overseas_supplementary_note(s):
            notes.append(s)
        else:
            rest.append(s)
    return rest, notes


def fill_learning_nested_tables(
    cell, edu_lines, work_lines, overseas_lines, name_set=None, experience_notes=None
):
    """按调整空白版中内嵌表格结构填充学习/工作/海外经历。
    experience_notes：抽取得到的「注：…」等，与海外表格外 o_notes 合并后写于海外嵌套表下方最后。
    """
    if len(cell.tables) < 3:
        return False

    edu_tbl, work_tbl, overseas_tbl = cell.tables[0], cell.tables[1], cell.tables[2]

    # 学习经历表：填写有数据的行，多余行直接删除（避免空行导致段间距翻倍）
    n_edu = max(len(edu_lines), 1)
    while len(edu_tbl.rows) < n_edu:
        _insert_blank_rows_after(edu_tbl, len(edu_tbl.rows) - 1, 1)
    for i in range(min(n_edu, len(edu_tbl.rows))):
        parts = split_experience_line(edu_lines[i] if i < len(edu_lines) else "")
        for j in range(min(4, len(edu_tbl.rows[i].cells))):
            if j == 2:
                write_cell_school_with_985211_highlight(
                    edu_tbl.rows[i].cells[j],
                    parts[j],
                    name_set,
                    font_cn="仿宋",
                    font_en="Times New Roman",
                    size_pt=12,
                )
            else:
                write_cell_keep_style(
                    edu_tbl.rows[i].cells[j],
                    parts[j],
                    font_cn="仿宋",
                    font_en="Times New Roman",
                    size_pt=12,
                )
    # 删除超出数据行数的模板行
    while len(edu_tbl.rows) > n_edu:
        tr = edu_tbl.rows[-1]._tr
        tr.getparent().remove(tr)

    # 工作经历表（多行 x4）：每行一条；模板默认 1 行，不足则克隆插入，多余则删行
    work_items = [str(x).strip() for x in (work_lines or []) if str(x).strip()]
    if not work_items:
        work_items = ["无"]
    n_work = max(len(work_items), 1)
    while len(work_tbl.rows) < n_work:
        _insert_blank_rows_after(work_tbl, len(work_tbl.rows) - 1, 1)
    while len(work_tbl.rows) > n_work:
        tr = work_tbl.rows[-1]._tr
        tr.getparent().remove(tr)
    for i in range(n_work):
        w_parts = split_experience_line(work_items[i] if i < len(work_items) else "")
        for j in range(min(4, len(work_tbl.rows[i].cells))):
            write_cell_keep_style(
                work_tbl.rows[i].cells[j],
                w_parts[j],
                font_cn="仿宋",
                font_en="Times New Roman",
                size_pt=12,
            )

    # 海外经历表（多行 x4）：序号正文逐行入表；「补充说明」类附注不入表，写在嵌套表下方段落
    o_list = [str(x).strip() for x in (overseas_lines or []) if str(x).strip()]
    o_notes = [x for x in o_list if _is_overseas_supplementary_note(x)]
    o_table_items = [x for x in o_list if not _is_overseas_supplementary_note(x)]
    if not o_table_items:
        o_table_items = ["无"]
    n_os = max(len(o_table_items), 1)
    while len(overseas_tbl.rows) < n_os:
        _insert_blank_rows_after(overseas_tbl, len(overseas_tbl.rows) - 1, 1)
    while len(overseas_tbl.rows) > n_os:
        tr = overseas_tbl.rows[-1]._tr
        tr.getparent().remove(tr)
    for i in range(n_os):
        o_parts = split_overseas_line(o_table_items[i] if i < len(o_table_items) else "")
        for j in range(min(4, len(overseas_tbl.rows[i].cells))):
            write_cell_keep_style(
                overseas_tbl.rows[i].cells[j],
                o_parts[j],
                font_cn="仿宋",
                font_en="Times New Roman",
                size_pt=12,
            )
    extra = [str(x).strip() for x in (experience_notes or []) if str(x).strip()]
    _append_overseas_notes_below_table(cell, overseas_tbl, o_notes + extra)

    return True


def _tight_paragraph(paragraph):
    """与正确001接近：段前段后零间距、单倍行距，减少「空行」观感。"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0


def _table_cell_vertical_center_paragraphs_left(cell):
    """表格正文格：垂直居中 + 各段水平左对齐（思想政治、岗位任务与学院支撑等）。"""
    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    except Exception:
        pass
    for p in cell.paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _ps_n(ps, key):
    try:
        v = float(ps.get(key, 0))
        return int(v) if v == int(v) else v
    except Exception:
        return 0


def _ps_first_raw(ps, fk):
    """Return 0 if first-author count was blank in source, else numeric value."""
    raw = ps.get(fk, 0)
    if raw == "":
        return 0
    return _ps_n(ps, fk)


# papers_summary 中数值字段（与 batch_extract_json SCHEMA 一致）；用于判断是否出现「半篇」等非整数
_PAPER_SUMMARY_NUM_KEYS = (
    "total",
    "jcr_q1",
    "jcr_q1_first",
    "jcr_q2",
    "jcr_q2_first",
    "jcr_q3",
    "jcr_q3_first",
    "jcr_q4",
    "jcr_q4_first",
    "jcr_other",
    "jcr_other_first",
    "school_a",
    "school_a_first",
    "school_b",
    "school_b_first",
    "school_c",
    "school_c_first",
    "school_a_star",
    "school_a_star_first",
    "school_other",
    "school_other_first",
    "ccf_a",
    "ccf_a_first",
    "ccf_b",
    "ccf_b_first",
    "ccf_c",
    "ccf_c_first",
    "ccf_other",
    "ccf_other_first",
    "abs_4",
    "abs_4_first",
    "abs_3",
    "abs_3_first",
    "abs_2",
    "abs_2_first",
    "abs_1",
    "abs_1_first",
    "abs_other",
    "abs_other_first",
)


def _papers_summary_has_fraction(ps):
    """任一篇数为非整数（如 2.5、折合半篇）时，统计区内整数亦显示为一位小数（如 1.0）以对齐。"""
    if not ps:
        return False
    for k in _PAPER_SUMMARY_NUM_KEYS:
        if k not in ps:
            continue
        try:
            v = float(ps[k])
        except (TypeError, ValueError):
            continue
        if abs(v - round(v)) > 1e-9:
            return True
    return False


def _stats_display_num(val, align_decimal):
    """align_decimal 时统一一位小数；否则整数不带 .0。"""
    try:
        v = float(val)
    except (TypeError, ValueError):
        return str(val)
    if align_decimal:
        return f"{v:.1f}"
    if abs(v - round(v)) < 1e-9:
        return str(int(round(v)))
    return str(v)


def _jcr_other_pair(ps):
    """
    JCR「其他」两行数字：仅以 papers_summary.jcr_other / jcr_other_first 为准（抽取或手工填写），不做推算。
    第一作者数为空时返回 0。
    """
    return max(_ps_n(ps, "jcr_other"), 0), _ps_first_raw(ps, "jcr_other_first")


def _paper_stats_jcr_parts(ps):
    """
    模式B：仅非零分区；JCR「其他」完全依赖 papers_summary 中的 jcr_other / jcr_other_first。
    无制表符的纯文本条（供兼容）；版面填充见 _paper_stats_jcr_segments_tabbed。
    """
    j_other, j_other_first = _jcr_other_pair(ps)
    dec = _papers_summary_has_fraction(ps)

    parts = []
    for lab, tk, fk in [
        ("Q1区", "jcr_q1", "jcr_q1_first"),
        ("Q2区", "jcr_q2", "jcr_q2_first"),
        ("Q3区", "jcr_q3", "jcr_q3_first"),
        ("Q4区", "jcr_q4", "jcr_q4_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_first_raw(ps, fk)
        if tn <= 0:
            continue
        parts.append(
            f"{lab}： {_stats_display_num(tn, dec)} 篇（ {_stats_display_num(fn, dec)} 篇为第一作者）"
        )
    if j_other > 0:
        parts.append(
            f"其他： {_stats_display_num(j_other, dec)} 篇（ {_stats_display_num(j_other_first, dec)} 篇为第一作者）"
        )
    return parts


def _fmt_jcr_tabbed(lab, tn, fn, ps):
    """标签后经制表符对齐篇数；「篇」与「（」紧接，避免多一个 \\t 在 Word 中拉成大片空白。"""
    dec = _papers_summary_has_fraction(ps)
    return f"{lab}：\t{_stats_display_num(tn, dec)} 篇（ {_stats_display_num(fn, dec)} 篇为第一作者）"


def _fmt_school_ccf_tabbed(lab, tn, fn, ps):
    """校级/CCF：「类 ：」后制表对齐篇数；篇与括号之间不用制表符。"""
    dec = _papers_summary_has_fraction(ps)
    return f"{lab} ：\t{_stats_display_num(tn, dec)} 篇（ {_stats_display_num(fn, dec)} 篇为第一作者）"


def _paper_stats_jcr_has_content(ps):
    """JCR 分区统计是否有任一非零条目（含「其他」）。"""
    return bool(_paper_stats_jcr_segments_tabbed(ps))


def _paper_stats_jcr_segments_tabbed(ps):
    """JCR 明细条列表（标签后单制表符），顺序 Q1～Q4、其他（仅非零）。"""
    j_other, j_other_first = _jcr_other_pair(ps)
    segs = []
    for lab, tk, fk in [
        ("Q1区", "jcr_q1", "jcr_q1_first"),
        ("Q2区", "jcr_q2", "jcr_q2_first"),
        ("Q3区", "jcr_q3", "jcr_q3_first"),
        ("Q4区", "jcr_q4", "jcr_q4_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_first_raw(ps, fk)
        if tn <= 0:
            continue
        segs.append(_fmt_jcr_tabbed(lab, tn, fn, ps))
    if j_other > 0:
        segs.append(_fmt_jcr_tabbed("其他", j_other, j_other_first, ps))
    return segs


def _paper_stats_school_segments_tabbed(ps):
    segs = []
    for lab, tk, fk in [
        ("A类", "school_a", "school_a_first"),
        ("B类", "school_b", "school_b_first"),
        ("C类", "school_c", "school_c_first"),
        ("A*类", "school_a_star", "school_a_star_first"),
        ("其他", "school_other", "school_other_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_first_raw(ps, fk)
        if tn <= 0:
            continue
        segs.append(_fmt_school_ccf_tabbed(lab, tn, fn, ps))
    return segs


def _paper_stats_ccf_segments_tabbed(ps):
    segs = []
    for lab, tk, fk in [
        ("A类", "ccf_a", "ccf_a_first"),
        ("B类", "ccf_b", "ccf_b_first"),
        ("C类", "ccf_c", "ccf_c_first"),
        ("其他", "ccf_other", "ccf_other_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_first_raw(ps, fk)
        if tn <= 0:
            continue
        segs.append(_fmt_school_ccf_tabbed(lab, tn, fn, ps))
    return segs


def _fmt_abs_tabbed(lab, tn, fn, ps):
    """ABS 分级：制表仅用于标签后对齐篇数；篇与括号紧接。"""
    dec = _papers_summary_has_fraction(ps)
    return f"{lab}：\t{_stats_display_num(tn, dec)} 篇（ {_stats_display_num(fn, dec)} 篇为第一作者）"


def _paper_stats_abs_segments_tabbed(ps):
    segs = []
    for lab, tk, fk in [
        ("ABS4", "abs_4", "abs_4_first"),
        ("ABS3", "abs_3", "abs_3_first"),
        ("ABS2", "abs_2", "abs_2_first"),
        ("ABS1", "abs_1", "abs_1_first"),
        ("其他", "abs_other", "abs_other_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_n(ps, fk)
        if tn <= 0:
            continue
        segs.append(_fmt_abs_tabbed(lab, tn, fn, ps))
    return segs


def _stats_join_n_per_line(segments, n=2):
    """每行最多 n 条（分号拼接），条数超出则换行（单段内 \\n，便于同一段落共用制表位）。"""
    if not segments:
        return ""
    lines = ["；".join(segments[i : i + n]) for i in range(0, len(segments), n)]
    return "\n".join(lines)


def _set_stat_detail_tab_stops(paragraph):
    """
    每条仅「标签：\\t篇数 篇（…）」一个制表符；一行最多 2 条则最多 2 个 \\t。
    不再为「篇」与括号之间设制表位，避免出现大块空白。
    """
    ts = paragraph.paragraph_format.tab_stops
    ts.add_tab_stop(Pt(52), WD_TAB_ALIGNMENT.LEFT)
    ts.add_tab_stop(Pt(230), WD_TAB_ALIGNMENT.LEFT)


def _paper_stats_jcr_line(ps):
    return "；".join(_paper_stats_jcr_parts(ps))


def _paper_stats_school_parts(ps):
    """与正确 Word 一致：「A类 ： 3 篇」类与冒号之间有一空格。"""
    dec = _papers_summary_has_fraction(ps)
    parts = []
    for lab, tk, fk in [
        ("A类", "school_a", "school_a_first"),
        ("B类", "school_b", "school_b_first"),
        ("C类", "school_c", "school_c_first"),
        ("A*类", "school_a_star", "school_a_star_first"),
        ("其他", "school_other", "school_other_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_n(ps, fk)
        if tn <= 0:
            continue
        parts.append(
            f"{lab} ： {_stats_display_num(tn, dec)} 篇（ {_stats_display_num(fn, dec)} 篇为第一作者）"
        )
    return parts


def _paper_stats_school_line(ps):
    return "；".join(_paper_stats_school_parts(ps))


def _paper_stats_ccf_parts(ps):
    """
    中国计算机学会推荐目录（CCF）：papers_summary 中选填非零项即可输出第四统计块。
    字段：ccf_a / ccf_b / ccf_c / ccf_other 及 ccf_*_first（与校级统计写法一致；不需则全 0 或不写）。
    """
    parts = []
    dec = _papers_summary_has_fraction(ps)
    for lab, tk, fk in [
        ("A类", "ccf_a", "ccf_a_first"),
        ("B类", "ccf_b", "ccf_b_first"),
        ("C类", "ccf_c", "ccf_c_first"),
        ("其他", "ccf_other", "ccf_other_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_n(ps, fk)
        if tn <= 0:
            continue
        parts.append(
            f"{lab} ： {_stats_display_num(tn, dec)} 篇（ {_stats_display_num(fn, dec)} 篇为第一作者）"
        )
    return parts


def _paper_stats_ccf_line(ps):
    return "；".join(_paper_stats_ccf_parts(ps))


def _paper_stats_abs_parts(ps):
    parts = []
    dec = _papers_summary_has_fraction(ps)
    for lab, tk, fk in [
        ("ABS4", "abs_4", "abs_4_first"),
        ("ABS3", "abs_3", "abs_3_first"),
        ("ABS2", "abs_2", "abs_2_first"),
        ("ABS1", "abs_1", "abs_1_first"),
        ("其他", "abs_other", "abs_other_first"),
    ]:
        tn, fn = _ps_n(ps, tk), _ps_n(ps, fk)
        if tn <= 0:
            continue
        parts.append(
            f"{lab}： {_stats_display_num(tn, dec)} 篇（ {_stats_display_num(fn, dec)} 篇为第一作者）"
        )
    return parts


def _paper_stats_abs_line(ps):
    return "；".join(_paper_stats_abs_parts(ps))


def _paper_marked_lead_author(authors):
    """作者串以 [**…**] 标记开头时视为第一作者/通讯作者署名（与 JSON 约定一致）。"""
    s = (authors or "").lstrip()
    return bool(s.startswith("[**"))


def _journal_display_parts(p):
    """
    期刊单元格展示用字段。
    - journal 可合并写「浙江大学学报(人文社会科学版)」，自动拆成刊名 + (人文社会科学版)；
      也可分字段 journal + journal_subtitle（后者优先）。
    - volume 形如 55(03) 且存在 pages 时，若无 volume_pages_line，则生成「55(03)：60-75」一行（中文期刊版式）。
    """
    j_main = str(p.get("journal") or "").strip()
    j_sub = str(p.get("journal_subtitle") or "").strip()
    if not j_sub:
        m = re.match(r"^(.+?)(\([^)]+\))\s*$", j_main)
        if m:
            j_main, j_sub = m.group(1).strip(), m.group(2).strip()
    vol = str(p.get("volume") or "").strip()
    pgs = str(p.get("pages") or "").strip()
    vp_line = str(p.get("volume_pages_line") or "").strip()
    if not vp_line and vol and pgs and "(" in vol:
        vp_line = f"{vol}：{pgs}"
    return j_main, j_sub, vp_line


def _is_paren_page_count_line(s):
    """「（共14页）」类页数说明，不是期刊副题名。"""
    t = (s or "").strip()
    return bool(re.match(r"^[（(]\s*共\s*\d+\s*页\s*[）)]\s*$", t))


def _journal_bold_title_only(journal_main_text, issn):
    """
    刊名单独加粗：若单元格把 ISSN/日期 写在同一串 journal 里，只取 ISSN 之前的刊名部分。
    """
    s = (journal_main_text or "").strip()
    issn = (issn or "").strip()
    if not s:
        return ""
    if issn and issn in s:
        m_paren = re.search(r"\(\s*" + re.escape(issn) + r"\s*\)", s)
        if m_paren:
            name = s[: m_paren.start()].strip().rstrip(" ,，")
            if name:
                return name
        name = s[: s.index(issn)].rstrip(" ,，").strip()
        if name:
            return name
    m = re.match(r"^(.+?),\s*\d{4}-\d{3}[\dXx]", s)
    if m:
        return m.group(1).strip().rstrip(" ,，")
    return s


def _journal_flat_pipe_for_parse(pub_raw):
    """去掉期刊单元格样式标记，按行拼成 parse_pub_info 可用的 | 分段串。"""
    text = (pub_raw or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = []
    for ln in text.split("\n"):
        plain = "".join(x[0] for x in parse_marked_authors(ln))
        t = plain.strip()
        if t:
            lines.append(t)
    return "|".join(lines)


def _journal_citation_flat_for_parse(p):
    """
    供 parse_pub_info 使用的扁平串：优先 pub_raw（去标记后按行 | 拼接）；
    若无 pub_raw 但 journal 整格即一条英文逗号引文，则用 journal 换行→|。
    """
    raw = str(p.get("pub_raw") or "").strip()
    if raw:
        return _journal_flat_pipe_for_parse(raw)
    j = str(p.get("journal") or "").strip()
    if not j:
        return ""
    return j.replace("\r\n", "\n").replace("\r", "\n").replace("\n", "|")


def _journal_flat_has_valid_issn(flat):
    if find_valid_issn_in_text:
        return bool(find_valid_issn_in_text(flat or ""))
    s = flat or ""
    p_page = re.search(r"(?i)\bp\.?\s*\d", s)
    p_page_start = p_page.start() if p_page else len(s) + 1
    for m in re.finditer(r"\d{4}-\d{3}[\dXx]", s, re.I):
        st, en = m.start(), m.end()
        if st >= p_page_start or (st > 0 and s[st - 1].isdigit()):
            continue
        if en < len(s) and s[en].isdigit():
            continue
        if re.search(
            rf",\s*{re.escape(m.group())}\s*[,，|\n]",
            s[max(0, st - 1) : min(len(s), en + 2)],
        ):
            return True
    return False


def _journal_flat_looks_like_comma_citation(flat):
    """英文期刊常见「…, ISSN, Jul. 2020, vol, pages」：应用逗号解析并强制多行结构化排版。"""
    if not flat or not _journal_flat_has_valid_issn(flat):
        return False
    return bool(
        re.search(
            r"\b(Jan(?:uary)?\.?|Feb(?:ruary)?\.?|Mar(?:ch)?\.?|Apr(?:il)?\.?|May\.?|Jun(?:e)?\.?|Jul(?:y)?\.?|Aug(?:ust)?\.?|Sep(?:t(?:ember)?)?\.?|Oct(?:ober)?\.?|Nov(?:ember)?\.?|Dec(?:ember)?\.?)\s*\d{4}\b",
            flat,
            re.I,
        )
    )


_APA_PAREN_ISSN_YEAR_RE = re.compile(
    r"\(\s*\d{4}-\d{3}[\dXx]\s*\)\s*\.?\s*(?:19|20)\d{2}\b",
    re.I,
)


def _journal_flat_looks_like_apa_paren_citation(flat):
    """心理学等 APA 单行：刊名 (ISSN). 年, 卷(期), 页码。"""
    if not flat or not re.search(r"\d{4}-\d{3}[\dXx]", flat):
        return False
    if _journal_flat_looks_like_comma_citation(flat):
        return False
    s = flat.replace("|", " ").replace("\n", " ")
    return bool(_APA_PAREN_ISSN_YEAR_RE.search(s))


def _journal_verbatim_display_text(p):
    """回填时按原文展示：优先保留 pub_raw 中的 [**] 标记与标点。"""
    raw = str(p.get("pub_raw") or "").strip()
    if raw:
        return raw.replace("\r\n", "\n").replace("\r", "\n")
    return str(p.get("journal") or "").strip()


def _journal_has_online_or_doi_note(text):
    """括号 online 附注或 DOI（含中文冒号 DOI：）不宜拆成多行结构化刊名。"""
    s = text or ""
    if re.search(r"\bDOI\s*[：:]", s, re.I):
        return True
    if re.search(r"[（(]\s*已?\s*online", s, re.I):
        return True
    return False


def _journal_should_fill_verbatim(p, flat):
    """
    无法可靠走结构化多行版式时，原样输出 pub_raw / journal，避免信息丢失或刊名被截断。
    """
    if _journal_flat_looks_like_apa_paren_citation(flat):
        return True
    src = _journal_verbatim_display_text(p)
    if not src.strip():
        return False
    if _journal_has_online_or_doi_note(src) or _journal_has_online_or_doi_note(flat):
        return True
    if _journal_flat_looks_like_comma_citation(flat):
        return False
    j = str(p.get("journal") or "").strip()
    if len(j) < 25:
        return False
    if not re.search(r"\d{4}-\d{3}[\dXx]", j):
        return False
    if not re.search(r"\b(19|20)\d{2}\b", j):
        return False
    date_empty = not str(p.get("date") or "").strip()
    pages_empty = not str(p.get("pages") or "").strip()
    if date_empty and pages_empty and "," in j:
        return True
    return False


def _fill_journal_cell_verbatim(journal_cell, p):
    """按 pub_raw / journal 原文逐行写入单元格（不拆字段、不改标点）。"""
    text = _journal_verbatim_display_text(p)
    if not text.strip():
        _fill_journal_cell_structured(journal_cell, p)
        return
    clear_cell(journal_cell)
    first = True
    for raw_ln in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_ln.rstrip("\r")
        if not line.strip():
            continue
        para = journal_cell.paragraphs[0] if first else journal_cell.add_paragraph()
        first = False
        _tight_paragraph(para)
        _paragraph_keep_lines_together(para)
        for seg, b, u in parse_marked_authors(line):
            append_run(
                para,
                seg,
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=10.5,
                bold=b,
                underline=u,
            )
    if first:
        _fill_journal_cell_structured(journal_cell, p)


def _paper_journal_effective_fields(p):
    """
    合并 papers[] 中结构化字段与 parse_pub_info（对 pub_raw / 整段 journal 扁平串）。
    - 对「ISSN + 英文月年」类逗号引文：用解析结果覆盖 journal/date/volume/pages 等（修正旧 JSON 整段塞在 journal 里、
      或仅合并空字段导致仍走原样 pub_raw 无法分行的问题）。
    - 其余：仍仅向空字段补齐。
    """
    base = dict(p)
    if not _parse_pub_info_impl:
        return base
    flat = _journal_citation_flat_for_parse(p)
    if not flat.strip():
        return base
    pi = _parse_pub_info_impl(flat)
    cite_like = _journal_flat_looks_like_comma_citation(flat)
    protected_sub = str(base.get("journal_subtitle") or "").strip()
    for k, v in pi.items():
        if k == "journal_subtitle":
            continue
        if v is None:
            continue
        vs = str(v).strip()
        if not vs:
            continue
        if cite_like and k in (
            "journal",
            "issn",
            "date",
            "volume",
            "issue",
            "pages",
            "page_count",
        ):
            base[k] = v
        elif not str(base.get(k) or "").strip():
            base[k] = v
    jst_pi = str(pi.get("journal_subtitle") or "").strip()
    if cite_like:
        base["journal_subtitle"] = jst_pi or protected_sub
    elif jst_pi and not str(base.get("journal_subtitle") or "").strip():
        base["journal_subtitle"] = jst_pi
    return base


def _journal_pub_raw_has_style_markup(pub_raw):
    """仅当 pub_raw 含 read_cell_marked 的加粗/下划线标记时，才整格原样还原。"""
    s = pub_raw or ""
    return "[**" in s or "[[" in s


def _journal_pages_s_strip_redundant_volume_line(pages_s, *, eng_comma, eng_vol_line):
    """
    英文逗号引文已单独输出 Volume/Issue 行时，若 pages 仍带整段「Volume:… p.…」则只保留页码，
    避免再出现一行「Volume:160, Issue:0, p.105300」。
    亦处理 Elsevier 等「Volume 230, p. 109828」（Volume 与数字间为空格、无冒号）与 eng_vol_line「Volume:230」重复（错误001）。
    """
    s = (pages_s or "").strip()
    if not s or not eng_comma or not eng_vol_line:
        return s
    if re.search(r"(?i)volume\s*[：:]", s):
        m = re.search(r"(?i)p\.?\s*([\d\-—]+)", s)
        if m:
            return m.group(1).strip()
        return s
    # 「Volume 230, p. …」：与已输出的 Volume:230 行重复，只保留页码数字供单独一行 p.
    if re.search(r"(?i)volume\s+\d+", s):
        m = re.search(r"(?i)p\.?\s*([\d\-—]+)", s)
        if m:
            return m.group(1).strip()
    return s


def _journal_pages_display_line(pages_s, *, no_space_after_p=False):
    """
    页码行：纯数字/起止页加 p. 前缀；Just Accepted 等原文保持。
    no_space_after_p：英文逗号引文版式用「p.105300」无空格（与 Volume: 行一致）。
    """
    s = (pages_s or "").strip()
    if not s:
        return ""
    if re.fullmatch(r"[\d\s\-—]+", s):
        if no_space_after_p:
            compact = re.sub(r"\s+", "", s)
            return f"p.{compact}"
        return f"p. {s}"
    return s


def _journal_issn_or_isbn_match_start(s):
    """ISSN 或 978-… ISBN 在串中的起始下标；无则 None（不含页码误匹配的假 ISSN）。"""
    text = s or ""
    issn = find_valid_issn_in_text(text) if find_valid_issn_in_text else ""
    if issn:
        m = re.search(re.escape(issn), text, re.I)
        if m:
            return m.start()
    m = re.search(
        r"978-\d{1,5}-\d{1,7}-\d{1,7}-\d{1,3}[Xx]?",
        text,
        re.I,
    )
    return m.start() if m else None


def _journal_split_trailing_paren_pages(s):
    """从尾部拆出「（共N页）」「（约N页）」「（注：…）」等括号说明。"""
    t = (s or "").strip()
    m = re.search(r"(\s*[（(](?:约|共)\s*\d+\s*页\s*[）)])\s*$", t)
    if m:
        return t[: m.start()].strip(), m.group(1).strip()
    # 注释类括号说明（如「（注：目前还没有期卷号等）」）
    m2 = re.search(r"(\s*,?\s*[（(]\s*注\s*[：:].*?[）)])\s*$", t, re.S)
    if m2:
        body = t[: m2.start()].strip().rstrip(",")
        note = m2.group(1).strip().lstrip(",").strip()
        return body, note
    return t, ""


def _journal_line_page_keyword_to_p_dot(s):
    """
    将「, Page 106649」「,Page 123」转为「, p. …」，与既有「, p.」卷期拆行规则一致；
    行首「Page …」转为「p. …」。
    Volume/Issue 后的裸页码范围（如 Issue:9, 12967-12980）自动补 p：前缀以便与卷期行分开。
    """
    t = (s or "").strip()
    if not t:
        return t
    t = re.sub(r"(?i),\s*Page\s*", ", p. ", t)
    t = re.sub(r"(?i)^Page\s*", "p. ", t)
    # 裸页码范围：Volume/Issue:N 后直接跟 XXXX-XXXX 格式页码（无 p. 前缀）→ 补 p：
    t = re.sub(
        r"(?i)((?:Issue|Vol(?:ume)?)\s*[：:]\s*\d+)\s*,\s{0,5}(\d{3,}\s*[-–—]\s*\d{3,})(?=\s*[,(（]|$)",
        r"\1, p：\2",
        t,
    )
    return t


def _journal_split_volume_line_before_p_pages(line):
    """
    「Volume:…,Issue:…, p.123」「…, p：12967-12980」或「…, Page 106649」类：在页码段前切开，便于页码单独成行。
    返回 (卷期段, 页码段)；无页码段时第二项为空字符串。
    """
    t = _journal_line_page_keyword_to_p_dot(line)
    if not t:
        return "", ""
    m = re.search(r"(?i),\s*(?=p[.：]?\s*[\d\-—])", t)
    if m:
        head = t[: m.start()].strip().rstrip(",")
        tail = t[m.end() :].strip()
        # 补点：p1-17 → p.1-17
        tail = re.sub(r"^p(?=[0-9])", "p.", tail, flags=re.I)
        return head, tail
    m2 = re.search(r"(?i)(?<=[,，\s])(p[.：]?\s*[\d\-—\s]+)$", t)
    if m2 and re.search(r"(?i)volume\s*[：:]", t):
        head = t[: m2.start()].strip().rstrip(",")
        pg = m2.group(1).strip()
        pg = re.sub(r"^p(?=[0-9])", "p.", pg, flags=re.I)
        return head, pg
    return t, ""


def _journal_meta_break_without_issn(s, vol_start):
    """
    在 s[:vol_start] 内定位「meta（长数字码 / 起止页码 / 英文月年）」起始下标，用于无 ISSN 时刊名与 meta 分行。
    无匹配则返回 None（整块并入刊名行，仅 Volume 前再断行）。
    """
    prefix = (s[:vol_start] or "").strip()
    if not prefix:
        return None
    _mon_pat = r"(?:Jan(?:uary)?\.?|Feb(?:ruary)?\.?|Mar(?:ch)?\.?|Apr(?:il)?\.?|May\.?|Jun(?:e)?\.?|Jul(?:y)?\.?|Aug(?:ust)?\.?|Sep(?:t(?:ember)?)?\.?|Oct(?:ober)?\.?|Nov(?:ember)?\.?|Dec(?:ember)?\.?)"
    cuts = []
    for pat in (
        r",\s*\d{5,}\s*,",
        r",\s*\d+\s*[-—]\s*\d+\s*,",
        r",\s*" + _mon_pat + r"\s*\.?\s*\d{4}",
    ):
        for m in re.finditer(pat, prefix, re.I):
            cuts.append(m.start())
    mspace = re.search(
        r"\s+(?=" + _mon_pat + r"\.?\s*\d{4})",
        prefix,
        re.I,
    )
    if mspace:
        cuts.append(mspace.start())
    if not cuts:
        return None
    return min(cuts)


def _journal_volume_split_layout_blocks(p, flat, eng_comma):
    """
    混杂英文引文：刊名后换行、Volume 前换行；Volume 行与 p. 行分开；（共/约N页）单独一段。
    含标准 ISSN 或无 ISSN 仅有数字码+月年（如 134812, Aug. 2014）均处理。
    即使为英文月年逗号引文，只要 flat 中含 Volume: 仍可用本规则分行（解决整段塞在 journal 时刊名过长）。
    flat：_journal_citation_flat_for_parse 的扁平串（可含 |）。
    eng_comma：调用方传入以兼容旧接口，本函数内不再据此跳过。
    """
    if not (flat or "").strip():
        return None
    flat_nl = str(flat).replace("|", "\n").replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"\s+", " ", flat_nl.replace("\n", " ")).strip()

    page_count_s = str(p.get("page_count") or "").strip()
    vol_m = re.search(r"\bVolume\s*[：:]", s, re.I)
    ip = _journal_issn_or_isbn_match_start(s)

    if vol_m:
        vol_pos = vol_m.start()
        if ip is not None and ip >= vol_pos:
            ip = None
        if ip is not None:
            journal_line = s[:ip].strip().rstrip(",，")
            mid = s[ip:vol_pos].strip().strip(",，")
        else:
            mb = _journal_meta_break_without_issn(s, vol_pos)
            if mb is None:
                journal_line = s[:vol_pos].strip().rstrip(",，")
                mid = ""
            else:
                journal_line = s[:mb].strip().rstrip(",，")
                mid = s[mb:vol_pos].strip().strip(",，")
        mid = re.sub(r"^[\s,，]+|[,，]+$", "", (mid or "").strip()).strip()
        tail = s[vol_pos:].strip()
        tail_body, trailer = _journal_split_trailing_paren_pages(tail)
        vol_seg, pg_seg = _journal_split_volume_line_before_p_pages(tail_body)
        blocks = []
        if journal_line:
            blocks.append(("title", journal_line + ","))
        if mid:
            blocks.append(("plain", mid if mid.endswith(",") else mid + ","))
        if vol_seg:
            vline = vol_seg if vol_seg.endswith(",") else vol_seg + ","
            blocks.append(("plain", vline))
        if pg_seg:
            blocks.append(("plain", pg_seg))
        if trailer:
            blocks.append(("plain", trailer))
        elif page_count_s:
            blocks.append(("plain", f"（共{page_count_s}页）"))
        return blocks if len(blocks) >= 2 else None

    if not vol_m and ip is not None and ("录用待发表" in s or ("约" in s and "页" in s)):
        journal_line = s[:ip].strip().rstrip(",")
        rest = s[ip:].strip()
        rest_body, trailer = _journal_split_trailing_paren_pages(rest)
        blocks = []
        if journal_line:
            blocks.append(("title", journal_line + ","))
        if rest_body:
            blocks.append(("plain", rest_body))
        if trailer:
            blocks.append(("plain", trailer))
        elif page_count_s:
            blocks.append(("plain", f"（共{page_count_s}页）"))
        return blocks if len(blocks) >= 2 else None

    return None


def _journal_cn_bibliography_split_blocks(p, flat):
    """
    中文「刊名, ISSN, YYYY年M月, 卷期…」整段：刊名单独一行，ISSN+年月一行，卷期页一行，（共N页）单独一行。
    不与英文 Volume: / 英文月年逗号引文抢。
    """
    if not (flat or "").strip():
        return None
    flat_nl = str(flat).replace("|", "\n").replace("\r\n", "\n").replace("\r", "\n")
    if _journal_flat_looks_like_comma_citation(flat_nl):
        return None
    s = re.sub(r"\s+", " ", flat_nl.replace("\n", " ")).strip()
    if re.search(r"\bVolume\s*:", s, re.I):
        return None
    if not re.search(r"\d{4}年\s*\d+\s*月", s):
        return None
    ip = _journal_issn_or_isbn_match_start(s)
    if ip is None or ip < 2:
        return None
    title = s[:ip].strip().rstrip(",，")
    if not title:
        return None
    rest = s[ip:].strip()
    rest, trailer = _journal_split_trailing_paren_pages(rest)
    page_count_s = str(p.get("page_count") or "").strip()
    m = re.match(
        r"^(\d{3,4}-[\dXx][\dXx\-]*)\s*[,，]\s*(\d{4}年\s*\d+\s*月)\s*[,，]?\s*(.+)$",
        rest,
        re.I,
    )
    blocks = [("title", title + "，")]
    if m:
        blocks.append(("plain", f"{m.group(1)}, {m.group(2)}，"))
        vp = m.group(3).strip()
        if vp:
            # 卷期与页码分行：「37（05）, p. 20-30」→ 卷期行 + 页码行
            m_pg = re.search(r"[,，]\s*(?=p[.：]?\s*[\d\-—])", vp)
            if m_pg:
                vol_part = vp[: m_pg.start()].strip().rstrip(",，")
                pg_part = vp[m_pg.end() :].strip()
                pg_part = re.sub(r"^p(?=[0-9])", "p.", pg_part, flags=re.I)
                if vol_part:
                    blocks.append(("plain", vol_part + "，"))
                blocks.append(("plain", pg_part))
            else:
                blocks.append(("plain", vp))
    elif rest:
        blocks.append(("plain", rest))
    if trailer:
        blocks.append(("plain", trailer))
    elif page_count_s:
        blocks.append(("plain", f"（共{page_count_s}页）"))
    return blocks if len(blocks) >= 2 else None


def _journal_strip_p_pages_from_vol_issue_pages(vol_s, issue_s, pages_s):
    """volume 或 issue 误并入「, p.…」或「, Page …」时拆出到 pages_s（英文逗号引文常见抽取错位）。"""
    ps = _journal_line_page_keyword_to_p_dot((pages_s or "").strip())
    v = (vol_s or "").strip()
    isu = (issue_s or "").strip()
    v = re.sub(r"(?i),\s*Page\s+", ", p. ", v)
    isu = re.sub(r"(?i),\s*Page\s+", ", p. ", isu)
    if not re.search(r"(?i)p\.", ps):
        if re.search(r"(?i),\s*p\.", v):
            parts = re.split(r"(?i),\s*(?=p\.)", v, maxsplit=1)
            if len(parts) == 2:
                v = parts[0].strip().rstrip(",")
                ps = parts[1].strip()
        if not re.search(r"(?i)p\.", ps) and re.search(r"(?i),\s*p\.", isu):
            parts = re.split(r"(?i),\s*(?=p\.)", isu, maxsplit=1)
            if len(parts) == 2:
                isu = parts[0].strip().rstrip(",")
                ps = parts[1].strip()
    return v, isu, ps


def _fill_journal_cell_from_pub_raw(journal_cell, pub_raw):
    """
    按抽取时的 pub_raw（read_cell_marked_paragraphs：段间 \\n，[**…**] 等）原样还原，
    与 Word 原文的换行、加粗范围一致；无 pub_raw 时再用结构化版式。
    """
    clear_cell(journal_cell)
    first = True
    for raw_ln in (pub_raw or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        text = raw_ln.rstrip("\r")
        if not text.strip():
            continue
        para = journal_cell.paragraphs[0] if first else journal_cell.add_paragraph()
        first = False
        _tight_paragraph(para)
        _paragraph_keep_lines_together(para)
        for seg, b, u in parse_marked_authors(text):
            append_run(
                para,
                seg,
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=10.5,
                bold=b,
                underline=u,
            )


def _fill_journal_cell(journal_cell, p):
    """
    期刊列：
    - 「ISSN + 英文月年」逗号引文：一律多行结构化（刊名 / ISSN / 年月 / 卷期 / 页码 / 共N页），
      即使 pub_raw 含刊名加粗标记也如此（版式优先）。
    - 「刊名 / ISSN… / Volume:…」混杂引文：同样版式优先，可覆盖 pub_raw 的 [** 原样还原，以便在刊名后、Volume 前强制换行。
    - 「中文 ISSN + …年…月」整段：版式优先于 pub_raw 标记还原，刊名 / ISSN+年月 / 卷期页 分行。
    - APA「刊名 (ISSN). 年, …」等无法可靠结构化：按 pub_raw / journal 原文原样写入。
    - 其余且 pub_raw 含 [** 等：原样还原以保留 Word 局部样式。
    - 否则：结构化。
    """
    eff = _paper_journal_effective_fields(p)
    flat = _journal_citation_flat_for_parse(eff)
    pub_raw = str(eff.get("pub_raw") or "").strip()
    if (
        pub_raw
        and _journal_pub_raw_has_style_markup(pub_raw)
        and not _journal_flat_has_valid_issn(flat)
    ):
        _fill_journal_cell_from_pub_raw(journal_cell, pub_raw)
        return
    if _journal_should_fill_verbatim(eff, flat):
        _fill_journal_cell_verbatim(journal_cell, eff)
        return
    if _journal_flat_looks_like_comma_citation(flat):
        _fill_journal_cell_structured(journal_cell, eff)
        return
    eng_comma = bool(flat and _journal_flat_looks_like_comma_citation(flat))
    if _journal_volume_split_layout_blocks(eff, flat, eng_comma) or _journal_cn_bibliography_split_blocks(
        eff, flat
    ):
        _fill_journal_cell_structured(journal_cell, eff)
        return
    if pub_raw and _journal_pub_raw_has_style_markup(pub_raw):
        _fill_journal_cell_from_pub_raw(journal_cell, pub_raw)
    else:
        _fill_journal_cell_structured(journal_cell, eff)


def _fill_journal_cell_structured(journal_cell, p):
    """
    期刊列固定多行版式（每段单独一行）：
    1) 刊名（加粗，可与 journal_subtitle 同段）+ 逗号；中文刊用全角「，」；英文 ISSN+月年 类引文用半角「,」；
    2) ISSN 与发表年月同一行、半角逗号；英文引文无多余空格（例 0926-5805,Apr. 2020,）；中文可保留 ISSN 与年月间空格；
    3) 卷期：英文引文为「Volume:160,Issue:0,」单行（页码「p.…」单独一行，不与 Volume 段合并）；
    4) 页码单独一行（英文引文 p.105300；其余可为 p. 60-75）；
    5) （共N页）单独一行。
    """
    clear_cell(journal_cell)
    flat_for_cite = _journal_citation_flat_for_parse(p)
    eng_comma = bool(flat_for_cite and _journal_flat_looks_like_comma_citation(flat_for_cite))

    j_main, subtitle, vp_line = _journal_display_parts(p)
    if _is_paren_page_count_line(subtitle) or re.match(r"^[（(]\s*注\s*[：:]", subtitle or ""):
        subtitle = ""

    issn_s = str(p.get("issn") or "").strip()
    date_s = str(p.get("date") or "").strip()
    pages_s = str(p.get("pages") or "").strip()
    pages_s = _journal_line_page_keyword_to_p_dot(pages_s)
    vol_s = str(p.get("volume") or "").strip()
    issue_s = str(p.get("issue") or "").strip()
    page_count_s = str(p.get("page_count") or "").strip()

    if not pages_s and j_main:
        m_pg = re.search(r"[Pp]\.?\s*(\d+)", j_main)
        if m_pg:
            pages_s = m_pg.group(1).strip()

    if eng_comma:
        vol_s, issue_s, pages_s = _journal_strip_p_pages_from_vol_issue_pages(
            vol_s, issue_s, pages_s
        )
    pages_s, pages_paren_tail = _journal_split_trailing_paren_pages(pages_s)
    pages_note_s = ""
    if pages_paren_tail and not page_count_s:
        m_pc = re.search(r"(\d+)\s*页", pages_paren_tail)
        if m_pc:
            page_count_s = m_pc.group(1).strip()
        else:
            pages_note_s = pages_paren_tail

    title_bold = _journal_bold_title_only(j_main, issn_s)
    if not title_bold:
        title_bold = j_main

    eng_vol_line = ""
    if vol_s and "(" not in vol_s:
        if eng_comma:
            if issue_s != "":
                eng_vol_line = f"Volume:{vol_s},Issue:{issue_s},"
            else:
                eng_vol_line = f"Volume:{vol_s},"
        else:
            if issue_s:
                eng_vol_line = f"{vol_s}, {issue_s},"
            else:
                eng_vol_line = f"{vol_s},"

    pages_s = _journal_pages_s_strip_redundant_volume_line(
        pages_s, eng_comma=eng_comma, eng_vol_line=eng_vol_line
    )

    split_blocks = _journal_volume_split_layout_blocks(p, flat_for_cite, eng_comma)
    p_for_cn = dict(p)
    if page_count_s:
        p_for_cn["page_count"] = page_count_s
    cn_blocks = _journal_cn_bibliography_split_blocks(p_for_cn, flat_for_cite)
    if split_blocks:
        blocks = split_blocks
    elif cn_blocks:
        blocks = cn_blocks
    else:
        chinese_m = None
        if vp_line:
            chinese_m = re.match(r"^(\d+\(\d+\))\s*[：:]\s*(.+)$", vp_line.strip())

        blocks = []
        line1 = title_bold
        if subtitle:
            line1 += subtitle
        title_sep = "," if eng_comma else "，"
        if line1.strip():
            blocks.append(("title", line1.strip() + title_sep))
        if issn_s and date_s:
            if eng_comma:
                blocks.append(("plain", f"{issn_s},{date_s},"))
            else:
                blocks.append(("plain", f"{issn_s}, {date_s},"))
        elif issn_s:
            blocks.append(("plain", issn_s + ","))
        elif date_s:
            blocks.append(("plain", date_s + ","))

        if eng_vol_line:
            vol_only, pg_from_vol = _journal_split_volume_line_before_p_pages(
                eng_vol_line.rstrip(",").strip()
            )
            vol_disp = (
                vol_only
                if (vol_only.endswith(",") or not vol_only)
                else (vol_only + ",")
            )
            if not vol_disp:
                vol_disp = eng_vol_line
            blocks.append(("plain", vol_disp))
            pages_use = pages_s
            if not pages_use.strip() and pg_from_vol:
                pages_use = re.sub(r"(?i)^p\.\s*", "", pg_from_vol)
                pages_use = re.sub(r"\s+", "", pages_use)
            if pages_use:
                blocks.append(
                    ("plain", _journal_pages_display_line(pages_use, no_space_after_p=eng_comma))
                )
            elif pg_from_vol:
                blocks.append(("plain", pg_from_vol))
        elif chinese_m:
            blocks.append(("plain", chinese_m.group(1) + "，"))
            tail = chinese_m.group(2).strip()
            if tail:
                blocks.append(("plain", tail))
            elif pages_s:
                blocks.append(
                    ("plain", _journal_pages_display_line(pages_s, no_space_after_p=False))
                )
        elif vp_line:
            blocks.append(("plain", vp_line))
        elif pages_s:
            blocks.append(
                ("plain", _journal_pages_display_line(pages_s, no_space_after_p=False))
            )

        if page_count_s:
            blocks.append(("plain", f"（共{page_count_s}页）"))
        if pages_note_s:
            blocks.append(("plain", pages_note_s))

    if not blocks:
        clear_cell(journal_cell)
        return

    for i, (kind, text) in enumerate(blocks):
        para = journal_cell.paragraphs[0] if i == 0 else journal_cell.add_paragraph()
        _tight_paragraph(para)
        _paragraph_keep_lines_together(para)
        append_run(
            para,
            text,
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            bold=(kind == "title"),
        )


def _split_recommendation_mixed_body(raw):
    """
    推荐信正文中英文粘在一段时拆段；输出顺序为英文段在前、中文段在后（另起一段）。
    - 显式「\\n\\n」分段；
    - 单行「\\n」且第二段以中文起、首段含英文时，拆成两段；
    - 「句末 .!?（及可选引号）」后可直接接中文（无空格），在最后一次此类过渡处切开（避免误切 Dr. / modeling.In 等）；
    - 兼容旧数据：句末后 ≥2 空白再接中文。
    """
    raw = (raw or "").strip()
    if not raw:
        return []
    if "\n" in raw:
        parts = [p.strip() for p in raw.split("\n") if p.strip()]
        if len(parts) > 1:
            return parts

    # 英文句末标点之后紧跟中文（0 或多个空白均可）：取「head 最长」的切分点，对应整段英文结束后接中文正文
    candidates = []
    for m in re.finditer(r'[.!?…]["\'」』）]?\s*(?=[\u4e00-\u9fff])', raw):
        split_at = m.end()
        head = raw[:split_at].strip()
        tail = raw[split_at:].strip()
        if not tail or len(head) < 8:
            continue
        if not re.search(r"[a-zA-Z]{3,}", head):
            continue
        if not re.match(r"[\u4e00-\u9fff]", tail):
            continue
        candidates.append((len(head), head, tail))
    if candidates:
        _, head, tail = max(candidates, key=lambda x: x[0])
        return [head, tail]

    last_ok = None
    for m in re.finditer(r"[\.!?]\s{2,}", raw):
        tail = raw[m.end() :].lstrip()
        head = raw[: m.start() + 1].strip()
        if tail and re.match(r"[\u4e00-\u9fff]", tail) and re.search(r"[a-zA-Z]", head):
            last_ok = (head, tail)
    if last_ok:
        return [last_ok[0], last_ok[1]]
    return [raw]


def _recommendation_body_paragraphs(rec):
    """推荐信正文段落（不含「序号.姓名 职称」标题行）。优先 content_en / content_zh，否则解析 content。"""
    content_en = str(rec.get("content_en", "") or "").strip()
    content_zh = str(rec.get("content_zh", "") or "").strip()
    raw = str(rec.get("content", "") or "").strip()
    if content_en:
        out = [content_en]
        if content_zh:
            out.append(content_zh)
        return out
    if content_zh:
        return [content_zh]
    return _split_recommendation_mixed_body(raw)


# 推荐语正文粘在「…方面专家 / …领域专家」后且无冒号分隔时，用小标题结尾词切开（不依赖 ：）
_RE_REC_SUBTITLE_BODY = re.compile(r"^(\d+\..+?(?:方面专家|领域专家))(.+)$", re.S)


def _format_recommendation_heading(index_one_based, name, role, title_f):
    """
    推荐人首行加粗标题：与常见简历一致。
    - 有 role： «序号.姓名（角色），职称…»
    - 无 role： «序号.姓名 职称…»（空格分隔，兼容旧数据）
    另：若抽取把「（角色）」并进姓名、或 role 落在 title 开头，在此合并还原。
    """
    name = (name or "").strip()
    role = (role or "").strip()
    title_f = (title_f or "").strip()
    # title 以「（角色），职称…」开头而 role 为空（解析错位）时纠正
    if not role and title_f:
        m_t = re.match(r"^[（(]\s*([^（）()]+)\s*[）)]\s*[，,]\s*(.+)$", title_f, re.S)
        if m_t:
            role = m_t.group(1).strip()
            title_f = m_t.group(2).strip()
    # 仅姓名末尾带「（角色）」而 role 为空（第二套正则把括号并进 name）
    if not role and name:
        m_n = re.match(r"^(.+)[（(]([^（）()]+)[）)]\s*$", name)
        if m_n:
            name, role = m_n.group(1).strip(), m_n.group(2).strip()
    # name 与 role 同时含同一括号段时去重
    if role and name:
        for suf in (f"（{role}）", f"({role})"):
            if suf in name:
                name = name.replace(suf, "").strip()
                break
    if not name and not title_f:
        return ""
    if role:
        h = f"{index_one_based}.{name}（{role}）" if name else f"{index_one_based}.（{role}）"
        if title_f:
            h += f"，{title_f}"
        return h
    sub = f"{index_one_based}.{name} {title_f}".strip()
    return re.sub(r"\s+", " ", sub)


def _recommendation_subtitle_and_body_paragraphs(rec, index_one_based):
    """
    返回 (加粗小标题文本或 None, 正文段落列表)。
    - name / title 有值：小标题为「序号.name（role），title」或「序号.name title」；
      role 来自抽取；正文来自 content（或 content_en/zh）。
    - 二者皆空：若 content 匹配「序号…方面专家|领域专家」后紧跟推荐正文，则拆开；否则整段 content 作正文（无单独小标题行）。
    """
    name = str(rec.get("name", "") or "").strip()
    role = str(rec.get("role", "") or "").strip()
    title_f = str(rec.get("title", "") or "").strip()
    content_en = str(rec.get("content_en", "") or "").strip()
    content_zh = str(rec.get("content_zh", "") or "").strip()

    if content_en or content_zh:
        bodies = []
        if content_en:
            bodies.append(content_en)
        if content_zh:
            bodies.append(content_zh)
        if name or title_f:
            sub = _format_recommendation_heading(index_one_based, name, role, title_f)
            return sub, bodies
        return None, bodies

    raw = str(rec.get("content", "") or "").strip()
    if name or title_f:
        sub = _format_recommendation_heading(index_one_based, name, role, title_f)
        if raw:
            return sub, _recommendation_body_paragraphs(rec)
        return sub, []

    if not raw:
        return None, []

    m = _RE_REC_SUBTITLE_BODY.match(raw)
    if m:
        sub = m.group(1).strip()
        rest = m.group(2).strip()
        if not rest:
            return sub, []
        inner = _split_recommendation_mixed_body(rest)
        return sub, inner if inner else [rest]

    return None, _recommendation_body_paragraphs(rec)


def fill_projects_summary_like_template(cell, pjs):
    """
    与「调整空白版」一致：仅两行——主持总体 + 参与总体；整行加粗，数字下划线。
    不再追加「主持项目总体情况」「参与项目总体情况」简写行。
    """
    clear_cell(cell)
    ham, hak, ha, hb = (
        _ps_n(pjs, "host_a_major"),
        _ps_n(pjs, "host_a_key"),
        _ps_n(pjs, "host_a"),
        _ps_n(pjs, "host_b"),
    )
    pam, pak, pa, pb = (
        _ps_n(pjs, "participate_a_major"),
        _ps_n(pjs, "participate_a_key"),
        _ps_n(pjs, "participate_a"),
        _ps_n(pjs, "participate_b"),
    )
    line1 = (
        f"总体情况：主持A类重大项目  {ham}  项、A类重点项目  {hak}  项、"
        f"A类项目  {ha}  项、B类项目  {hb}  项。"
    )
    line2 = (
        f"          参与A类重大项目  {pam}  项、A类重点项目  {pak}  项、"
        f"A类项目  {pa}  项、B类项目  {pb}  项。"
    )
    p0 = cell.paragraphs[0]
    clear_paragraph_runs(p0)
    _tight_paragraph(p0)
    append_run_bold_underline_digits(p0, line1, font_cn="宋体", font_en="Times New Roman", size_pt=10.5)
    p1 = cell.add_paragraph()
    clear_paragraph_runs(p1)
    _tight_paragraph(p1)
    append_run_bold_underline_digits(p1, line2, font_cn="宋体", font_en="Times New Roman", size_pt=10.5)


def fill_stats_like_correct001(cell, ps):
    """
    与「调整空白版」统计条一致：每条「标签：\\t篇数 篇（…）」仅标签后一个制表符；
    每行最多 2 条（分号连接），多于 2 条则换行后继续。
    在「按学校期刊目录分级统计」之后，可按 JSON 非零项追加「按中国计算机学会分类统计」「按 ABS 分类统计」。
    """
    clear_cell(cell)
    total = _ps_n(ps, "total")
    jcr_body = _stats_join_n_per_line(_paper_stats_jcr_segments_tabbed(ps), 2)
    school_body = _stats_join_n_per_line(_paper_stats_school_segments_tabbed(ps), 2)
    stat_pt = 12

    def add_para(text, *, bold, tabbed=False, hide_decimal_suffix_white=False):
        nonlocal first
        if first:
            p = cell.paragraphs[0]
            clear_paragraph_runs(p)
            first = False
        else:
            p = cell.add_paragraph()
            clear_paragraph_runs(p)
        _tight_paragraph(p)
        # 换行后的明细段与首段对齐：顶格、无悬挂缩进（避免模板遗留首行缩进）
        set_paragraph_format(p, first_line_chars=0, line_spacing=1.0)
        p.paragraph_format.left_indent = Pt(0)
        if tabbed:
            _set_stat_detail_tab_stops(p)
        if hide_decimal_suffix_white:
            _append_stats_detail_text_runs(
                p,
                text,
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
                bold=bold,
            )
        else:
            append_run(
                p,
                text,
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
                bold=bold,
            )

    first = True
    add_para(
        f"总体情况：第一作者或通讯作者发表论文（ {_stats_display_num(total, False)} 篇）",
        bold=True,
        tabbed=False,
    )
    if jcr_body:
        add_para("按JCR分区统计：", bold=True, tabbed=False)
        add_para(jcr_body, bold=False, tabbed=True, hide_decimal_suffix_white=True)
        p_blank = cell.add_paragraph()
        clear_paragraph_runs(p_blank)
        _tight_paragraph(p_blank)
    add_para("按学校期刊目录分级统计：", bold=True, tabbed=False)
    if school_body:
        add_para(school_body, bold=False, tabbed=True, hide_decimal_suffix_white=True)
    _fill_stats_ccf_tail(cell, ps)
    _fill_stats_abs_tail(cell, ps)


def _fill_stats_ccf_tail(cell, ps):
    """与 fill_stats_like_correct001 末尾一致：追加「按中国计算机学会分类统计」。"""
    ccf_body = _stats_join_n_per_line(_paper_stats_ccf_segments_tabbed(ps), 2)
    if not ccf_body:
        return
    stat_pt = 12
    p_blank = cell.add_paragraph()
    clear_paragraph_runs(p_blank)
    _tight_paragraph(p_blank)
    p_title = cell.add_paragraph()
    clear_paragraph_runs(p_title)
    _tight_paragraph(p_title)
    set_paragraph_format(p_title, first_line_chars=0, line_spacing=1.0)
    p_title.paragraph_format.left_indent = Pt(0)
    append_run(
        p_title,
        "按中国计算机学会分类统计：",
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=stat_pt,
        bold=True,
    )
    p_body = cell.add_paragraph()
    clear_paragraph_runs(p_body)
    _tight_paragraph(p_body)
    set_paragraph_format(p_body, first_line_chars=0, line_spacing=1.0)
    p_body.paragraph_format.left_indent = Pt(0)
    _set_stat_detail_tab_stops(p_body)
    _append_stats_detail_text_runs(
        p_body,
        ccf_body,
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=stat_pt,
        bold=False,
    )


def _fill_stats_abs_tail(cell, ps):
    """papers_summary 中含 abs_* 非零时追加「按 ABS 分类统计」段落。"""
    abs_body = _stats_join_n_per_line(_paper_stats_abs_segments_tabbed(ps), 2)
    if not abs_body:
        return
    stat_pt = 12
    p_blank = cell.add_paragraph()
    clear_paragraph_runs(p_blank)
    _tight_paragraph(p_blank)
    p_title = cell.add_paragraph()
    clear_paragraph_runs(p_title)
    _tight_paragraph(p_title)
    set_paragraph_format(p_title, first_line_chars=0, line_spacing=1.0)
    p_title.paragraph_format.left_indent = Pt(0)
    append_run(
        p_title,
        "按ABS分类统计：",
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=stat_pt,
        bold=True,
    )
    p_body = cell.add_paragraph()
    clear_paragraph_runs(p_body)
    _tight_paragraph(p_body)
    set_paragraph_format(p_body, first_line_chars=0, line_spacing=1.0)
    p_body.paragraph_format.left_indent = Pt(0)
    _set_stat_detail_tab_stops(p_body)
    _append_stats_detail_text_runs(
        p_body,
        abs_body,
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=stat_pt,
        bold=False,
    )


def fill_stats_nested_adjust_blank(cell, ps):
    """
    「调整空白版」论文统计格：内嵌大表第 3～4 行为 Q1|Q2、Q3|Q4，第 5 行为「其他」占位。
    仅改写内嵌单元格，不 clear_cell，以便保留模板表格；JSON 中 jcr_other / jcr_other_first 写入第 5 行左三格。
    字号与段落模式统一为 12pt，避免同一格内大小不一显得杂乱。
    不满足结构时返回 False，由 fill_stats_like_correct001 回退。
    """
    stat_pt = 12
    if len(cell.tables) < 2:
        return False
    big = cell.tables[0]
    school_tbl = cell.tables[1]
    if len(big.rows) < 5 or len(big.rows[2].cells) < 8:
        return False
    if not _paper_stats_jcr_has_content(ps):
        return False

    j_other, j_other_first = _jcr_other_pair(ps)
    dec_align = _papers_summary_has_fraction(ps)

    def nz(v):
        x = _ps_n({"v": v}, "v")
        if x == 0 or x == 0.0:
            return ""
        return _stats_display_num(x, False)

    def clear_table(tbl):
        for row in tbl.rows:
            for c in row.cells:
                write_cell_keep_style(c, "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)

    def _write_jcr_new_row_zone(row, side, label, tk, fk, end_suffix):
        tn = _ps_n(ps, tk)
        fn = _ps_n(ps, fk)
        if tn <= 0:
            if side == 0:
                for ci in (0, 1, 2):
                    write_cell_keep_style(row.cells[ci], "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            else:
                for ci in (3, 6, 7, 8):
                    if ci < len(row.cells):
                        write_cell_keep_style(row.cells[ci], "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            return
        ts = _stats_display_num(tn, dec_align)
        fs = _stats_display_num(fn, dec_align)
        if side == 0:
            write_cell_keep_style(row.cells[0], f"{label}：", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            write_cell_keep_style_decimal_hidden(
                row.cells[1],
                f" {ts} 篇",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
            )
            write_cell_keep_style_decimal_hidden(
                row.cells[2],
                f"（ {fs} 篇为第一作者）{end_suffix}",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
            )
        else:
            write_cell_keep_style(row.cells[3], f"{label}：", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            write_cell_keep_style_decimal_hidden(
                row.cells[6],
                f" {ts} 篇（ {fs} 篇为第一作者）{end_suffix}",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
            )
            if len(row.cells) > 7:
                write_cell_keep_style(row.cells[7], "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            if len(row.cells) > 8:
                write_cell_keep_style(row.cells[8], "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)

    def _write_school_new_row_zone(row, side, label, tk, fk, end_suffix):
        tn = _ps_n(ps, tk)
        fn = _ps_n(ps, fk)
        if tn <= 0:
            if side == 0:
                for ci in (0, 1, 2):
                    write_cell_keep_style(row.cells[ci], "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            else:
                for ci in (3, 4, 5):
                    write_cell_keep_style(row.cells[ci], "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            return
        ts = _stats_display_num(tn, dec_align)
        fs = _stats_display_num(fn, dec_align)
        if side == 0:
            write_cell_keep_style(row.cells[0], f"{label}：", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            write_cell_keep_style_decimal_hidden(
                row.cells[1],
                f" {ts} 篇",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
            )
            write_cell_keep_style_decimal_hidden(
                row.cells[2],
                f"（ {fs} 篇为第一作者）{end_suffix}",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
            )
        else:
            write_cell_keep_style(row.cells[3], f"{label}：", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
            write_cell_keep_style_decimal_hidden(
                row.cells[4],
                f" {ts} 篇",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
            )
            write_cell_keep_style_decimal_hidden(
                row.cells[5],
                f"（ {fs} 篇为第一作者）{end_suffix}",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=stat_pt,
            )

    if len(big.rows[0].cells) > 4:
        write_cell_keep_style(
            big.rows[0].cells[4],
            nz(ps.get("total", 0)),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=stat_pt,
        )
    for ridx in (2, 3, 4):
        row = big.rows[ridx]
        for ci in range(len(row.cells)):
            write_cell_keep_style(row.cells[ci], "", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)

    q2_ok = _ps_n(ps, "jcr_q2") > 0
    q4_ok = _ps_n(ps, "jcr_q4") > 0
    _write_jcr_new_row_zone(big.rows[2], 0, "Q1区", "jcr_q1", "jcr_q1_first", "；" if q2_ok else "")
    _write_jcr_new_row_zone(big.rows[2], 1, "Q2区", "jcr_q2", "jcr_q2_first", "")
    _write_jcr_new_row_zone(big.rows[3], 0, "Q3区", "jcr_q3", "jcr_q3_first", "；" if q4_ok else "")
    _write_jcr_new_row_zone(big.rows[3], 1, "Q4区", "jcr_q4", "jcr_q4_first", "")

    row4 = big.rows[4]
    if j_other > 0:
        write_cell_keep_style(row4.cells[0], "其他：", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
        jo = _stats_display_num(j_other, dec_align)
        jof = _stats_display_num(j_other_first, dec_align)
        write_cell_keep_style_decimal_hidden(
            row4.cells[1],
            f" {jo} 篇",
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=stat_pt,
        )
        write_cell_keep_style_decimal_hidden(
            row4.cells[2],
            f"（ {jof} 篇为第一作者）；",
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=stat_pt,
        )
    else:
        write_cell_keep_style(row4.cells[0], "其他：", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
        write_cell_keep_style(row4.cells[1], "   篇", font_cn="宋体", font_en="Times New Roman", size_pt=stat_pt)
        write_cell_keep_style(
            row4.cells[2],
            "（   篇为第一作者）；",
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=stat_pt,
        )

    clear_table(school_tbl)
    b_ok = _ps_n(ps, "school_b") > 0
    o_ok = _ps_n(ps, "school_other") > 0
    if len(school_tbl.rows) >= 2:
        _write_school_new_row_zone(school_tbl.rows[0], 0, "A类", "school_a", "school_a_first", "；" if b_ok else "")
        _write_school_new_row_zone(school_tbl.rows[0], 1, "B类", "school_b", "school_b_first", "")
        _write_school_new_row_zone(school_tbl.rows[1], 0, "C类", "school_c", "school_c_first", "；" if o_ok else "")
        _write_school_new_row_zone(school_tbl.rows[1], 1, "其他", "school_other", "school_other_first", "")

    sas = _ps_n(ps, "school_a_star")
    sasf = _ps_n(ps, "school_a_star_first")
    if sas > 0:
        # 勿在 tbl 后 insert XML，以免与模板原有段落顺序错乱；接到单元格末尾，版式与段落模式一致。
        p = cell.add_paragraph()
        clear_paragraph_runs(p)
        _tight_paragraph(p)
        set_paragraph_format(p, first_line_chars=0, line_spacing=1.0)
        p.paragraph_format.left_indent = Pt(0)
        _set_stat_detail_tab_stops(p)
        _append_stats_detail_text_runs(
            p,
            _fmt_school_ccf_tabbed("A*类", sas, sasf, ps),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=stat_pt,
            bold=False,
        )

    return True


def fill_paper_stats_cell(cell, ps, *, use_nested_table=False):
    """
    默认用段落 + 制表符（fill_stats_like_correct001），整格版式统一、最稳定。
    若 JSON 顶层 fill_paper_stats_nested=true 且模板含 5 行 JCR 内嵌表，则写入内嵌单元格（含 jcr_other 行），再追加 CCF、ABS 段落。
    """
    if use_nested_table and fill_stats_nested_adjust_blank(cell, ps):
        _fill_stats_ccf_tail(cell, ps)
        _fill_stats_abs_tail(cell, ps)
    else:
        fill_stats_like_correct001(cell, ps)


def _safe_row(table, idx):
    return table.rows[idx] if idx < len(table.rows) else None


def _paper_cell_center(cell):
    """论文明细列：水平居中 + 垂直居中（表格属性）。"""
    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    except Exception:
        pass
    for p in cell.paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _keep_table_row_on_one_page(row):
    """
    表格行不允许跨页断开（w:cantSplit），使同一篇论文的序号/作者/题目/期刊/分级整行在同一页。
    """
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    for el in list(trPr.findall(qn("w:cantSplit"))):
        trPr.remove(el)
    trPr.append(OxmlElement("w:cantSplit"))


def _set_row_auto_height(row):
    """清除固定行高，交由内容自动撑开，避免页面出现大面积空白。"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        return
    for h in list(trPr.findall(qn("w:trHeight"))):
        trPr.remove(h)


def _paragraph_keep_lines_together(paragraph):
    """段落内不分页（与下一段同属一行内多段时，减少期刊信息被拆开）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def _p_el_snap_to_grid_off(p_el):
    """段落 XML：关闭「如果定义了文档网格，则对齐到网格」，与调优空白版一致。"""
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    snap = p_pr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        p_pr.append(snap)
    snap.set(qn("w:val"), "0")


def _document_snap_to_grid_off(doc):
    """全文段落关闭对齐到网格（含表格内所有嵌套段落）。"""
    for p_el in doc.part.element.iter(qn("w:p")):
        _p_el_snap_to_grid_off(p_el)


def _set_cell_text(cell, text, *, font_cn, font_en, size_pt, bold=False, first_line_chars=0, paragraph_align=None):
    clear_cell(cell)
    p = cell.paragraphs[0]
    clear_paragraph_runs(p)
    if paragraph_align is not None:
        p.alignment = paragraph_align
    append_run(
        p,
        str(text),
        font_cn=font_cn,
        font_en=font_en,
        size_pt=size_pt,
        bold=bold,
    )


# 现职称 cell：末尾「（YYYY.MM）」类评审时间与职称拆开两行；时间为六号（7.5pt）
_CURRENT_TITLE_REVIEW_FW = re.compile(
    r"^(.+?)\s*（\s*(\d{4}\.\d{1,2}(?:\.\d{1,2})?)\s*）\s*$"
)
_CURRENT_TITLE_REVIEW_HW = re.compile(
    r"^(.+?)\s*\(\s*(\d{4}\.\d{1,2}(?:\.\d{1,2})?)\s*\)\s*$"
)


def _split_current_title_review(text):
    """拆成 (职称正文, 评审时间含括号)；无法识别时间尾则第二段为空。"""
    s = str(text or "").strip()
    if not s:
        return "", ""
    # 将「YYYY年MM月」格式转为「YYYY.MM」
    s = re.sub(r"（(\d{4})年(\d{1,2})月）", lambda m: f"（{m.group(1)}.{int(m.group(2)):02d}）", s)
    s = re.sub(r"\((\d{4})年(\d{1,2})月\)", lambda m: f"({m.group(1)}.{int(m.group(2)):02d})", s)
    # 处理「职称\n（日期）」换行格式
    if "\n" in s:
        parts = [p.strip() for p in s.split("\n") if p.strip()]
        s = " ".join(parts)
    m = _CURRENT_TITLE_REVIEW_FW.match(s)
    if m:
        return m.group(1).strip(), f"（{m.group(2)}）"
    m = _CURRENT_TITLE_REVIEW_HW.match(s)
    if m:
        return m.group(1).strip(), f"（{m.group(2)}）"
    return s, ""


def _fill_basic_current_title_cell(cell, raw_text, *, main_size_pt=12, review_size_pt=7.5):
    """
    「现职称（评审时间）」：有评审时间时在时间前换行；职称行为 main_size_pt，评审时间为六号（默认 7.5pt）。
    原文为空时填「无」（与宗教信仰格一致）。
    """
    main, review = _split_current_title_review(raw_text)
    clear_cell(cell)
    if not main and not review:
        p = cell.paragraphs[0]
        clear_paragraph_runs(p)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _tight_paragraph(p)
        append_run(
            p,
            "无",
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=main_size_pt,
            bold=False,
        )
        _paper_cell_center(cell)
        return
    if not review:
        p = cell.paragraphs[0]
        clear_paragraph_runs(p)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _tight_paragraph(p)
        append_run(
            p,
            main,
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=main_size_pt,
            bold=False,
        )
        _paper_cell_center(cell)
        return

    p0 = cell.paragraphs[0]
    clear_paragraph_runs(p0)
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _tight_paragraph(p0)
    append_run(
        p0,
        main,
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=main_size_pt,
        bold=False,
    )

    p1 = cell.add_paragraph()
    clear_paragraph_runs(p1)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _tight_paragraph(p1)
    append_run(
        p1,
        review,
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=review_size_pt,
        bold=False,
    )
    _paper_cell_center(cell)


def _inflate_row_height_by_paragraphs(row, cols, spacer_count):
    """通过在指定列追加空段，轻微抬高整行高度，用于将后续块推到下一页。"""
    n = max(0, int(spacer_count or 0))
    if n <= 0:
        return
    for c in cols:
        if c >= len(row.cells):
            continue
        cell = row.cells[c]
        for _ in range(n):
            p = cell.add_paragraph("")
            _tight_paragraph(p)


def _insert_blank_rows_after(table, after_row_idx, count):
    """在指定行后插入 count 行（克隆当前行结构）。"""
    n = max(0, int(count or 0))
    if n <= 0:
        return
    anchor = table.rows[after_row_idx]._tr
    tmpl = table.rows[after_row_idx]._tr
    for _ in range(n):
        new_tr = deepcopy(tmpl)
        anchor.addnext(new_tr)
        anchor = new_tr


def _bump_row_height(row, boost_pt):
    """将行高在原模板基础上增加 boost_pt（pt），用于视觉填充页底空白。"""
    if boost_pt <= 0:
        return
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    h = trPr.find(qn("w:trHeight"))
    if h is None:
        h = OxmlElement("w:trHeight")
        h.set(qn("w:val"), str(int(round(boost_pt * 20))))
        h.set(qn("w:hRule"), "exact")
        trPr.append(h)
        return
    try:
        cur = int(h.get(qn("w:val")) or "0")
    except Exception:
        cur = 0
    h.set(qn("w:val"), str(max(0, cur + int(round(boost_pt * 20)))))
    h.set(qn("w:hRule"), "exact")


def split_table_before_row_with_page_break(doc, table, split_row_idx):
    """
    将一个大表在指定行前硬拆分为两个表，中间插入分页符。
    用于保证后半表（如推荐人意见）从新页页首开始。
    """
    tr_nodes = list(table._tbl.findall(qn("w:tr")))
    if split_row_idx <= 0 or split_row_idx >= len(tr_nodes):
        return

    first_tbl = table._tbl
    # 勿用 copy.deepcopy(oxml)：部分环境下子节点仍与首表共享，首表 remove 行时会牵连第二表，
    # Word 中表现为推荐人意见等后半表内容丢失。序列化再 parse 可得到独立 DOM。
    tbl_xml = etree.tostring(first_tbl, encoding="utf-8", xml_declaration=False, with_tail=False)
    second_tbl = parse_xml(tbl_xml)

    # first_tbl 保留 [0, split_row_idx)
    first_rows = list(first_tbl.findall(qn("w:tr")))
    for i in range(len(first_rows) - 1, split_row_idx - 1, -1):
        first_tbl.remove(first_rows[i])

    # second_tbl 保留 [split_row_idx, end)
    second_rows = list(second_tbl.findall(qn("w:tr")))
    for i in range(split_row_idx - 1, -1, -1):
        second_tbl.remove(second_rows[i])

    # 在两个表之间插入分页段落
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)

    first_tbl.addnext(p)
    p.addnext(second_tbl)


def fill_template(template_path, data_json, output_path=None, compact=False):
    """读取 JSON，填充空白版模板，输出规范 Word。compact=True 时删除无数据的论文/项目空行。"""
    with open(data_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("模板未检测到表格。")
    table = doc.tables[0]

    # ---------------------------
    # 表格0（基本信息）-> 大表 row 0~1
    # ---------------------------
    bi = data.get("basic_info", {})
    # 空白模板中的值区域是“标签后面的占位区”，避免写入标签单元格。
    basic_map = {
        (0, 4): bi.get("name", ""),
        (0, 8): bi.get("gender", ""),
        (0, 12): bi.get("hometown", ""),
        (0, 17): bi.get("marital_status", ""),
        (0, 22): normalize_birth_text(bi.get("birth", "")),
        (1, 8): bi.get("discipline", ""),
        (1, 17): (str(bi.get("religion", "") or "").strip() or "无"),
        (1, 22): bi.get("target_college", ""),
    }
    for (r, c), v in basic_map.items():
        if r < len(table.rows) and c < len(table.rows[r].cells):
            _set_cell_text(
                table.rows[r].cells[c],
                v,
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=12,
                bold=False,
                paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            )
            _paper_cell_center(table.rows[r].cells[c])

    if len(table.rows) > 1 and 4 < len(table.rows[1].cells):
        _fill_basic_current_title_cell(table.rows[1].cells[4], bi.get("current_title", ""))

    # --------------------------------
    # 表格1（学习/工作/海外经历）-> 大表 row 2
    # --------------------------------
    edu_lines = format_with_double_space(data.get("education", []))
    work_lines = format_with_double_space(data.get("work", []))
    overseas_lines = format_with_double_space(data.get("overseas", []))
    edu_lines, ne = _partition_experience_notes_from_lines(edu_lines)
    work_lines, nw = _partition_experience_notes_from_lines(work_lines)
    overseas_lines, no = _partition_experience_notes_from_lines(overseas_lines)
    if not edu_lines:
        edu_lines = ["无"]
    if not work_lines:
        work_lines = ["无"]
    if not overseas_lines:
        overseas_lines = ["无"]
    experience_notes = list(data.get("experience_notes") or []) + ne + nw + no
    _univ_985211 = _load_985211_name_set(Path(__file__).resolve().parent)

    exp_row = _safe_row(table, 2)
    if exp_row is not None:
        cell = exp_row.cells[2]
        if not fill_learning_nested_tables(
            cell, edu_lines, work_lines, overseas_lines, _univ_985211, experience_notes=experience_notes
        ):
            clear_cell(cell)
            p = cell.paragraphs[0]
            clear_paragraph_runs(p)
            append_run(
                p,
                "学习经历：",
                font_cn="宋体",
                font_en="Times New Roman",
                size_pt=12,
                bold=True,
                underline=True,
                italic=True,
            )
            for line in edu_lines:
                p = cell.add_paragraph()
                clear_paragraph_runs(p)
                _append_education_line_runs(p, line, _univ_985211)
            cell.add_paragraph("")
            p = cell.add_paragraph()
            clear_paragraph_runs(p)
            append_run(p, "工作经历：", font_cn="宋体", font_en="Times New Roman", size_pt=12, bold=True, underline=True, italic=True)
            for line in work_lines:
                p = cell.add_paragraph()
                clear_paragraph_runs(p)
                append_run(p, line, font_cn="仿宋", font_en="Times New Roman", size_pt=12)
            cell.add_paragraph("")
            p = cell.add_paragraph()
            clear_paragraph_runs(p)
            append_run(p, "海外经历：", font_cn="宋体", font_en="Times New Roman", size_pt=12, bold=True, underline=True, italic=True)
            for line in overseas_lines:
                p = cell.add_paragraph()
                clear_paragraph_runs(p)
                append_run(p, line, font_cn="仿宋", font_en="Times New Roman", size_pt=12)
            if any(str(n).strip() for n in (experience_notes or [])):
                p = cell.add_paragraph()
                clear_paragraph_runs(p)
                _tight_paragraph(p)
            for n in experience_notes or []:
                ns = str(n).strip()
                if not ns:
                    continue
                p = cell.add_paragraph()
                clear_paragraph_runs(p)
                append_run(p, ns, font_cn="仿宋", font_en="Times New Roman", size_pt=12)

    # ----------------------------
    # 表格2（研究信息）-> 大表 row 3~4
    # ----------------------------
    _set_cell_text(
        table.rows[3].cells[6],
        _normalize_piped_table_text(data.get("research_area", "")),
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=12,
        paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
    )
    _paper_cell_center(table.rows[3].cells[6])
    _set_cell_text(
        table.rows[3].cells[18],
        _normalize_piped_table_text(data.get("research_team", ""), ""),
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=12,
        bold=True,
        paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
    )
    _paper_cell_center(table.rows[3].cells[18])
    _set_cell_text(
        table.rows[4].cells[18],
        _normalize_piped_table_text(data.get("team_achievement", "")),
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=12,
        paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
    )
    _paper_cell_center(table.rows[4].cells[18])

    # ----------------------------
    # 表格3（论文统计）-> 大表 row 5
    # papers_summary 原样显示（不从明细推算）；段落填充；fill_paper_stats_nested=true 时用内嵌表。
    # 注意：row 13 为「项目总体情况」占位，勿再写入校级论文统计（此前会重复且撑出空行）。
    # ----------------------------
    # papers_summary 数字与 JSON 完全一致，不从 papers 明细推算。
    ps = deepcopy(data.get("papers_summary") or {})
    use_nested = bool(data.get("fill_paper_stats_nested", False))
    r5 = table.rows[5]
    c_stats = r5.cells[3]
    fill_paper_stats_cell(c_stats, ps, use_nested_table=use_nested)
    if r5.cells[15]._tc is not c_stats._tc:
        fill_paper_stats_cell(r5.cells[15], ps, use_nested_table=use_nested)

    # ----------------------------
    # 表格4（论文明细）-> 大表 row 7 开始（模板默认 6 行：7~12）
    # 论文条数超过模板行数时，在末行后插入与末行同结构的行（与「科研项目」区逻辑一致）。
    # ----------------------------
    papers = data.get("papers", [])
    paper_start = 7
    paper_end = 12
    paper_data_rows_inserted = 0
    while len(papers) > (paper_end - paper_start + 1):
        _insert_blank_rows_after(table, paper_end, 1)
        paper_end += 1
        paper_data_rows_inserted += 1
    max_papers = max(0, paper_end - paper_start + 1)
    for i in range(max_papers):
        row_idx = paper_start + i
        row = table.rows[row_idx]
        if i >= len(papers):
            for c in (3, 5, 10, 15, 21):
                _set_cell_text(
                    row.cells[c],
                    "",
                    font_cn="宋体",
                    font_en="Times New Roman",
                    size_pt=10.5,
                    paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                )
                _paper_cell_center(row.cells[c])
            _keep_table_row_on_one_page(row)
            continue

        p = papers[i]
        _set_cell_text(
            row.cells[3],
            str(p.get("seq", i + 1)),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[3])

        # 作者信息：与 read_cell_marked 标记一致（加粗/下划线/组合）
        author_cell = row.cells[5]
        clear_cell(author_cell)
        ap = author_cell.paragraphs[0]
        _tight_paragraph(ap)
        for seg, b, u in parse_marked_authors(str(p.get("authors", ""))):
            append_run(ap, seg, font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=b, underline=u)
        _paper_cell_center(author_cell)

        _set_cell_text(
            row.cells[10],
            p.get("title", ""),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[10])

        # 期刊列：结构化字段与 pub_raw 解析合并后，按固定多行版式写入。
        journal_cell = row.cells[15]
        _fill_journal_cell(journal_cell, p)
        _paper_cell_center(journal_cell)

        # 期刊分级：每项一行。papers[] 支持可选 level_note / fold_note（多行用 \\n），例「折为1/2」「折合半篇」；
        # 从 Word 抽 JSON 时 batch_extract_json.parse_level_info 会把无法归入 JCR/校级/中科院/IF 的段落写入 level_note，回填时可还原多行。
        level_cell = row.cells[21]
        clear_cell(level_cell)
        levels = [p.get("jcr", ""), p.get("school_level", ""), p.get("cas_zone", ""), p.get("if", "")]
        note_raw = str(p.get("level_note") or p.get("fold_note") or "").strip()
        if note_raw:
            for ln in note_raw.splitlines():
                t = ln.strip()
                if t:
                    levels.append(t)
        non_empty = [x for x in levels if str(x).strip()]
        if not non_empty:
            non_empty = [""]
        for idx, ln in enumerate(non_empty):
            lp = level_cell.paragraphs[0] if idx == 0 else level_cell.add_paragraph()
            _tight_paragraph(lp)
            _paragraph_keep_lines_together(lp)
            append_run(lp, ln, font_cn="宋体", font_en="Times New Roman", size_pt=10.5)
        _paper_cell_center(level_cell)

        _keep_table_row_on_one_page(row)

    # compact 模式的实际删除推迟到所有内容填写完毕后执行（见 salary 填写后的 compact 块），
    # 避免行数变化导致后续 proj_start 等硬编码索引指向错误行。
    compact_paper_deleted = 0

    # 可选：通过增加论文区行高来推迟「推荐人意见」开始页（0~3 建议）
    paper_spacers = int(data.get("paper_row_spacers", 0) or 0)
    if paper_spacers > 0:
        _inflate_row_height_by_paragraphs(table.rows[paper_end], (3, 5, 10, 15, 21), min(paper_spacers, 3))

    # ----------------------------
    # 表格5（项目明细）表头 -> 大表 row 14（超论文行后顺延）
    # ----------------------------
    _proj_detail_cols = (3, 5, 9, 13, 16, 18, 20, 23)
    proj_header_row_idx = 14 + paper_data_rows_inserted
    for c in _proj_detail_cols:
        if c < len(table.rows[proj_header_row_idx].cells):
            _paper_cell_center(table.rows[proj_header_row_idx].cells[c])

    # ----------------------------
    # 表格5（项目明细）-> 大表 row 15 开始（超论文行后顺延）
    # ----------------------------
    projects = data.get("projects", [])
    proj_start = 15 + paper_data_rows_inserted
    proj_end = 17 + paper_data_rows_inserted
    proj_data_rows_inserted = 0  # 为容纳额外项目数据而插入的行数
    # 若项目数超过模板行数，先插入足够的行
    while len(projects) > proj_end - proj_start + 1:
        _insert_blank_rows_after(table, proj_end, 1)
        proj_end += 1
        proj_data_rows_inserted += 1
    max_projects = max(0, proj_end - proj_start + 1)
    for i in range(max_projects):
        row_idx = proj_start + i
        row = table.rows[row_idx]
        if i >= len(projects):
            for c in _proj_detail_cols:
                if c < len(row.cells):
                    _set_cell_text(
                        row.cells[c],
                        "",
                        font_cn="宋体",
                        font_en="Times New Roman",
                        size_pt=10.5,
                        paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    )
                    _paper_cell_center(row.cells[c])
            continue
        p = projects[i]
        _set_cell_text(
            row.cells[3],
            str(p.get("seq", i + 1)),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[3])
        _set_cell_text(
            row.cells[5],
            p.get("date", ""),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[5])
        _set_cell_text(
            row.cells[9],
            p.get("name", ""),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[9])
        _fill_project_institution_cell(row.cells[13], p.get("institution", ""))
        _set_cell_text(
            row.cells[16],
            p.get("category", ""),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[16])
        _set_cell_text(
            row.cells[18],
            p.get("level", ""),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[18])
        _set_cell_text(
            row.cells[20],
            p.get("fund", ""),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[20])
        _set_cell_text(
            row.cells[23],
            p.get("ranking", ""),
            font_cn="宋体",
            font_en="Times New Roman",
            size_pt=10.5,
            paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )
        _paper_cell_center(row.cells[23])

    # compact 模式的实际删除同样推迟（见后），此处仅占位。
    compact_proj_deleted = 0

    # 可选：通过增加项目区行高来推迟「推荐人意见」开始页（0~3 建议）
    proj_spacers = int(data.get("project_row_spacers", 0) or 0)
    if proj_spacers > 0:
        _inflate_row_height_by_paragraphs(table.rows[proj_end], _proj_detail_cols, min(proj_spacers, 3))

    # 项目总体情况：仅两行（与空白版一致），整行加粗、数字下划线
    pjs = data.get("projects_summary", {})
    fill_projects_summary_like_template(table.rows[13 + paper_data_rows_inserted].cells[1], pjs)

    # 通过增加「代表性论文」「科研项目」空白行来填充页面空白（用于将推荐人意见稳定推到第三页页首）
    extra_paper_rows = max(0, int(data.get("extra_paper_blank_rows", 0) or 0))
    extra_project_rows = max(0, int(data.get("extra_project_blank_rows", 0) or 0))
    if extra_paper_rows > 0:
        _insert_blank_rows_after(table, paper_end, extra_paper_rows)
        paper_blank_rows = []
        for i in range(extra_paper_rows):
            r = table.rows[paper_end + 1 + i]
            paper_blank_rows.append(r)
            for c in (3, 5, 10, 15, 21):
                if c < len(r.cells):
                    _set_cell_text(
                        r.cells[c],
                        "",
                        font_cn="宋体",
                        font_en="Times New Roman",
                        size_pt=10.5,
                        paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    )
                    _paper_cell_center(r.cells[c])
            _keep_table_row_on_one_page(r)
    else:
        paper_blank_rows = []

    # 插入论文空白行后，项目区起始索引顺延
    proj_start += extra_paper_rows
    proj_end += extra_paper_rows

    if extra_project_rows > 0:
        _insert_blank_rows_after(table, proj_end, extra_project_rows)
        proj_blank_rows = []
        for i in range(extra_project_rows):
            r = table.rows[proj_end + 1 + i]
            proj_blank_rows.append(r)
            for c in _proj_detail_cols:
                if c < len(r.cells):
                    _set_cell_text(
                        r.cells[c],
                        "",
                        font_cn="宋体",
                        font_en="Times New Roman",
                        size_pt=10.5,
                        paragraph_align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    )
                    _paper_cell_center(r.cells[c])
    else:
        proj_blank_rows = []

    # 视觉填充：将页底剩余空白均分到已插入的空白行高度上（不改变行数，仅增高空白行）
    if bool(data.get("auto_visual_fill_blank_rows", True)):
        total_blank = len(paper_blank_rows) + len(proj_blank_rows)
        if total_blank > 0:
            total_gap_pt = float(data.get("visual_fill_total_gap_pt", 0.0) or 0.0)
            ratio = float(data.get("visual_fill_ratio", 0.9) or 0.9)
            per_row_cap = float(data.get("visual_fill_per_row_cap_pt", 24.0) or 24.0)
            if total_gap_pt > 0:
                add_pt = min((total_gap_pt * ratio) / total_blank, per_row_cap)
                for r in paper_blank_rows + proj_blank_rows:
                    _bump_row_height(r, add_pt)

    # ----------------------------
    # 表格6（其他成果）-> 大表 row 18
    # ----------------------------
    # compact 删除在 salary 填写后才执行，此处不减；rec_start_row_idx 在 compact 块中单独更新。
    _row_delta = extra_paper_rows + extra_project_rows + proj_data_rows_inserted + paper_data_rows_inserted
    other_row_idx = 18 + _row_delta
    rec_start_row_idx = 19 + _row_delta
    pol_row_idx = 22 + _row_delta
    task_row_idx = 23 + _row_delta
    salary_row_idx = 24 + _row_delta

    other_cell = table.rows[other_row_idx].cells[3]
    other_raw = _normalize_other_achievements_structure(
        str(data.get("other", "")).strip()
    )
    applicant_name = str((data.get("basic_info") or {}).get("name", "")).strip()
    clear_cell(other_cell)
    if not other_raw:
        _set_cell_text(other_cell, "", font_cn="宋体", font_en="Times New Roman", size_pt=10.5)
    elif "\n" in other_raw:
        first_para = True
        for ln in other_raw.splitlines():
            text = ln.rstrip("\r")
            if not text.strip():
                continue
            p = other_cell.paragraphs[0] if first_para else other_cell.add_paragraph()
            first_para = False
            clear_paragraph_runs(p)
            plain_det = _other_plain_text_for_subtitle_detect(text)
            is_sub = _is_other_section_subtitle_line(plain_det)
            _apply_other_paragraph_layout(p, is_subtitle=is_sub)
            _other_append_line_runs(p, text, applicant_name, default_bold=is_sub)
    else:
        # 排除引文格式中的分号（「出版社;,年份:」），不作为列表分隔符
        parts = [x.strip() for x in re.split(r"[。；]|;(?!\s*,\s*\d{4}\s*:)", other_raw) if x.strip()]
        # 若第一段已有编号（「1.」「（1）」等）但后续段均无编号，说明「。」是句内标点而非条目分隔符，整段保持单条
        _NUM_PAT = re.compile(r"^(?:\d+[\.\．]|[（(]\s*\d+\s*[）)])")
        if (len(parts) >= 2
                and _NUM_PAT.match(parts[0])
                and not any(_NUM_PAT.match(p) for p in parts[1:])):
            parts = [other_raw.strip()]
        if not parts:
            _set_cell_text(other_cell, "", font_cn="宋体", font_en="Times New Roman", size_pt=10.5)
        elif len(parts) < 2:
            # 无换行且按句读拆分不足两段：不自动加「（1）」，整段按单行逻辑写入（避免「（1）2。」等）
            p = other_cell.paragraphs[0]
            clear_paragraph_runs(p)
            plain_det = _other_plain_text_for_subtitle_detect(other_raw)
            is_sub = _is_other_section_subtitle_line(plain_det)
            _apply_other_paragraph_layout(p, is_subtitle=is_sub)
            _other_append_line_runs(p, other_raw, applicant_name, default_bold=is_sub)
        else:
            for idx, txt in enumerate(parts, start=1):
                p = other_cell.paragraphs[0] if idx == 1 else other_cell.add_paragraph()
                clear_paragraph_runs(p)
                _tight_paragraph(p)
                end_mark = "。" if idx == len(parts) else "；"
                set_paragraph_format(p, first_line_chars=2, line_spacing=1.0)
                if not _other_fragment_already_numbered(txt):
                    append_run(
                        p,
                        f"（{idx}）",
                        font_cn="宋体",
                        font_en="Times New Roman",
                        size_pt=10.5,
                        bold=False,
                    )
                _other_append_line_runs(p, txt, applicant_name, default_bold=False)
                append_run(
                    p,
                    end_mark,
                    font_cn="宋体",
                    font_en="Times New Roman",
                    size_pt=10.5,
                    bold=False,
                )

    # ----------------------------
    # 表格7（推荐人意见）-> 大表 row 19~21
    # ----------------------------
    recs = data.get("recommendations", [])
    for i in range(3):
        row_idx = rec_start_row_idx + i
        row = table.rows[row_idx]
        cell = row.cells[1]
        clear_cell(cell)
        if i >= len(recs):
            continue
        r = recs[i]
        sub, body_paras = _recommendation_subtitle_and_body_paragraphs(r, i + 1)
        p0 = cell.paragraphs[0]
        clear_paragraph_runs(p0)
        _tight_paragraph(p0)
        if sub:
            set_paragraph_format(p0, first_line_chars=0, line_spacing=1.0)
            append_run(p0, sub, font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)
        started = bool(sub)
        for para_text in body_paras:
            if not para_text:
                continue
            if not started:
                p_body = p0
                started = True
            else:
                p_body = cell.add_paragraph()
                _tight_paragraph(p_body)
            set_paragraph_format(p_body, first_line_chars=2, line_spacing=1.0)
            append_run(p_body, para_text, font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=False)

        # 当启用“拆表+分页”强制新页时，不再在单元格内插分页，避免标题前出现空白段。
        if i == 0 and not bool(data.get("force_recommendations_new_page", True)):
            n_pb = int(data.get("recommendation_page_breaks", 2) or 2)
            n_pb = max(1, min(n_pb, 2))
            for _ in range(n_pb):
                add_page_break_before_paragraph(p0)

    # ----------------------------
    # 表格8（思想政治）-> 大表 row 22
    # ----------------------------
    # 左侧标题列保留「思想政治与品德综合评价」，正文写入右侧内容列
    pol_cell = table.rows[pol_row_idx].cells[1]
    pol_raw = str(data.get("political", "") or "").strip()
    pol_paras = split_political_paragraphs(pol_raw)
    clear_cell(pol_cell)
    if not pol_paras:
        _set_cell_text(pol_cell, "", font_cn="宋体", font_en="Times New Roman", size_pt=10.5)
    else:
        for idx, para in enumerate(pol_paras):
            pp = pol_cell.paragraphs[0] if idx == 0 else pol_cell.add_paragraph()
            _tight_paragraph(pp)
            # 与正确006 等原文一致：各段均首行空两格（按换行拆成多段后每段单独缩进）
            set_paragraph_format(pp, first_line_chars=2, line_spacing=1.0)
            append_run(pp, para, font_cn="宋体", font_en="Times New Roman", size_pt=10.5)
    _table_cell_vertical_center_paragraphs_left(pol_cell)

    # ----------------------------
    # 表格9（岗位任务/学院支持）-> 大表 row 23
    # 1. 岗位任务 / 2. 学院帮助…：之间已空一行；条目为全角圆括号（1），并规范 1）、(1) 等
    # ----------------------------
    task_cell = table.rows[task_row_idx].cells[1]
    clear_cell(task_cell)
    job_lines = [
        normalize_task_section_parens(x).strip()
        for x in str(data.get("job_tasks", "") or "").splitlines()
        if str(x).strip()
    ]
    sup_lines = [
        normalize_task_section_parens(x).strip()
        for x in str(data.get("college_support", "") or "").splitlines()
        if str(x).strip()
    ]
    p_head = task_cell.paragraphs[0]
    _tight_paragraph(p_head)
    set_paragraph_format(p_head, first_line_chars=0, line_spacing=1.0)
    append_run(p_head, "1. 岗位任务：", font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)
    for ln in job_lines:
        pj = task_cell.add_paragraph()
        _tight_paragraph(pj)
        is_sub_header = bool(re.match(r"^\d+[.．]\S", ln))
        set_paragraph_format(pj, first_line_chars=0 if is_sub_header else 2, line_spacing=1.0)
        append_run(pj, ln, font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=is_sub_header)
    task_cell.add_paragraph("")
    p_sec = task_cell.add_paragraph()
    _tight_paragraph(p_sec)
    set_paragraph_format(p_sec, first_line_chars=0, line_spacing=1.0)
    append_run(
        p_sec,
        "2. 学院帮助和支持其完成岗位任务的具体举措：",
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=10.5,
        bold=True,
    )
    for ln in sup_lines:
        ps = task_cell.add_paragraph()
        _tight_paragraph(ps)
        set_paragraph_format(ps, first_line_chars=2, line_spacing=1.0)
        append_run(ps, ln, font_cn="宋体", font_en="Times New Roman", size_pt=10.5)
    _table_cell_vertical_center_paragraphs_left(task_cell)

    # ----------------------------
    # 表格10（相关待遇）-> 大表 row 24
    # ----------------------------
    salary_cell = table.rows[salary_row_idx].cells[1]
    clear_cell(salary_cell)
    p = salary_cell.paragraphs[0]
    set_paragraph_format(p, first_line_chars=0, line_spacing=1.0)
    append_run(p, "年薪：", font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)
    _sal = str(data.get("salary", "") or "").strip()
    _sal_text = _sal if "万" in _sal else (f"{_sal} 万元/年" if _sal else "")
    if _sal_text and not _sal_text.endswith("；") and not _sal_text.endswith(";"):
        _sal_text += "；"
    append_run(
        p,
        _sal_text,
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=10.5,
    )
    append_run(p, "科研启动经费：", font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)
    _rf = str(data.get("research_fund", "") or "").strip()
    _rf_text = _rf if "万" in _rf else f"{_rf} 万元；"
    if _rf_text and not _rf_text.endswith("；") and not _rf_text.endswith(";"):
        _rf_text += "；"
    append_run(
        p,
        _rf_text,
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=10.5,
    )
    append_run(p, "安家费：", font_cn="宋体", font_en="Times New Roman", size_pt=10.5, bold=True)
    append_run(
        p,
        f"{data.get('relocation_fee', '')} 万元。",
        font_cn="宋体",
        font_en="Times New Roman",
        size_pt=10.5,
    )

    # ----------------------------
    # compact 模式：所有内容已填写完毕，现在统一删除空白行
    # 必须在所有 section 填写后执行，以保证 paper_start/proj_start 等索引在填写时不受干扰
    # ----------------------------
    if compact:
        # 1. 删除多余的论文空行（从后往前）
        if max_papers > len(papers):
            compact_paper_deleted = max_papers - max(len(papers), 1)
            for i in range(max_papers - 1, max(len(papers), 1) - 1, -1):
                tr = table.rows[paper_start + i]._tr
                tr.getparent().remove(tr)
        # 2. 删除多余的项目空行；此时 proj_start/proj_end 已含 extra_paper_rows 偏移，
        #    还须减去刚才删掉的论文行数（表行整体上移了 compact_paper_deleted 行）
        adj_proj_start = proj_start - compact_paper_deleted
        adj_proj_end = proj_end - compact_paper_deleted
        adj_max_projects = adj_proj_end - adj_proj_start + 1
        if adj_max_projects > len(projects):
            compact_proj_deleted = adj_max_projects - max(len(projects), 1)
            for i in range(adj_max_projects - 1, max(len(projects), 1) - 1, -1):
                tr = table.rows[adj_proj_start + i]._tr
                tr.getparent().remove(tr)
        # 3. 更新 rec_start_row_idx 供下方拆表使用
        rec_start_row_idx -= compact_paper_deleted + compact_proj_deleted

    # 强制推荐人意见从新页页首开始：在 row 19 前把大表硬拆分并插入分页符。
    # 说明：表格单元格内分页在 Word 中常被忽略，拆表是稳定可控方案。
    if bool(data.get("force_recommendations_new_page", True)):
        split_table_before_row_with_page_break(doc, table, rec_start_row_idx)

    # python-docx 改行距/缩进时可能留下 snapToGrid=1，与空白版不一致；保存前统一关闭。
    _document_snap_to_grid_off(doc)

    # 输出文件名：{姓名}_学术业绩简表.docx
    name = bi.get("name", "未命名")
    out = output_path or f"{name}_学术业绩简表.docx"
    try:
        doc.save(out)
        return out
    except PermissionError:
        # 文件被 Word 占用时，自动落到带时间戳的新文件名，避免再次冲突。
        alt = f"{name}_学术业绩简表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(alt)
        return alt


def _get_rec1_page_via_word_com(doc_path):
    """
    使用 Word COM 读取“1.张录法”所在页码（用于自动校验推荐人意见起始页）。
    返回 int，失败返回 None。
    """
    safe_path = str(Path(doc_path)).replace("'", "''")
    ps = (
        "$ErrorActionPreference='Stop';"
        f"$docPath = '{safe_path}';"
        "$word = New-Object -ComObject Word.Application;"
        "$word.Visible = $false;"
        "$doc = $word.Documents.Open($docPath, $false, $true);"
        "$doc.Repaginate();"
        "$rng = $doc.Content;"
        "$found = $rng.Find.Execute('1.张录法');"
        "if ($found) { $p = $rng.Information(3); Write-Output ('REC1_PAGE=' + $p) } else { Write-Output 'REC1_PAGE=0' };"
        "$doc.Close($false);"
        "$word.Quit();"
    )
    try:
        proc = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            check=False,
            timeout=25,
        )
        out = (proc.stdout or "") + "\n" + (proc.stderr or "")
        m = re.search(r"REC1_PAGE=(\d+)", out)
        return int(m.group(1)) if m else None
    except Exception:
        return None


def _get_layout_stats_via_word_com(doc_path):
    """
    返回版面统计：
    - rec_page: 第一位推荐人所在页
    - total_pages: 总页数
    - tbl1_end_page: 第一张表结束所在页
    - gap_to_bottom_pt: 第一张表末尾到页面正文底部的距离（pt）
    """
    # 优先使用 pywin32，避免 PowerShell 编码/转义导致的统计失真。
    if win32com is not None:
        try:
            app = win32com.client.Dispatch("Word.Application")
            app.Visible = False
            doc = app.Documents.Open(str(Path(doc_path).resolve()), False, True)
            doc.Repaginate()
            total_pages = int(doc.ComputeStatistics(2))
            if int(doc.Tables.Count) >= 2:
                rec_page = int(doc.Tables.Item(2).Range.Information(3))
            else:
                rec_page = 0
            tbl1 = doc.Tables.Item(1)
            er = tbl1.Range.Duplicate
            er.Collapse(0)
            tbl1_end_page = int(er.Information(3))
            y = float(er.Information(6))
            gap = float(doc.PageSetup.PageHeight - doc.PageSetup.BottomMargin - y)
            doc.Close(False)
            app.Quit()
            return {
                "rec_page": rec_page,
                "total_pages": total_pages,
                "tbl1_end_page": tbl1_end_page,
                "gap_to_bottom_pt": round(gap, 2),
            }
        except Exception:
            try:
                app.Quit()
            except Exception:
                pass

    safe_path = str(Path(doc_path)).replace("'", "''")
    ps = (
        "$ErrorActionPreference='Stop';"
        f"$docPath = '{safe_path}';"
        "$word = New-Object -ComObject Word.Application;"
        "$word.Visible = $false;"
        "$doc = $word.Documents.Open($docPath, $false, $true);"
        "$doc.Repaginate();"
        # 推荐人页码：优先取第二张表起始页（当前模板已将推荐区拆到第二张表）
        "if ($doc.Tables.Count -ge 2) { $rec = $doc.Tables.Item(2).Range.Information(3) } "
        "else { $rng = $doc.Content; $f = $rng.Find.Execute('推荐人意见'); if ($f) { $rec = $rng.Information(3) } else { $rec = 0 } };"
        "Write-Output ('REC_PAGE=' + $rec);"
        "Write-Output ('TOTAL=' + $doc.ComputeStatistics(2));"
        "$tbl = $doc.Tables.Item(1);"
        "$er = $tbl.Range.Duplicate;"
        "$er.Collapse(0);"
        "Write-Output ('TBL1_END_PAGE=' + $er.Information(3));"
        "$y = $er.Information(6);"
        "$gap = $doc.PageSetup.PageHeight - $doc.PageSetup.BottomMargin - $y;"
        "Write-Output ('GAP_PT=' + [math]::Round($gap, 2));"
        "$doc.Close($false);"
        "$word.Quit();"
    )
    try:
        proc = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            check=False,
            timeout=30,
        )
        out = (proc.stdout or "") + "\n" + (proc.stderr or "")
        m_rec = re.search(r"REC_PAGE=(\d+)", out)
        m_total = re.search(r"TOTAL=(\d+)", out)
        m_end = re.search(r"TBL1_END_PAGE=(\d+)", out)
        m_gap = re.search(r"GAP_PT=([-\d\.]+)", out)
        return {
            "rec_page": int(m_rec.group(1)) if m_rec else 0,
            "total_pages": int(m_total.group(1)) if m_total else 0,
            "tbl1_end_page": int(m_end.group(1)) if m_end else 0,
            "gap_to_bottom_pt": float(m_gap.group(1)) if m_gap else 0.0,
        }
    except Exception:
        return None


def incremental_tune_blank_rows(template_path, data_json, output_path=None):
    """
    逐行试加空白行：
    - 每次 +1 行后生成并校验
    - 约束：总页数=3，推荐人第3页，第一页表结束仍在第2页
    - 当剩余距离不足以再加一行（或再加会破坏约束）时停止
    """
    with open(data_json, "r", encoding="utf-8") as f:
        cfg = json.load(f)

    base = deepcopy(cfg)
    # 每次都从 0 开始自动寻优，避免“上次示例”的行数污染“本次示例”。
    cur_pap = 0
    cur_prj = 0
    max_rows = int(cfg.get("blank_row_max_rows", 20) or 20)
    order = cfg.get("blank_row_tune_order", "project_first")
    keys = ["extra_project_blank_rows", "extra_paper_blank_rows"] if order == "project_first" else ["extra_paper_blank_rows", "extra_project_blank_rows"]

    def _save_and_gen(pap, prj):
        cand = deepcopy(base)
        cand["extra_paper_blank_rows"] = max(0, pap)
        cand["extra_project_blank_rows"] = max(0, prj)
        with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False, encoding="utf-8") as tf:
            json.dump(cand, tf, ensure_ascii=False, indent=2)
            tmp_json = tf.name
        try:
            out = fill_template(template_path, tmp_json, output_path)
            st = _get_layout_stats_via_word_com(out)
            return out, st
        finally:
            try:
                Path(tmp_json).unlink(missing_ok=True)
            except Exception:
                pass

    # 先确保当前参数至少可用
    out, st = _save_and_gen(cur_pap, cur_prj)
    if not st:
        return out, cur_pap, cur_prj

    def valid(stat):
        return bool(
            stat
            and stat["rec_page"] == 3
            and stat["total_pages"] == 3
            and stat["tbl1_end_page"] == 2
        )

    # 贪心寻优：每次只尝试 +1 行（项目或论文），选“页底空白最小”的可行解。
    for _ in range(max_rows):
        candidates = []
        for k in keys:
            npap, nprj = cur_pap, cur_prj
            if k == "extra_project_blank_rows":
                nprj += 1
            else:
                npap += 1
            out2, st2 = _save_and_gen(npap, nprj)
            if valid(st2):
                candidates.append((st2["gap_to_bottom_pt"], npap, nprj, out2, st2))

        if not candidates:
            break
        candidates.sort(key=lambda x: x[0])  # gap 越小越接近“视觉无空白”
        best_gap, bp, br, bout, bst = candidates[0]
        # 没有继续改善则停止
        if best_gap >= st["gap_to_bottom_pt"]:
            break
        cur_pap, cur_prj, out, st = bp, br, bout, bst

    cfg["extra_paper_blank_rows"] = cur_pap
    cfg["extra_project_blank_rows"] = cur_prj
    # 把当前页底剩余距离传给填充阶段，用于拉高空白行（视觉更接近“无空白”）
    if st and (cur_pap + cur_prj) > 0:
        cfg["visual_fill_total_gap_pt"] = max(0.0, float(st.get("gap_to_bottom_pt", 0.0)))
    with open(data_json, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)
    final_out = fill_template(template_path, data_json, output_path)
    return final_out, cur_pap, cur_prj


def auto_tune_blank_rows(template_path, data_json, output_path=None):
    """
    自动校验并调优：
    - 目标1：推荐人意见落在第3页（首位推荐人所在页 == 3）
    - 目标2：在不跳到第4页前提下，尽量增大论文/项目空白行，减少第2页末空白
    """
    with open(data_json, "r", encoding="utf-8") as f:
        base = json.load(f)

    best = None
    max_total = int(base.get("auto_tune_max_total_blank_rows", 8) or 8)
    max_total = max(0, min(max_total, 10))

    # 快速寻优：每个 total 只测 3 组分配，避免 Word COM 反复打开导致长时间卡住
    for total in range(max_total, -1, -1):
        candidates = []
        # 项目优先（更接近用户视觉预期）
        candidates.append((0, total))
        # 均分
        candidates.append((total // 2, total - total // 2))
        # 论文优先
        candidates.append((total, 0))
        # 去重保持顺序
        uniq = []
        seen = set()
        for pap, prj in candidates:
            if (pap, prj) not in seen:
                uniq.append((pap, prj))
                seen.add((pap, prj))

        for pap, prj in uniq:
            cand = deepcopy(base)
            cand["extra_paper_blank_rows"] = pap
            cand["extra_project_blank_rows"] = prj
            with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False, encoding="utf-8") as tf:
                json.dump(cand, tf, ensure_ascii=False, indent=2)
                tmp_json = tf.name
            try:
                out = fill_template(template_path, tmp_json, output_path)
                page = _get_rec1_page_via_word_com(out)
            finally:
                try:
                    Path(tmp_json).unlink(missing_ok=True)
                except Exception:
                    pass
            if page == 3:
                best = (pap, prj, out)
                break
        if best:
            break

    # 若全部候选都失败，回退到当前配置直接生成
    if not best:
        out = fill_template(template_path, data_json, output_path)
        return out, None, None

    pap, prj, out = best
    # 将最优参数写回数据文件，后续直接复用
    base["extra_paper_blank_rows"] = pap
    base["extra_project_blank_rows"] = prj
    with open(data_json, "w", encoding="utf-8") as f:
        json.dump(base, f, ensure_ascii=False, indent=2)
    return out, pap, prj


def tune_and_fill_one_command(template_path, data_json, output_path=None, max_steps=24):
    """
    一键模式：
    1) 从 0/0 空白行开始自动寻优
    2) 目标：总页数=3、推荐人第3页、第一页表结束在第2页、且页底空白最小
    3) 回写最佳参数并输出最终文档
    """
    with open(data_json, "r", encoding="utf-8") as f:
        cfg = json.load(f)

    base = deepcopy(cfg)
    base["force_recommendations_new_page"] = True
    base["auto_incremental_blank_rows"] = False
    base["auto_tune_page2_blank"] = False
    base["extra_paper_blank_rows"] = 0
    base["extra_project_blank_rows"] = 0

    def _gen_with(pap, prj):
        cand = deepcopy(base)
        cand["extra_paper_blank_rows"] = max(0, int(pap))
        cand["extra_project_blank_rows"] = max(0, int(prj))
        with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False, encoding="utf-8") as tf:
            json.dump(cand, tf, ensure_ascii=False, indent=2)
            tmp_json = tf.name
        try:
            out = fill_template(template_path, tmp_json, output_path)
            st = _get_layout_stats_via_word_com(out)
            return out, st
        finally:
            try:
                Path(tmp_json).unlink(missing_ok=True)
            except Exception:
                pass

    def _valid(st):
        return bool(st and st["rec_page"] == 3 and st["total_pages"] == 3 and st["tbl1_end_page"] == 2)

    cur_pap, cur_prj = 0, 0
    out, st = _gen_with(cur_pap, cur_prj)
    best = (cur_pap, cur_prj, st["gap_to_bottom_pt"] if st else 1e9, out, st)

    for _ in range(max_steps):
        candidates = []
        for target in ("project", "paper"):
            npap, nprj = cur_pap, cur_prj
            if target == "project":
                nprj += 1
            else:
                npap += 1
            out2, st2 = _gen_with(npap, nprj)
            if _valid(st2):
                candidates.append((st2["gap_to_bottom_pt"], npap, nprj, out2, st2))
        if not candidates:
            break
        candidates.sort(key=lambda x: x[0])
        gap, pap, prj, out2, st2 = candidates[0]
        if gap >= best[2]:
            break
        cur_pap, cur_prj = pap, prj
        best = (pap, prj, gap, out2, st2)

    cfg["extra_paper_blank_rows"] = best[0]
    cfg["extra_project_blank_rows"] = best[1]
    cfg["force_recommendations_new_page"] = True
    cfg["auto_incremental_blank_rows"] = False
    cfg["auto_tune_page2_blank"] = False
    with open(data_json, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)

    # 若最终仍超过3页，启用 compact 模式（删除无数据的论文/项目空行）再试一次
    final_out = fill_template(template_path, data_json, output_path)
    final_st = _get_layout_stats_via_word_com(final_out)
    if final_st and final_st["total_pages"] > 3:
        final_out = fill_template(template_path, data_json, output_path, compact=True)
        final_st = _get_layout_stats_via_word_com(final_out)
    return final_out, best[0], best[1], final_st


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="一键读取JSON并输出调好版式的Word")
    parser.add_argument("--template", default="调整空白版.docx")
    parser.add_argument("--json", dest="data_file", default="test_data.json")
    parser.add_argument("--output", dest="out_file", default=None)
    parser.add_argument("--max-steps", type=int, default=24)
    parser.add_argument("--no-tune", action="store_true", help="仅填充，不自动调优空白行")
    args = parser.parse_args()

    if args.no_tune:
        result = fill_template(args.template, args.data_file, args.out_file)
        print(f"已生成：{result}")
    else:
        result, pap, prj, st = tune_and_fill_one_command(args.template, args.data_file, args.out_file, args.max_steps)
        if st:
            print(
                f"已生成：{result}（一键调优：extra_paper_blank_rows={pap}, extra_project_blank_rows={prj}, "
                f"rec_page={st['rec_page']}, total={st['total_pages']}, gap={st['gap_to_bottom_pt']}）"
            )
        else:
            print(f"已生成：{result}（一键调优完成：extra_paper_blank_rows={pap}, extra_project_blank_rows={prj}）")

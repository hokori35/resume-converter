import copy
import json
import re
from pathlib import Path
from typing import Tuple

from docx import Document


SCHEMA_TEMPLATE = {
    "basic_info": {
        "name": "", "gender": "", "hometown": "", "marital_status": "",
        "birth": "", "current_title": "", "discipline": "", "religion": "",
        "target_college": "",
    },
    "education": [], "work": [], "overseas": [],
    # 「注：…」「补充说明」等：无论出现在学习/工作/海外哪一段，均不入嵌套表，回填在海外表格外最后
    "experience_notes": [],
    "research_area": "", "research_team": "", "team_achievement": "",
    "papers_summary": {
        "total": 0,
        "jcr_q1": 0, "jcr_q1_first": 0,
        "jcr_q2": 0, "jcr_q2_first": 0,
        "jcr_q3": 0, "jcr_q3_first": 0,
        "jcr_q4": 0, "jcr_q4_first": 0,
        # JCR 分区统计块内的「其他」（与校级「其他」分列；仅填抽取值，回填不做算术推算）
        "jcr_other": 0,
        "jcr_other_first": 0,
        "school_a": 0, "school_a_first": 0,
        "school_b": 0, "school_b_first": 0,
        "school_c": 0, "school_c_first": 0,
        "school_a_star": 0, "school_a_star_first": 0,
        "school_other": 0, "school_other_first": 0,
        # 中国计算机学会推荐目录（选填；与 fill_template「按中国计算机学会分类统计」对应）
        "ccf_a": 0, "ccf_a_first": 0,
        "ccf_b": 0, "ccf_b_first": 0,
        "ccf_c": 0, "ccf_c_first": 0,
        "ccf_other": 0, "ccf_other_first": 0,
        # ABS 期刊分级（选填；与「按 ABS 分类统计」对应）
        "abs_4": 0, "abs_4_first": 0,
        "abs_3": 0, "abs_3_first": 0,
        "abs_2": 0, "abs_2_first": 0,
        "abs_1": 0, "abs_1_first": 0,
        "abs_other": 0, "abs_other_first": 0,
    },
    "papers": [],
    "projects_summary": {
        "host_a_major": 0, "host_a_key": 0, "host_a": 0, "host_b": 0,
        "participate_a_major": 0, "participate_a_key": 0, "participate_a": 0, "participate_b": 0,
    },
    "projects": [],
    "recommendation_page_breaks": 2,
    "paper_row_spacers": 0,
    "project_row_spacers": 0,
    "force_recommendations_new_page": True,
    "auto_incremental_blank_rows": False,
    "blank_row_tune_order": "project_first",
    "blank_row_min_gap_pt": 8,
    "blank_row_max_rows": 16,
    "auto_visual_fill_blank_rows": True,
    "visual_fill_ratio": 0.95,
    "visual_fill_per_row_cap_pt": 28,
    "visual_fill_total_gap_pt": 0.0,
    "auto_tune_page2_blank": False,
    "auto_tune_max_total_blank_rows": 8,
    "extra_paper_blank_rows": 0,
    "extra_project_blank_rows": 5,
    "other": "",
    "recommendations": [],
    "political": "",
    "job_tasks": "",
    "college_support": "",
    "salary": "",
    "research_fund": "",
    "relocation_fee": "",
}

# 行内已知标签，用于 other 行内容过滤（Word 模板可能用「其它」作小节标题）
_SECTION_LABELS = {
    "其他", "其它", "近5年代表性成果", "近5年代表性成果|",
    "代表性论文", "科研项目", "推荐人意见",
}


def consec_unique(row):
    """连续去重：只去掉相邻重复单元格，保留同行不同位置的相同值（如两个'无'）。"""
    out, prev = [], None
    for cell in row.cells:
        t = cell.text.replace("\n", "|").strip()
        if t and t != prev:
            out.append(t)
            prev = t
    return out


def row_cells_all(row):
    """
    逐格读取表格行文本（含空字符串），列位置与 Word 一致。
    「科研项目」等表若用 consec_unique，相邻空单元格会被跳过，导致项目级别为空时
    后续经费、排序列整体左移错位。
    """
    return [c.text.replace("\n", "|").strip() for c in row.cells]


def row_cells_collapse_merged_duplicates_non_empty(values):
    """
    Word 横向合并后，python-docx 同行连续多格文本完全相同；去掉相邻重复（仅非空），
    以恢复「序号 / 日期 / 项目名称…」逻辑列，避免把同一日期当成多列。
    空单元格不去重，避免丢失「项目级别」空列。
    """
    out = []
    for t in values:
        if out and t and t == out[-1]:
            continue
        out.append(t)
    return out


def _project_tail_compress_empty_runs(tail):
    """tail 中连续空串合并为一个（模板常见「级别」空占多格）。"""
    out = []
    prev_empty = False
    for x in tail:
        empty = not str(x).strip()
        if empty:
            if not prev_empty:
                out.append("")
            prev_empty = True
        else:
            out.append(x)
            prev_empty = False
    return out


def consec_unique_with_cells(row):
    """同 consec_unique，但同时返回 (text, cell_obj) 对，用于需要读取 run 格式的场景。"""
    out, prev = [], None
    for cell in row.cells:
        t = cell.text.replace("\n", "|").strip()
        if t and t != prev:
            out.append((t, cell))
            prev = t
    return out


def read_cell_marked(cell):
    """
    读取 cell 中所有 run，按样式编码为标记，供 fill_template 原样还原（与「正确」Word 对齐）。
    相邻且样式相同的 run 会合并，避免 Word 拆 run 导致 JSON 碎片化。
    多段合并为一段（段间无换行），与历史行为一致；多段需保留换行请用 read_cell_marked_paragraphs。
    - 加粗+下划线 -> [**text**]{.underline}
    - 仅加粗 -> [**text**]
    - 仅下划线 -> [[text]]{.underline}
    - 其余 -> 原文
    """
    return "".join(_paragraph_to_marked(p) for p in cell.paragraphs)


def read_cell_marked_paragraphs(cell):
    """多段单元格：段之间用 \\n，段内样式同 read_cell_marked（用于「其他成果」等）。"""
    lines = [_paragraph_to_marked(p) for p in cell.paragraphs]
    return "\n".join(lines).strip()


def _extract_other_achievements_from_row(row):
    """
    「其他成果」行：横向多格时不能只取第一个非空格（常为孤立序号「2」），须合并正文格。
    去掉格首「其他/其它」标签；合并后丢弃在存在长正文时的纯「n」「n.」短噪格。
    """
    texts = row_cells_all(row)
    cells = list(row.cells)
    n = min(len(texts), len(cells))

    def only_labels(tx):
        parts = [p.strip() for p in (tx or "").split("|") if p.strip()]
        return not parts or all(p in _SECTION_LABELS for p in parts)

    chunks = []
    seen = set()
    for i in range(n):
        tx = (texts[i] or "").strip()
        if not tx or only_labels(tx):
            continue
        body = read_cell_marked_paragraphs(cells[i]).strip()
        if not body:
            continue
        body = re.sub(
            r"^\s*(?:\[\*\*)?(其他|其它)(?![一-鿿])(?:\*\*\])?\s*[：:\n|]?\s*",
            "",
            body,
            count=1,
        ).strip()
        if not body:
            continue
        if body in seen:
            continue
        seen.add(body)
        chunks.append(body)

    if not chunks:
        return ""
    mx = max(len(c) for c in chunks)
    if mx > 80:
        chunks = [
            c
            for c in chunks
            if len(c) > 40 or not re.fullmatch(r"\d+\.?\s*", c.strip())
        ]
    return "\n".join(chunks).strip()


def _paragraph_to_marked(para):
    """将段落转为标记文本，支持直接 <w:r> 和 <w:hyperlink> 内嵌 <w:r>（超链接内的 runs）。"""
    from docx.oxml.ns import qn

    chunks = []

    def _rpr_style(rpr):
        """从 rPr 元素读取 bold/underline，返回样式键。"""
        if rpr is None:
            return "p"
        b_el = rpr.find(qn("w:b"))
        u_el = rpr.find(qn("w:u"))
        b = b_el is not None and b_el.get(qn("w:val"), "1").lower() not in ("false", "0", "off")
        u = u_el is not None and u_el.get(qn("w:val"), "single").lower() not in ("none", "false", "0")
        if b and u:
            return "bu"
        if b:
            return "b"
        if u:
            return "u"
        return "p"

    def _process_run_el(r_el):
        t_el = r_el.find(qn("w:t"))
        if t_el is None:
            return
        text = t_el.text or ""
        if not text:
            return
        st = _rpr_style(r_el.find(qn("w:rPr")))
        if chunks and chunks[-1][0] == st:
            chunks[-1] = (st, chunks[-1][1] + text)
        else:
            chunks.append((st, text))

    for child in para._p:
        if child.tag == qn("w:r"):
            _process_run_el(child)
        elif child.tag == qn("w:hyperlink"):
            for r_el in child.findall(qn("w:r")):
                _process_run_el(r_el)

    parts = []
    for st, text in chunks:
        if st == "bu":
            parts.append(f"[**{text}**]{{.underline}}")
        elif st == "b":
            parts.append(f"[**{text}**]")
        elif st == "u":
            parts.append(f"[[{text}]]{{.underline}}")
        else:
            parts.append(text)
    return "".join(parts)


def _pick_other_achievements_row_index(
    row_texts, t, *, paper_hdr=-1, paper_body_end=-1, proj_hdr=-1
):
    """
    定位「其他成果」所在行。

    策略分两级：
    1. 行内有「其他成果」明文，或「其他|1.」结构 → 立即返回（无歧义）。
    2. 用单元格结构判断：第一个不重复格 = 「其他/其它」→ 返回该行。
       此条件基于表格布局而非文本内容，天然排除两类假阳性：
       - 科研项目表头行：「其他」在中间列，不是第一格；
       - 论文统计行：「其他」嵌于长文本内，不独占首格。
       因此不再需要针对具体内容特征的启发式过滤链，避免新模板触发新的误排除。

    paper_hdr / paper_body_end：论文表数据区内的行跳过（避免误把引文格误判为「其他」标签）。
    proj_hdr：保留参数签名兼容性，新逻辑不再需要该参数。
    """
    if not row_texts:
        return -1

    # 级别1：行内有明确「其他成果」文本或「其他|1.」结构，立即返回
    for i, txt in enumerate(row_texts):
        if "其他成果" in txt or "其它成果" in txt:
            return i
        if re.search(r"(?:^|\|)(?:其他|其它)\|1\.", txt):
            return i

    # 级别2：前3个不重复格之一恰好是「其他/其它」（含竖排「其|他」形式）
    # 之所以检查前3格而非仅第一格：「近5年代表性成果」为跨行合并的大标签，占据第一格，
    # 「其他」子节标签通常在第二格；但无大标签的简单模板中「其他」也可能在第一格。
    # 假阳性排除：科研项目表头的「其他」列出现在第4格以后（前面有「序号/起止时间/项目名称」等）；
    # 论文统计区的「其他」嵌于长文本内（不精确匹配）；论文数据区由 _in_paper_band 排除。
    _other_label = re.compile(r"^(?:其他|其它|其\|他|其\|它)\s*$")

    def _in_paper_band(idx):
        return (
            paper_hdr >= 0
            and paper_body_end > paper_hdr + 1
            and paper_hdr < idx < paper_body_end
        )

    for i in range(5, len(row_texts)):
        if _in_paper_band(i):
            continue
        cells = consec_unique(t.rows[i])
        if any(_other_label.match(c.strip()) for c in cells[:3]):
            return i

    return -1


def find_after(label, items, stops=()):
    """在 items 中找第一个包含 label 子串的项，返回其后第一个非空项。
    stops：若下一个非空项本身是已知字段标签，则视为无值返回空串。"""
    stop_norms = {_norm_basic_label(x) for x in stops}
    for i, item in enumerate(items):
        if label in item:
            for j in range(i + 1, len(items)):
                if items[j].strip():
                    if stop_norms and _norm_basic_label(items[j]) in stop_norms:
                        return ""
                    return items[j]
    return ""


def _norm_basic_label(s):
    return re.sub(r"[\s　]+", "", s or "")


# 基本信息第 0 行：某标签后的「下一项标题」出现时，说明本项无独立取值格
_BASIC_R0_STOPS = {
    "姓  名": ("性别", "籍贯", "婚姻状况", "出生年月"),
    "姓名": ("性别", "籍贯", "婚姻状况", "出生年月"),
    "性别": ("籍贯", "婚姻状况", "出生年月"),
    "籍贯": ("婚姻状况", "出生年月"),
    "婚姻状况": ("出生年月",),
    "出生年月": (),
}

# 第 1 行：职称、学科、宗教、拟聘
_BASIC_R1_STOPS = {
    "一级学科": ("宗教信仰", "拟聘学院", "拟聘单位", "现职称"),
    "现职称": ("一级学科", "宗教信仰", "拟聘学院", "拟聘单位"),
    "宗教信仰": ("拟聘学院", "拟聘单位"),
    "拟聘学院": ("拟聘单位",),
    "拟聘单位": (),
}


def find_after_basic_cell(label, items, next_headers, *, max_scan=24):
    """
    用于 row_cells_all 逐格列表：合并列会产生多格与 label 文本完全相同的「标题格」。
    - 默认：下一非空格若已是下一字段标题（stops），则本项无取值，返回空串。
    - 「婚姻状况」「宗教信仰」：下一格可能是「出生年月」「拟聘学院」等标题，取值格在其后，
      故遇 stops 时继续扫描（最多 max_scan 个非空格），与旧版 consec_unique 行为对齐。
    """
    lab_n = _norm_basic_label(label)
    stops = {_norm_basic_label(x) for x in (next_headers or ())}
    flex = lab_n in {
        _norm_basic_label("婚姻状况"),
        _norm_basic_label("宗教信仰"),
    }

    for i, item in enumerate(items):
        if not item:
            continue
        if lab_n not in _norm_basic_label(item):
            continue
        non_empty_seen = 0
        for j in range(i + 1, len(items)):
            c = (items[j] or "").strip()
            if not c:
                continue
            non_empty_seen += 1
            if non_empty_seen > max_scan:
                break
            cn = _norm_basic_label(c)
            if cn == lab_n:
                continue
            if cn in stops:
                if flex:
                    continue
                return ""
            return items[j]
    return ""


def _current_title_value_is_placeholder(s: str) -> bool:
    """合并格后下一格仍是表头模版、无真实职称时视为空。"""
    raw = (s or "").replace("|", "\n").strip()
    if not raw:
        return True
    title_norms = {
        _norm_basic_label(x)
        for x in (
            "现职称",
            "（评审时间）",
            "现职称（评审时间）",
            "(评审时间)",
        )
    }
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    if not lines:
        return True
    return all(_norm_basic_label(ln) in title_norms for ln in lines)


def extract_current_title_from_row(row):
    """
    「现职称（评审时间）」取值：
    1) 用 row_cells_all 向后逐格扫描：跳过标签重复格、模版占位格（_current_title_value_is_placeholder），
       遇到下一字段标题（stops）则停止，取第一个真实值（如「讲师，特聘副教授（2021年10月）」）。
       此策略解决「（评审时间）」占位格夹在标签与真实值之间导致漏提的问题（错误004）。
    2) 仅当上述为空时，再读「现职称」所在单元格内多段正文（值嵌入同一格的情形）。
    原文有内容则写入 JSON；原文空白则 current_title 为空串（回填不填「无」）。
    """
    r1_all = row_cells_all(row)
    stops = _BASIC_R1_STOPS["现职称"]
    stops_n = {_norm_basic_label(x) for x in stops}
    nl = _norm_basic_label("现职称")
    nl2 = _norm_basic_label("现职称（评审时间）")

    # Step 1: 向后扫描，跳过标签重复格与占位格，取第一个真实值
    for i, item in enumerate(r1_all):
        if not item:
            continue
        n = _norm_basic_label(item)
        if nl not in n and nl2 not in n:
            continue
        for j in range(i + 1, len(r1_all)):
            c = (r1_all[j] or "").strip()
            if not c:
                continue
            cn = _norm_basic_label(c)
            if nl in cn or nl2 in cn:
                continue  # 合并格标签重复
            if cn in stops_n:
                break  # 已到下一字段标题，无取值
            if _current_title_value_is_placeholder(c):
                continue  # 跳过模版占位格
            return c.replace("|", "\n")
        break  # 找到标签但无取值格

    # Step 2: 段落回退——值嵌入「现职称」同一单元格内
    idx = None
    for i, t in enumerate(r1_all):
        if nl in _norm_basic_label(t):
            idx = i
            break
    if idx is None:
        return ""
    cell = row.cells[idx]
    lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    body_lines = []
    skip_head = frozenset(
        ("现职称", "（评审时间）", "现职称（评审时间）", "(评审时间)")
    )
    for ln in lines:
        s = ln.strip()
        if not s or s in skip_head:
            continue
        if _norm_basic_label(s) in {nl, nl2}:
            continue
        body_lines.append(ln)
    body = "\n".join(body_lines).strip()
    return body.replace("|", "") if body else ""


def normalize_piped_table_text(s, joiner="、"):
    """
    Word 表格同行多格经 consec_unique 会变成用 | 拼接；写入 JSON 时改为 joiner 连接。
    joiner 默认顿号；拟加入科研团队等需无缝拼接时传 joiner=""。
    若前一段以连字符结尾（如「调和分析-」），则无缝拼接而不插入 joiner。
    """
    s = (s or "").strip()
    if not s or "|" not in s:
        return s
    parts = [p.strip() for p in s.split("|") if p.strip()]
    if not parts:
        return s
    result = parts[0]
    for part in parts[1:]:
        if result.endswith(("-", "–", "—")):
            result += part
        else:
            result += joiner + part
    return result


def _to_num(s):
    """将数字字符串转为 int（整数）或 float（小数）。"""
    try:
        f = float(s)
        return int(f) if f == int(f) else f
    except Exception:
        return 0


def _is_edu_work_supplementary_line(line):
    """
    学习及工作简历中的附注行：不入经历表，归入 experience_notes，回填时写在海外嵌套表下方。
    与 fill_template._is_overseas_supplementary_note 判定一致。
    """
    s = (line or "").strip()
    if not s:
        return False
    if re.match(r"^[\s\u3000]*注\s*[：:]", s):
        return True
    if "补充说明" in s:
        return True
    if s.startswith(("（", "(")) and ("补充" in s or "不连续" in s):
        return True
    if "须完整" in s or "须连续" in s:
        return True
    if re.search(r"学习及工作经历连续|经历连续", s):
        return True
    if re.search(r"临时工作", s):
        return True
    return False


_EXP_DATE_SEP = r"[-～–—－]"


def normalize_experience_date_jin(text):
    """
    学习/工作/海外经历：结束写「今」时规范为「至今」，如 2022.09-今 → 2022.09-至今。
  已是「至今」的不改。
    """
    s = re.sub(r"\s+", " ", (text or "").strip())
    s = re.sub(
        rf"(\d{{4}}\.\d{{1,2}})\s*{_EXP_DATE_SEP}\s*今(?!至)",
        r"\1-至今",
        s,
    )
    s = re.sub(r"(\d{4}\.\d{1,2})\s+今(?!至)", r"\1-至今", s)
    s = re.sub(r"(\d{4}\.\d{1,2})\s+至今", r"\1-至今", s)
    return s


def parse_edu_work_overseas(text):
    sec = {"education": [], "work": [], "overseas": [], "experience_notes": []}
    marker = None
    for ln in [x.strip() for x in text.replace("：", ":").split("|") if x.strip()]:
        if "学习经历" in ln:
            marker = "education"; continue
        if "工作经历" in ln:
            marker = "work"; continue
        if "海外经历" in ln:
            marker = "overseas"; continue
        if marker:
            n = normalize_experience_date_jin(ln)
            if not n:
                continue
            if _is_edu_work_supplementary_line(n):
                sec["experience_notes"].append(n)
                marker = "experience_notes"  # 注：之后的条目也归入附注
                continue
            sec[marker].append(n)
    for k in ("education", "work", "overseas"):
        if not sec[k]:
            sec[k] = ["无"]
    return sec


# 匹配带可选小数的篇数，如 "3.5 篇" 或 "3 篇"（解析前会将全角数字转为半角）
_NUM = r"[\d.]+"

# 科研项目表体行常见「起止时间」形态（与 extract_one 内 _proj_date_re 一致）
_PROJ_DATE_RANGE_BODY = re.compile(
    r"\d{4}[.\-]\d{2}(?:[.\-]\d{2})?\s*[-–—至]\s*(?:\d{4}[.\-]\d{2}(?:[.\-]\d{2})?|至今)?"
)


def _looks_like_project_table_body_row(txt):
    """合并单元格里含「|其他|」列名时，勿把项目数据行误判为「其他成果」行。"""
    t = txt or ""
    if not _PROJ_DATE_RANGE_BODY.search(t):
        return False
    compact = re.sub(r"\s+", "", t)
    # 排名「1/1」「3/8」；经费格常只有数字「50」，「万」可能在表头行而不在本行合并文本里（错误009）
    if re.search(r"\d+/\d+", compact):
        return True
    if "万" in t and re.search(r"(纵向|横向|国家级|省部级|校级|课题|基金)", t):
        return True
    return False

_FW_DIGIT_TRANS = str.maketrans("０１２３４５６７８９", "0123456789")


def _normalize_summary_text_for_parse(text):
    """全角数字→半角，避免「其他：１ 篇」等无法被 [\\d.]+ 匹配。"""
    return (text or "").translate(_FW_DIGIT_TRANS)


def _sum_num_expression(tok):
    """
    括号内篇数如「6.5+0.5」「3 ＋ 0.5」按加号求和；无加号则按单数解析。
    """
    t = re.sub(r"\s+", "", (tok or "").strip()).replace("＋", "+")
    if not t:
        return 0
    if "+" in t:
        parts = [p for p in t.split("+") if p.strip()]
        try:
            return round(sum(float(p) for p in parts), 10)
        except ValueError:
            return _to_num(t)
    return _to_num(t)


def parse_papers_summary(text):
    text = _normalize_summary_text_for_parse(text)
    s = {k: 0 for k in SCHEMA_TEMPLATE["papers_summary"]}
    m = re.search(
        rf"(?:发表论文|第一作者或通讯作者发表论文)\s*[（(]\s*([\d.\s＋+]+?)\s*篇[）)]",
        text,
    )
    if m:
        s["total"] = _sum_num_expression(m.group(1))

    sp_school = re.search(r"按\s*学校期刊目录分级统计", text)
    sp_ccf = re.search(r"按\s*中国计算机学会分类统计", text)
    sp_abs = re.search(r"按\s*ABS\s*分类统计", text)

    jcr_text = text[: sp_school.start()] if sp_school else text

    end_school = len(text)
    if sp_school:
        if sp_ccf and sp_ccf.start() > sp_school.start():
            end_school = min(end_school, sp_ccf.start())
        if sp_abs and sp_abs.start() > sp_school.start():
            end_school = min(end_school, sp_abs.start())
        school_only = text[sp_school.start() : end_school]
    else:
        school_only = ""

    end_ccf = len(text)
    if sp_ccf:
        if sp_abs and sp_abs.start() > sp_ccf.start():
            end_ccf = sp_abs.start()
        ccf_only = text[sp_ccf.start() : end_ccf]
    else:
        ccf_only = ""

    abs_only = text[sp_abs.start() :] if sp_abs else ""

    _OTHER_LABEL = r"(?:其\s*他|其他)"
    _PAREN_FIRST_COMM = (
        r"[（(]\s*(@@A@@)\s*篇为第一作者[，,、]\s*(@@B@@)\s*篇为通讯作者\s*[）)]"
    )

    for q in [1, 2, 3, 4]:
        m2 = re.search(
            rf"Q{q}\s*区[：:]\s*({_NUM})\s*篇"
            + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
            jcr_text,
            re.I,
        )
        m1 = re.search(
            rf"Q{q}\s*区[：:]\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            jcr_text,
            re.I,
        )
        m0 = re.search(
            rf"Q{q}\s*区[：:]\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            jcr_text,
            re.I,
        )
        m_bare = re.search(
            rf"Q{q}\s*区[：:]\s*({_NUM})\s*篇",
            jcr_text,
            re.I,
        )
        if m2:
            s[f"jcr_q{q}"] = _to_num(m2.group(1))
            s[f"jcr_q{q}_first"] = _to_num(m2.group(2))
        elif m1:
            s[f"jcr_q{q}"] = _to_num(m1.group(1))
            s[f"jcr_q{q}_first"] = _to_num(m1.group(2))
        elif m0:
            s[f"jcr_q{q}"] = _to_num(m0.group(1))
            s[f"jcr_q{q}_first"] = ""
        elif m_bare:
            s[f"jcr_q{q}"] = _to_num(m_bare.group(1))
            s[f"jcr_q{q}_first"] = ""

    # JCR 统计块内的「其他」（同一 jcr_text 内；兼容 Word 中「其 他」断字、数字与「篇」紧贴如 2.5篇）
    m_jo2 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇"
        + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
        jcr_text,
        re.I,
    )
    m_jo = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        jcr_text,
        re.I,
    )
    m_jo0 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        jcr_text,
        re.I,
    )
    m_jo_bare = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇",
        jcr_text,
        re.I,
    )
    if m_jo2:
        s["jcr_other"] = _to_num(m_jo2.group(1))
        s["jcr_other_first"] = _to_num(m_jo2.group(2))
    elif m_jo:
        s["jcr_other"] = _to_num(m_jo.group(1))
        s["jcr_other_first"] = _to_num(m_jo.group(2))
    elif m_jo0:
        s["jcr_other"] = _to_num(m_jo0.group(1))
        s["jcr_other_first"] = ""
    elif m_jo_bare:
        s["jcr_other"] = _to_num(m_jo_bare.group(1))
        s["jcr_other_first"] = ""

    # A*类：单独字段（勿写入「其他」，否则填充时会显示成「其他」）
    m_as2 = re.search(
        rf"A\s*[*＊]\s*类\s*[：:]\s*({_NUM})\s*篇"
        + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
        school_only,
        re.I,
    )
    m_as = re.search(
        rf"A\s*[*＊]\s*类\s*[：:]\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        school_only,
        re.I,
    )
    m_as0 = re.search(
        rf"A\s*[*＊]\s*类\s*[：:]\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        school_only,
        re.I,
    )
    if m_as2:
        s["school_a_star"] = _to_num(m_as2.group(1))
        s["school_a_star_first"] = _to_num(m_as2.group(2))
    elif m_as:
        s["school_a_star"] = _to_num(m_as.group(1))
        s["school_a_star_first"] = _to_num(m_as.group(2))
    elif m_as0:
        s["school_a_star"] = _to_num(m_as0.group(1))
        s["school_a_star_first"] = ""

    # A类/B类/C类：允许「类」与冒号之间有空格（如 Word 中「A类 ： 3 篇」）
    for pat, key in [
        (r"A\s*类\s*[：:]", "school_a"),
        (r"B\s*类\s*[：:]", "school_b"),
        (r"C\s*类\s*[：:]", "school_c"),
    ]:
        m2 = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇"
            + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
            school_only,
            re.I,
        )
        m1 = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            school_only,
            re.I,
        )
        m0 = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            school_only,
            re.I,
        )
        m_bare = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇",
            school_only,
            re.I,
        )
        if m2:
            s[key] = _to_num(m2.group(1))
            s[f"{key}_first"] = _to_num(m2.group(2))
        elif m1:
            s[key] = _to_num(m1.group(1))
            s[f"{key}_first"] = _to_num(m1.group(2))
        elif m0:
            s[key] = _to_num(m0.group(1))
            s[f"{key}_first"] = ""
        elif m_bare:
            s[key] = _to_num(m_bare.group(1))
            s[f"{key}_first"] = ""

    # 校级「其他」（仅 school_only，避免与 CCF/ABS 段落的「其他」混用）
    m_so2 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇"
        + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
        school_only,
        re.I,
    )
    m_so = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        school_only,
        re.I,
    )
    m_so0 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        school_only,
        re.I,
    )
    # 兜底：括号内为描述性文字（如「包含一篇北核」），只取总篇数
    m_so_bare = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇",
        school_only,
        re.I,
    )
    if m_so2:
        s["school_other"] = _to_num(m_so2.group(1))
        s["school_other_first"] = _to_num(m_so2.group(2))
    elif m_so:
        s["school_other"] = _to_num(m_so.group(1))
        s["school_other_first"] = _to_num(m_so.group(2))
    elif m_so0:
        s["school_other"] = _to_num(m_so0.group(1))
        s["school_other_first"] = ""
    elif m_so_bare:
        s["school_other"] = _to_num(m_so_bare.group(1))
        s["school_other_first"] = 0

    # 中国计算机学会（CCF）：仅在 ccf_only 段匹配 A/B/C/其他 类
    for pat, key in [
        (r"A\s*类\s*[：:]", "ccf_a"),
        (r"B\s*类\s*[：:]", "ccf_b"),
        (r"C\s*类\s*[：:]", "ccf_c"),
    ]:
        m2 = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇"
            + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
            ccf_only,
            re.I,
        )
        m1 = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            ccf_only,
            re.I,
        )
        m0 = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            ccf_only,
            re.I,
        )
        m_bare = re.search(
            rf"(?<!\*)(?<![A-Z]){pat}\s*({_NUM})\s*篇",
            ccf_only,
            re.I,
        )
        if m2:
            s[key] = _to_num(m2.group(1))
            s[f"{key}_first"] = _to_num(m2.group(2))
        elif m1:
            s[key] = _to_num(m1.group(1))
            s[f"{key}_first"] = _to_num(m1.group(2))
        elif m0:
            s[key] = _to_num(m0.group(1))
            s[f"{key}_first"] = ""
        elif m_bare:
            s[key] = _to_num(m_bare.group(1))
            s[f"{key}_first"] = ""
    m_co2 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇"
        + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
        ccf_only,
        re.I,
    )
    m_co = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        ccf_only,
        re.I,
    )
    m_co0 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        ccf_only,
        re.I,
    )
    m_co_bare = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇",
        ccf_only,
        re.I,
    )
    if m_co2:
        s["ccf_other"] = _to_num(m_co2.group(1))
        s["ccf_other_first"] = _to_num(m_co2.group(2))
    elif m_co:
        s["ccf_other"] = _to_num(m_co.group(1))
        s["ccf_other_first"] = _to_num(m_co.group(2))
    elif m_co0:
        s["ccf_other"] = _to_num(m_co0.group(1))
        s["ccf_other_first"] = ""
    elif m_co_bare:
        s["ccf_other"] = _to_num(m_co_bare.group(1))
        s["ccf_other_first"] = ""

    # ABS：ABS4 / ABS3 / ABS2 / ABS1 / 其他（兼容「ABS4」「ABS 4」）
    for n, key_base in [(4, "abs_4"), (3, "abs_3"), (2, "abs_2"), (1, "abs_1")]:
        m2 = re.search(
            rf"(?:ABS\s*{n}|ABS{n})\s*[：:]\s*({_NUM})\s*篇"
            + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
            abs_only,
            re.I,
        )
        m1 = re.search(
            rf"(?:ABS\s*{n}|ABS{n})\s*[：:]\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            abs_only,
            re.I,
        )
        m0 = re.search(
            rf"(?:ABS\s*{n}|ABS{n})\s*[：:]\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
            abs_only,
            re.I,
        )
        m_bare = re.search(
            rf"(?:ABS\s*{n}|ABS{n})\s*[：:]\s*({_NUM})\s*篇",
            abs_only,
            re.I,
        )
        if m2:
            s[key_base] = _to_num(m2.group(1))
            s[f"{key_base}_first"] = _to_num(m2.group(2))
        elif m1:
            s[key_base] = _to_num(m1.group(1))
            s[f"{key_base}_first"] = _to_num(m1.group(2))
        elif m0:
            s[key_base] = _to_num(m0.group(1))
            s[f"{key_base}_first"] = ""
        elif m_bare:
            s[key_base] = _to_num(m_bare.group(1))
            s[f"{key_base}_first"] = ""
    m_ao2 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇"
        + _PAREN_FIRST_COMM.replace("@@A@@", _NUM).replace("@@B@@", _NUM),
        abs_only,
        re.I,
    )
    m_ao = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*({_NUM})\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        abs_only,
        re.I,
    )
    m_ao0 = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇[（(]\s*篇(?:为一作|为第一作者)[^）)]*[）)]",
        abs_only,
        re.I,
    )
    m_ao_bare = re.search(
        rf"{_OTHER_LABEL}\s*[：:]\s*({_NUM})\s*篇",
        abs_only,
        re.I,
    )
    if m_ao2:
        s["abs_other"] = _to_num(m_ao2.group(1))
        s["abs_other_first"] = _to_num(m_ao2.group(2))
    elif m_ao:
        s["abs_other"] = _to_num(m_ao.group(1))
        s["abs_other_first"] = _to_num(m_ao.group(2))
    elif m_ao0:
        s["abs_other"] = _to_num(m_ao0.group(1))
        s["abs_other_first"] = ""
    elif m_ao_bare:
        s["abs_other"] = _to_num(m_ao_bare.group(1))
        s["abs_other_first"] = ""

    return s


_ISSN_TOKEN = re.compile(r"^\d{4}-\d{3}[\dXx]$", re.I)
_ISSN_IN_TEXT_RE = re.compile(r"\d{4}-\d{3}[\dXx]", re.I)
_MON_YEAR_TOKEN = re.compile(
    r"^(Jan(?:uary)?\.?|Feb(?:ruary)?\.?|Mar(?:ch)?\.?|Apr(?:il)?\.?|May\.?|Jun(?:e)?\.?|Jul(?:y)?\.?|Aug(?:ust)?\.?|Sep(?:t(?:ember)?)?\.?|Oct(?:ober)?\.?|Nov(?:ember)?\.?|Dec(?:ember)?\.?)\s*\d{4}$",
    re.I,
)


_MON_YEAR_HEAD = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
)


def find_valid_issn_in_text(text):
    """
    识别逗号/竖线/括号分隔的独立 ISSN；排除页码子串（16316-16324→6316-1632）及 p./pp. 之后的匹配。
    """
    s = text or ""
    if not s.strip():
        return ""
    p_page = re.search(r"(?i)\bpp?\.?\s*\d", s)
    p_page_start = p_page.start() if p_page else len(s) + 1
    for m in _ISSN_IN_TEXT_RE.finditer(s):
        issn = m.group()
        st, en = m.start(), m.end()
        if st >= p_page_start:
            continue
        if st > 0 and s[st - 1].isdigit():
            continue
        if en < len(s) and s[en].isdigit():
            continue
        before = s[max(0, st - 14) : st]
        if re.search(r"(?i)(?:pp?\.?\s*|Page\s*)$", before):
            continue
        ctx = s[max(0, st - 25) : en + 8]
        if re.search(rf"97[89][-\s\d.]*{re.escape(issn)}", ctx, re.I):
            continue
        if re.search(rf"[\d\)]\s*:\s*{re.escape(issn)}\b", s):
            continue
        tail = s[en : en + 80]
        if re.search(rf"[,，]\s*{re.escape(issn)}\s*[,，]", s):
            return issn
        if re.search(rf",\s*{re.escape(issn)}\s*[,，]", s):
            return issn
        if re.search(rf"\]\s*{re.escape(issn)}\s*[,，]", s):
            return issn
        if re.search(rf"(?:^|[|\n])\s*{re.escape(issn)}\s*[,，]", s):
            return issn
        if re.search(rf"\(\s*{re.escape(issn)}\s*\)", s):
            return issn
        if re.search(
            rf"ISSN[\s\u00a0\u3000]*号?[\s\u00a0\u3000]*[：:][\s\u00a0\u3000]*{re.escape(issn)}\b",
            s,
            re.I,
        ):
            return issn
        if re.search(
            rf"{re.escape(issn)}\s*[,，]\s*(?:{_MON_YEAR_HEAD}|20\d{{2}}|Volume\s*[:：])",
            s,
            re.I,
        ):
            return issn
        if re.search(rf"{re.escape(issn)}\s+Vol(?:ume)?\.?\s+\d", s, re.I):
            return issn
        if re.search(
            rf"{re.escape(issn)}\s*(?:\[\*\*[^\*]*\*\*\]\s*)?[,，\s]*(?:{_MON_YEAR_HEAD}\.?\s*\d{{4}}|Volume\s*[:：])",
            tail,
            re.I,
        ):
            return issn
        if re.search(
            rf"\]\s*{re.escape(issn)}\s*(?:\[\*\*[^\*]*\*\*\]\s*)?[,，\s]*(?:{_MON_YEAR_HEAD}|Volume\s*[:：])",
            s[max(0, st - 2) : en + 80],
            re.I,
        ):
            return issn
    return ""


def _parse_pub_info_comma_tail(pub, out):
    """
    英文化学类常见「刊名, ISSN, Jul. 2020, 49, 6002-6038」逗号分隔格式：
    补全 journal（去掉误并入的 ISSN 段）、volume、pages（及 Just Accepted 等尾段）。
    """
    work = (pub or "").replace("|", ",").replace("，", ",").replace("：", ":")
    work = re.sub(r"[（(]\s*共\s*\d+\s*页\s*[）)]", "", work, flags=re.IGNORECASE)
    work = work.strip().rstrip(".,;； ")
    if not work:
        return
    segs = [x.strip().rstrip(".,;；") for x in re.split(r"\s*,\s*", work) if x.strip()]
    if len(segs) < 4:
        return
    issn_idx = None
    for i, seg in enumerate(segs):
        if _ISSN_TOKEN.match(seg.replace(" ", "")):
            issn_idx = i
            break
    if issn_idx is None:
        return
    journal_join = "，".join(segs[:issn_idx]).strip(" ，,")
    date_idx = None
    for j in range(issn_idx + 1, len(segs)):
        sg = segs[j].strip()
        if _MON_YEAR_TOKEN.match(sg) or _MON_YEAR_TOKEN.match(sg.replace(" ", "")):
            date_idx = j
            break
    if date_idx is None:
        return
    i = date_idx + 1
    while i < len(segs) and _ISSN_TOKEN.match(segs[i].strip().replace(" ", "")):
        i += 1
    vol_val = ""
    if i < len(segs):
        seg_i = segs[i].strip()
        if re.fullmatch(r"\d{1,4}", seg_i):
            vol_val = seg_i
            i += 1
        else:
            # 「Volume NNN」格式（无冒号）：抽取纯数字并跳过该段
            m_vw = re.match(r"(?i)^Volume\s+(\d+)$", seg_i)
            if m_vw:
                vol_val = m_vw.group(1)
                i += 1
    pages_val = ""
    if i < len(segs):
        pages_val = ", ".join(segs[i:]).strip()
    if pages_val and re.search(r"(?i)\bvolume\b", pages_val):
        # 含 Volume 关键字：实为引文整段，非纯页码，避免覆盖主解析得到的页码
        pages_val = ""
    if journal_join:
        old_j = (out.get("journal") or "").strip()
        issn_txt = (out.get("issn") or "").strip()
        if (
            not old_j
            or issn_txt in old_j
            or ("," in old_j and len(old_j) > len(journal_join) + 5)
        ):
            out["journal"] = journal_join
    if vol_val:
        old_v = str(out.get("volume") or "").strip()
        if not old_v or not re.search(r"\(\d+\)", old_v):
            out["volume"] = vol_val
    if pages_val:
        old_p = str(out.get("pages") or "").strip()
        if not old_p or len(pages_val) >= len(old_p):
            out["pages"] = pages_val


def _parse_pub_info_apa_paren(pub, out):
    """
    APA 单行：刊名 (ISSN). 年, 卷(期), 页码（如错误073心理学简历）。
    含 DOI 的尾段不拆，留给回填原样展示。
    """
    work = re.sub(r"\s*\|\s*", ", ", (pub or "").strip())
    work = re.sub(r"\s+", " ", work)
    m = re.match(
        r"^(.+?)\s*\(\s*(\d{4}-\d{3}[\dXx])\s*\)\s*\.?\s*((?:19|20)\d{2})\s*,\s*(\d+\(\d+\))\s*,\s*(.+)$",
        work,
        re.I,
    )
    if not m:
        return False
    tail = m.group(5).strip()
    if re.search(r"\bDOI\s*:", tail, re.I):
        return False
    pages = re.sub(r"\s+", "", tail.replace("–", "-").replace("—", "-"))
    out["journal"] = m.group(1).strip()
    out["issn"] = m.group(2)
    out["date"] = m.group(3)
    out["volume"] = m.group(4)
    out["pages"] = pages
    return True


def parse_pub_info(pub):
    """
    期刊列经 consec_unique 后，单元格内换行会变成 | 分段。
    若第二段为单独一行的「(人文社会科学版)」类副题名，须写入 journal_subtitle，
    否则回填时只会输出「浙江大学学报」而丢了副刊名。
    """
    apa_out = {
        "journal": "",
        "issn": "",
        "date": "",
        "volume": "",
        "issue": "",
        "pages": "",
        "page_count": "",
    }
    if _parse_pub_info_apa_paren(pub, apa_out):
        page_count = re.search(r"共\s*(\d+)\s*页", pub or "")
        apa_out["page_count"] = page_count.group(1) if page_count else ""
        return apa_out

    parts = [p.strip(" ，,") for p in pub.split("|") if p.strip(" ，,")]
    journal = re.sub(r"[，,。\s]+$", "", parts[0]) if parts else ""
    journal_subtitle = ""
    if len(parts) >= 2:
        p1 = parts[1].strip()
        # 整段仅为括号副题名（全角/半角括号），不与 ISSN 等混淆
        # 「（共14页）」等为页数说明，须写入 page_count 而非 journal_subtitle，否则回填会重复一行且误加粗
        if re.match(r"^[（(][^（）()]*[）)]\s*$", p1) and not re.search(
            r"共\s*\d+\s*页", p1
        ):
            journal_subtitle = p1

    issn_val = find_valid_issn_in_text(pub)
    page_count = re.search(r"共\s*(\d+)\s*页", pub)

    date = ""
    # 英文月份：允许全称（January）或缩写（Jan）或缩写加句点（Jan.）
    m = re.search(
        r"\b(Jan(?:uary)?\.?|Feb(?:ruary)?\.?|Mar(?:ch)?\.?|Apr(?:il)?\.?|May\.?|Jun(?:e)?\.?|Jul(?:y)?\.?|Aug(?:ust)?\.?|Sep(?:t(?:ember)?)?\.?|Oct(?:ober)?\.?|Nov(?:ember)?\.?|Dec(?:ember)?\.?)\s*\d{4}\b",
        pub,
        re.IGNORECASE,
    )
    if m:
        date = re.sub(r"\s+", " ", m.group(0)).strip()
    else:
        m = re.search(r"\d{4}年\d+月", pub)
        if m:
            date = m.group(0)

    volume = ""
    m = re.search(r"Volume[：:]\s*([^\s,，|]+)", pub, re.IGNORECASE)
    if m:
        volume = m.group(1).strip()
    else:
        m = re.search(r"\bVolume\s+(\d+)\b", pub, re.IGNORECASE)
        if m:
            volume = m.group(1).strip()
        else:
            m = re.search(r"(\d+\(\d+\))", pub)
            if m:
                volume = m.group(1)

    issue = ""
    m = re.search(r"Issue[：:]\s*([^\s,，|]+)", pub, re.IGNORECASE)
    if m:
        issue = m.group(1).strip().rstrip(",，")
    else:
        m = re.search(r"\bIssue\s+(\d+)\b", pub, re.IGNORECASE)
        if m:
            issue = m.group(1).strip()

    pages = ""
    m = re.search(r"[Pp]\.?\s*(\d+\s*[-—]\s*\d+)", pub)
    if m:
        pages = re.sub(r"\s+", "", m.group(1))
    else:
        m = re.search(r"[Pp]\.?\s*(\d+)", pub)
        if m:
            pages = m.group(1).strip()
    if not pages:
        m = re.search(r"(?i)\bPage\s+(\d+(?:\s*[-—]\s*\d+)?)", pub)
        if m:
            pages = re.sub(r"\s+", "", m.group(1))
    if not pages:
        m = re.search(r"[：:]\s*(\d+\s*[-—]\s*\d+)", pub)
        if m:
            pages = re.sub(r"\s+", "", m.group(1))

    out = {
        "journal": journal,
        "issn": issn_val,
        "date": date,
        "volume": volume,
        "issue": issue,
        "pages": pages,
        "page_count": page_count.group(1) if page_count else "",
    }
    if journal_subtitle:
        out["journal_subtitle"] = journal_subtitle
    _parse_pub_info_comma_tail(pub, out)
    return out


def parse_level_info(level):
    """
    解析论文「期刊分级」单元格（竖线或单元格内换行拆出的多段）。
    标准四段：JCR / 校级 / 中科院 / IF；其余无法归类的段落写入 level_note（如「折为1/2」），供 fill_template 原样输出。
    """
    parts = [p.strip() for p in re.split(r"[|\n]+", level or "") if p.strip()]
    out = {"jcr": "", "school_level": "", "cas_zone": "", "if": "", "level_note": ""}
    extras = []
    for p in parts:
        if re.search(r"JCR\s*Q[1-4]", p, re.IGNORECASE):
            out["jcr"] = p
        elif (
            re.search(r"校级[A-Da-d类其他]", p)
            or re.search(r"校级A[＊*]类", p)
            or p in ("A类", "B类", "C类", "D类")
            or p.strip() in ("校级其他",)
        ):
            out["school_level"] = p
        elif "中科院" in p:
            out["cas_zone"] = p
        elif re.search(r"\bIF\b", p, re.IGNORECASE):
            out["if"] = p
        else:
            extras.append(p)
    if extras:
        out["level_note"] = "\n".join(extras)
    return out


def parse_projects_summary(text):
    out = {k: 0 for k in SCHEMA_TEMPLATE["projects_summary"]}
    host_text = text
    part_text = ""
    # 在「参与A类/B类项目」之前切分，避免误匹配文本中其他位置的「参与」
    m_part = re.search(r"参与(?:A类|B类|\d+)", text)
    if m_part:
        host_text = text[:m_part.start()]
        part_text = text[m_part.start():]

    def n(pat, t):
        m = re.search(pat, t)
        return int(m.group(1)) if m else 0

    out["host_a_major"]      = n(r"主持A类重大项目\s*(\d+)\s*项", host_text)
    out["host_a_key"]        = n(r"A类重点项目\s*(\d+)\s*项",    host_text)
    out["host_a"]            = n(r"(?<!重点)A类项目\s*(\d+)\s*项", host_text)
    out["host_b"]            = n(r"B类项目\s*(\d+)\s*项",         host_text)
    out["participate_a_major"] = n(r"参与A类重大项目\s*(\d+)\s*项", part_text)
    out["participate_a_key"]   = n(r"A类重点项目\s*(\d+)\s*项",    part_text)
    out["participate_a"]       = n(r"(?<!重点)A类项目\s*(\d+)\s*项", part_text)
    out["participate_b"]       = n(r"B类项目\s*(\d+)\s*项",         part_text)
    return out


def parse_recommendation(text):
    """
    解析「推荐人意见」单元格整段文本。
    典型：1.张三（本人博士生导师），某某大学教授，某某领域专家：正文
    兼容：）与逗号之间有空格；序号点后空格；正文前为半角冒号。
    """
    text = text.strip()
    # Word 偶发零宽字符，会导致「姓名」与「（角色）」断开使正则漏匹配
    for zw in ("\u200b", "\u200c", "\u200d", "\ufeff"):
        text = text.replace(zw, "")
    m = re.match(
        r"\d+\.\s*(.+?)\s*[（(]\s*(.+?)\s*[）)]\s*[，,]\s*(.+?)[:：]\s*\|?(.*)",
        text,
        re.S,
    )
    if m:
        return {
            "name": m.group(1).strip(),
            "role": m.group(2).strip(),
            "title": m.group(3).strip().rstrip("，,"),
            "content": re.sub(r"\n{2,}", "\n", m.group(4).replace("|", "\n")).strip(),
        }
    m2 = re.match(r"\d+\.\s*(.+?)\s*[，,]\s*(.+?)[:：]\s*\|?(.*)", text, re.S)
    if m2:
        return {
            "name": m2.group(1).strip(),
            "role": "",
            "title": m2.group(2).strip().rstrip("，,"),
            "content": re.sub(r"\n{2,}", "\n", m2.group(3).replace("|", "\n")).strip(),
        }
    # 无序号：「姓名（角色），机构，职称：正文」
    m3 = re.match(
        r"(.+?)\s*[（(]\s*(.+?)\s*[）)]\s*[，,]\s*(.+?)[:：]\s*\|?(.*)",
        text,
        re.S,
    )
    if m3:
        return {
            "name": m3.group(1).strip(),
            "role": m3.group(2).strip(),
            "title": m3.group(3).strip().rstrip("，,"),
            "content": re.sub(r"\n{2,}", "\n", m3.group(4).replace("|", "\n")).strip(),
        }
    # 无序号无角色：「姓名，机构职称：正文」
    m4 = re.match(r"(.+?)\s*[，,]\s*(.+?)[:：]\s*\|?(.*)", text, re.S)
    if m4:
        return {
            "name": m4.group(1).strip(),
            "role": "",
            "title": m4.group(2).strip().rstrip("，,"),
            "content": re.sub(r"\n{2,}", "\n", m4.group(3).replace("|", "\n")).strip(),
        }
    return {"name": "", "role": "", "title": "", "content": re.sub(r"\n{2,}", "\n", text.replace("|", "\n")).strip()}


def _trim_college_support_tail(tail: str) -> str:
    """去掉误入的「相关待遇」等后续节；截掉再次出现的「1. 岗位任务」重复块。"""
    if not tail:
        return ""
    m_cut = re.search(
        r"\r?\n\s*(?:相关待遇|年薪[：:]|本人签名|填表日期|\d+[.．]其他)",
        tail,
    )
    if m_cut:
        tail = tail[: m_cut.start()]
    m_dup = re.search(r"\r?\n\s*1[.．]\s*岗位任务", tail)
    if m_dup:
        tail = tail[: m_dup.start()]
    return tail.strip()


def _cell_text_with_auto_num(cell, doc):
    """
    读取单元格内所有段落文本，并补全 Word 自动编号（numPr）前缀。
    python-docx 的 cell.text / paragraph.text 均不含自动列表编号，需从 XML 中读取。
    支持格式：decimal（1.2.3.）和 「（N）」（lvlText 含 %N 的中文括号格式）。
    """
    from lxml import etree

    nsmap = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }

    # 构建 numId -> absNumId -> ilvl -> lvlText/numFmt/start 的映射
    num_fmt_map = {}  # (numId, ilvl) -> (numFmt, lvlText, start_val)
    counters = {}     # (numId, ilvl) -> current_count

    try:
        num_part = doc.part.numbering_part
        num_xml = num_part._element
        abs_nums = {}
        for abs_num in num_xml.findall(".//w:abstractNum", nsmap):
            abs_id = abs_num.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId")
            lvl_map = {}
            for lvl in abs_num.findall("w:lvl", nsmap):
                ilvl = lvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl", "0")
                fmt_el = lvl.find("w:numFmt", nsmap)
                txt_el = lvl.find("w:lvlText", nsmap)
                start_el = lvl.find("w:start", nsmap)
                fmt = fmt_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "decimal") if fmt_el is not None else "decimal"
                txt = txt_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "%1") if txt_el is not None else "%1"
                start = int(start_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "1")) if start_el is not None else 1
                lvl_map[ilvl] = (fmt, txt, start)
            abs_nums[abs_id] = lvl_map

        for num_el in num_xml.findall("w:num", nsmap):
            num_id = num_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
            abs_ref = num_el.find("w:abstractNumId", nsmap)
            if abs_ref is not None:
                abs_id = abs_ref.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if abs_id in abs_nums:
                    for ilvl, info in abs_nums[abs_id].items():
                        num_fmt_map[(num_id, ilvl)] = info
                        counters[(num_id, ilvl)] = info[2] - 1  # start - 1, incremented before use
    except Exception:
        pass

    lines = []
    for para in cell.paragraphs:
        ptext = para.text
        # 检查是否有 numPr
        p_el = para._element
        num_pr = p_el.find(".//w:numPr", nsmap)
        prefix = ""
        if num_pr is not None:
            ilvl_el = num_pr.find("w:ilvl", nsmap)
            num_id_el = num_pr.find("w:numId", nsmap)
            if ilvl_el is not None and num_id_el is not None:
                ilvl = ilvl_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")
                num_id = num_id_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0")
                key = (num_id, ilvl)
                if key in num_fmt_map:
                    counters[key] = counters.get(key, num_fmt_map[key][2] - 1) + 1
                    n = counters[key]
                    fmt, lvl_txt, _ = num_fmt_map[key]
                    # 用序号替换 %N 占位符
                    formatted = lvl_txt.replace(f"%{int(ilvl)+1}", str(n))
                    prefix = formatted + " "
        lines.append(prefix + ptext)
    return "\n".join(lines)


def _extract_job_tasks_college_support_block(block: str) -> Tuple[str, str]:
    """
    从「岗位任务及支撑保障」大段中拆出 job_tasks / college_support。

    支持多种版式：
    - 岗位任务节标题：「1. 岗位任务：」「1．岗位任务」「岗位任务：」（有无序号均可）
    - 学院支撑节标题：「2. 学院帮助…」「3. 学院支持…」（任意数字前缀，关键词含「学院」）
    - 节标题后内容可在冒号同行，也可在下一行
    - 条目编号可为「1.」「（1）」等任意格式
    """
    block = (block or "").replace("|", "\n").strip()
    job, sup = "", ""
    if "岗位任务" not in block:
        return job, sup

    # ── 定位「岗位任务」节标题（可带也可不带序号前缀）────────────────────
    # 冒号或序号至少须有其一，避免误匹配节标题行「岗位任务及支撑保障」
    m_job_hdr = re.search(
        r"(?:^|\n)[ \t]*(?:\d+[.．、][ \t]*岗位任务[：:]?|岗位任务[：:])[ \t]*",
        block,
        re.M,
    )

    _COLLEGE_PAT = r"(?:^|\n)[ \t]*(?:\d+[.．、]?[ \t]*)?学院(?:帮助|支持|协助|提供)[^\n]*"

    if not m_job_hdr:
        # 无「1.岗位任务：」子标题，但块内可能直接以条目开头；
        # 若能找到学院支撑节，则取其前内容（去掉行标题行）为 job_tasks
        m_college_fb = re.search(_COLLEGE_PAT, block, re.M)
        if m_college_fb:
            first_nl = block.find("\n")
            content_start = first_nl + 1 if first_nl >= 0 else 0
            job = block[content_start : m_college_fb.start()].strip()
            college_hdr_line = m_college_fb.group(0).lstrip("\n").rstrip()
            college_body_abs = m_college_fb.end()
            if re.search(r"[：:]\s*$", college_hdr_line):
                sup = block[college_body_abs:].strip()
            else:
                m_inline = re.search(r"[：:]\s*([\s\S]+)", college_hdr_line)
                if m_inline:
                    inline = m_inline.group(1).strip()
                    rest = block[college_body_abs:].strip()
                    sup = (inline + "\n" + rest).strip() if rest else inline
                else:
                    sup = block[college_body_abs:].strip()
            sup = _trim_college_support_tail(sup)
            return job.strip(), (sup or "").strip()
        return job, sup

    job_body_start = m_job_hdr.end()

    # ── 定位「学院」支撑节（可为 2./3. 等任意序号；序号可选；分隔符可选）──
    m_college_hdr = re.search(
        _COLLEGE_PAT,
        block[job_body_start:],
        re.M,
    )

    if m_college_hdr:
        # job_tasks = 岗位任务节标题之后、学院节标题之前
        job = block[job_body_start : job_body_start + m_college_hdr.start()].strip()

        # 学院节标题行（去掉前导换行）
        college_hdr_line = m_college_hdr.group(0).lstrip("\n").rstrip()
        college_body_abs = job_body_start + m_college_hdr.end()

        # 判断内容是否紧跟在标题行末冒号后，还是在下一行
        if re.search(r"[：:]\s*$", college_hdr_line):
            # 标题以冒号结尾，内容在下一行
            sup = block[college_body_abs:].strip()
        else:
            # 冒号后有行内内容（或无冒号）
            m_inline = re.search(r"[：:]\s*([\s\S]+)", college_hdr_line)
            if m_inline:
                inline = m_inline.group(1).strip()
                rest = block[college_body_abs:].strip()
                sup = (inline + "\n" + rest).strip() if rest else inline
            else:
                sup = block[college_body_abs:].strip()
    else:
        # 无学院节，全部内容归 job_tasks
        job = block[job_body_start:].strip()

    sup = _trim_college_support_tail(sup)
    return job.strip(), (sup or "").strip()


def _strip_vertical_table_noise(s):
    """去掉单元格首尾竖线/空白，便于「||2019.10-2023.10|」等格内匹配日期（错误009）。"""
    return re.sub(r"^[|｜\s　]+|[|｜\s　]+$", "", (s or "").strip())


def extract_one(docx_path):
    data = copy.deepcopy(SCHEMA_TEMPLATE)
    doc = Document(str(docx_path))
    if not doc.tables:
        raise ValueError(f"Word 文件无表格: {docx_path}")
    t = doc.tables[0]
    row_count = len(t.rows)
    if row_count < 6:
        raise ValueError(f"主表行数不足（仅 {row_count} 行）: {docx_path}")

    # 基本信息行：用逐格文本（含空），避免横向合并/空格导致 consec_unique 丢列错位
    r0 = row_cells_all(t.rows[0])
    r1 = row_cells_all(t.rows[1])
    r2 = consec_unique(t.rows[2])
    r3 = consec_unique(t.rows[3])
    r4 = consec_unique(t.rows[4])

    # ── 基本信息 ──────────────────────────────────────────────
    bi = data["basic_info"]
    bi["name"] = (
        find_after_basic_cell("姓  名", r0, _BASIC_R0_STOPS["姓  名"])
        or find_after_basic_cell("姓名", r0, _BASIC_R0_STOPS["姓名"])
        or find_after("姓  名", consec_unique(t.rows[0]))
        or find_after("姓名", consec_unique(t.rows[0]))
    )
    bi["gender"] = find_after_basic_cell("性别", r0, _BASIC_R0_STOPS["性别"]) or find_after(
        "性别", consec_unique(t.rows[0])
    )
    bi["hometown"] = find_after_basic_cell("籍贯", r0, _BASIC_R0_STOPS["籍贯"]) or find_after(
        "籍贯", consec_unique(t.rows[0])
    )
    bi["marital_status"] = find_after_basic_cell(
        "婚姻状况", r0, _BASIC_R0_STOPS["婚姻状况"]
    ) or find_after("婚姻状况", consec_unique(t.rows[0]))
    bi["birth"] = (
        find_after_basic_cell("出生年月", r0, _BASIC_R0_STOPS["出生年月"])
        or find_after("出生年月", consec_unique(t.rows[0]))
    ).replace("|", "")
    bi["current_title"]  = extract_current_title_from_row(t.rows[1])
    bi["discipline"]     = find_after_basic_cell(
        "一级学科", r1, _BASIC_R1_STOPS["一级学科"]
    ) or find_after("一级学科", consec_unique(t.rows[1]), _BASIC_R1_STOPS["一级学科"])
    bi["religion"]       = find_after_basic_cell(
        "宗教信仰", r1, _BASIC_R1_STOPS["宗教信仰"]
    ) or find_after("宗教信仰", consec_unique(t.rows[1]), _BASIC_R1_STOPS["宗教信仰"])
    bi["target_college"] = (
        find_after_basic_cell("拟聘学院", r1, _BASIC_R1_STOPS["拟聘学院"])
        or find_after_basic_cell("拟聘单位", r1, _BASIC_R1_STOPS["拟聘单位"])
        or find_after("拟聘学院", consec_unique(t.rows[1]), _BASIC_R1_STOPS["拟聘学院"])
        or find_after("拟聘单位", consec_unique(t.rows[1]), _BASIC_R1_STOPS["拟聘单位"])
    )

    # ── 学习/工作/海外经历 ────────────────────────────────────
    ewot = next((x for x in r2 if "学习经历" in x), "")
    sec = parse_edu_work_overseas(ewot)
    data["education"] = sec["education"]
    data["work"]      = sec["work"]
    data["overseas"]  = sec["overseas"]
    data["experience_notes"] = sec.get("experience_notes") or []

    # ── 研究信息 ──────────────────────────────────────────────
    data["research_area"]     = normalize_piped_table_text(find_after("研究领域及专长", r3))
    data["research_team"]     = normalize_piped_table_text(find_after("拟加入科研团队", r3), "")
    data["team_achievement"]  = normalize_piped_table_text(find_after("团队人才培育成效", r4))

    # ── 论文统计 ──────────────────────────────────────────────
    # 统计行通常在 row5，但部分模板在 row4；动态查找含「总体情况」的行，
    # 同时拼接整行所有格，避免因多列拆分丢失「校级其他」等字段。
    summary_row_idx = next(
        (i for i in range(3, min(10, row_count))
         if "总体情况" in "|".join(consec_unique(t.rows[i]))),
        5,
    )
    r5 = consec_unique(t.rows[summary_row_idx])
    summary_text = "\n".join(r5) if r5 else ""
    data["papers_summary"] = parse_papers_summary(summary_text)

    # 预构建每行文本（用于行定位）
    row_texts = ["|".join(consec_unique(t.rows[i])) for i in range(row_count)]

    # ── 论文明细 ──────────────────────────────────────────────
    paper_hdr = next((i for i, txt in enumerate(row_texts)
                      if "作者信息" in txt and "论文题目" in txt), -1)
    proj_sum_idx = next((i for i, txt in enumerate(row_texts)
                         if "主持A类重大项目" in txt), -1)
    pend = (
        proj_sum_idx
        if paper_hdr >= 0 and proj_sum_idx > paper_hdr
        else (row_count if paper_hdr >= 0 else -1)
    )
    papers = []
    last_paper_seq = 0  # 上一行论文序号；用于「序号列纵向合并」时续行无数字格
    if paper_hdr >= 0:
        for ridx in range(paper_hdr + 1, pend):
            cells_obj = consec_unique_with_cells(t.rows[ridx])
            cells = [c[0] for c in cells_obj]
            seq = next((x for x in cells if re.fullmatch(r"\d+", x.strip())), "")
            authors, title, pub, level = "", "", "", ""
            if seq and int(seq) > 0 and int(seq) < 500:
                try:
                    si = next(j for j, c in enumerate(cells) if c.strip() == seq)
                    last_paper_seq = int(seq)
                    # 作者信息：读取 run 格式，将加粗+下划线编码为标记
                    authors = read_cell_marked(cells_obj[si + 1][1]) if si + 1 < len(cells_obj) else ""
                    title   = cells[si + 2] if si + 2 < len(cells) else ""
                    pub     = cells[si + 3] if si + 3 < len(cells) else ""
                    level   = cells[si + 4] if si + 4 < len(cells) else ""
                    pub_cell = cells_obj[si + 3][1] if si + 3 < len(cells_obj) else None
                except (StopIteration, IndexError):
                    continue
            elif (
                last_paper_seq > 0
                and len(cells) == 6
                and ("代表性" in (cells[1] or "") or "论文" in (cells[1] or ""))
            ):
                # Word 将序号列纵向合并后，本行无独立序号格；作者列从索引 2 开始（少一格序号）
                si_eff = 2
                if si_eff + 3 >= len(cells) or si_eff >= len(cells_obj):
                    continue
                last_paper_seq += 1
                seq = str(last_paper_seq)
                authors = read_cell_marked(cells_obj[si_eff][1])
                title   = cells[si_eff + 1]
                pub     = cells[si_eff + 2]
                level   = cells[si_eff + 3]
                pub_cell = cells_obj[si_eff + 2][1] if si_eff + 2 < len(cells_obj) else None
            else:
                continue
            pi = parse_pub_info(pub)
            li = parse_level_info(level)
            pub_raw = (
                read_cell_marked_paragraphs(pub_cell).strip()
                if pub_cell is not None
                else pub.replace("|", "\n").strip()
            )
            row_paper = {
                "seq": int(seq), "authors": authors, "title": title,
                "journal": pi["journal"], "issn": pi["issn"],
                "date": pi["date"], "volume": pi["volume"],
                "issue": pi.get("issue", ""),
                "pages": pi["pages"], "page_count": pi["page_count"],
                "pub_raw": pub_raw,
                "jcr": li["jcr"], "school_level": li["school_level"],
                "cas_zone": li["cas_zone"], "if": li["if"],
            }
            if pi.get("journal_subtitle"):
                row_paper["journal_subtitle"] = pi["journal_subtitle"]
            if li.get("level_note"):
                row_paper["level_note"] = li["level_note"]
            papers.append(row_paper)
    data["papers"] = sorted(papers, key=lambda x: x["seq"])

    # ── 项目总体情况 ──────────────────────────────────────────
    if proj_sum_idx >= 0:
        r_ps = consec_unique(t.rows[proj_sum_idx])
        ps_text = next((x for x in r_ps if "总体情况" in x and "主持" in x), "")
        data["projects_summary"] = parse_projects_summary(ps_text)

    # ── 项目明细 ──────────────────────────────────────────────
    proj_hdr = next((i for i, txt in enumerate(row_texts)
                     if "起止时间" in txt and "项目名称" in txt), -1)
    # 「其他成果」行：须与科研项目表中含「其他」的列区分，见 _pick_other_achievements_row_index
    other_idx = _pick_other_achievements_row_index(
        row_texts,
        t,
        paper_hdr=paper_hdr if paper_hdr >= 0 else -1,
        paper_body_end=pend if paper_hdr >= 0 else -1,
        proj_hdr=proj_hdr,
    )
    proj_end = other_idx if other_idx > proj_hdr >= 0 else row_count
    # 起止时间：支持多种格式
    #   YYYY.MM-YYYY.MM、YYYY.MM.DD-YYYY.MM.DD（句点分隔，连字符范围）
    #   YYYY-MM 至 YYYY-MM、YYYY-MM至YYYY-MM（连字符分隔，「至」范围）
    #   任何格式均支持以「至今」结尾
    _proj_date_re = re.compile(
        r"\d{4}[.\-]\d{2}(?:[.\-]\d{2})?\s*(?:[-–—]|至)\s*(?:\d{4}[.\-]\d{2}(?:[.\-]\d{2})?|至今)?"
    )
    projects = []
    if proj_hdr >= 0:
        for ridx in range(proj_hdr + 1, proj_end):
            raw = row_cells_all(t.rows[ridx])
            cells = row_cells_collapse_merged_duplicates_non_empty(raw)
            if not any(str(x).strip() for x in cells):
                continue
            jd = next(
                (
                    j
                    for j, c in enumerate(cells)
                    if _proj_date_re.match(_strip_vertical_table_noise(c))
                ),
                None,
            )
            if jd is None:
                continue
            date = _strip_vertical_table_noise(cells[jd])
            seq_str = ""
            if jd > 0:
                prev = _strip_vertical_table_noise(cells[jd - 1])
                if re.fullmatch(r"\d+", prev) and int(prev) <= 500:
                    seq_str = prev
            if not seq_str:
                seq_str = str(len(projects) + 1)
            tail = cells[jd + 1 : jd + 16]
            tail = _project_tail_compress_empty_runs(tail)
            while len(tail) < 6:
                tail.append("")
            name, institution, category, level, fund, ranking = (
                tail[0],
                tail[1],
                tail[2],
                tail[3],
                tail[4],
                tail[5],
            )
            projects.append({
                "seq": int(seq_str),
                "date": date,
                "name": name,
                "institution": institution,
                "category": category,
                "level": level,
                "fund": fund,
                "ranking": ranking,
            })
    data["projects"] = sorted(projects, key=lambda x: x["seq"])

    # ── 其他成果 ──────────────────────────────────────────────
    if other_idx >= 0:
        data["other"] = _extract_other_achievements_from_row(t.rows[other_idx])

    # ── 推荐人意见 ────────────────────────────────────────────
    # 推荐人意见永远在「其他」之后，从 other_idx+1 开始扫描，避免「其他」区的
    # "1.代表性成果…" 等内容被误匹配
    rec_scan_start = other_idx + 1 if other_idx >= 0 else 0
    recs = []
    for ridx in range(rec_scan_start, row_count):
        txt = row_texts[ridx]
        # 「推荐人意见」单元格可能逐字分段存储（每段一字做竖排效果），
        # consec_unique 将换行转为「|」后变为「推|荐|人|意|见」，需同时兼容两种格式
        if "推荐人意见" not in txt and "推|荐|人|意|见" not in txt:
            continue
        cells = consec_unique(t.rows[ridx])
        # \d+\.\D 确保点号后不是数字，避免把日期（2022.01-...）误匹配为推荐人条目
        rec_text = next((x for x in cells if re.match(r"\d+\.\D", x.strip())), "")
        if not rec_text:
            # 无序号格式：取非标题标签的内容格（如「陈菁（导师），...：正文」）
            _REC_LABELS = {"推荐人意见", "推|荐|人|意|见"}
            rec_text = next(
                (x for x in cells if x.strip() and x.strip() not in _REC_LABELS
                 and "推荐人意见" not in x and re.search(r"[，,]", x)),
                "",
            )
        if rec_text:
            recs.append(parse_recommendation(rec_text))
    data["recommendations"] = recs

    # ── 思想政治 ──────────────────────────────────────────────
    pol_idx = next((i for i, txt in enumerate(row_texts)
                    if "思想政治与品德综合评价" in txt), -1)
    if pol_idx >= 0:
        from docx.oxml.ns import qn as _qn
        def _tc_full_text(tc_el):
            """递归提取单元格文字，含嵌套表格。"""
            parts = []
            for child in tc_el.iter():
                if child.tag == _qn('w:t') and child.text:
                    parts.append(child.text)
                elif child.tag == _qn('w:br'):
                    parts.append('\n')
            return ''.join(parts)
        row_el = t.rows[pol_idx]._tr
        tcs_el = row_el.findall('.//' + _qn('w:tc'))
        r_pol_full = []
        seen = set()
        for tc_el in tcs_el:
            txt = _tc_full_text(tc_el).strip()
            if txt and txt not in seen:
                seen.add(txt)
                r_pol_full.append(txt)
        label = "思想政治与品德综合评价"
        chunks = [
            x for x in r_pol_full
            if x and x.strip("|").strip() not in _SECTION_LABELS
        ]
        # 勿用「标题 substring not in x」：合并格常把标题+正文写在同一单元格，且文末会再出现标题，整段会被误丢。
        pol = ""
        if chunks:
            pol = max(chunks, key=len)
            if pol.strip() == label.strip():
                pol = ""
            else:
                if pol.startswith(label):
                    pol = pol[len(label) :].lstrip("：:|\n ")
                for suf in ("|" + label, "\n" + label):
                    if pol.endswith(suf):
                        pol = pol[: -len(suf)].rstrip("| \n")
                if pol.endswith(label):
                    pol = pol[: -len(label)].rstrip("| \n")
        data["political"] = pol.replace("|", "\n").strip()

    # ── 岗位任务 & 学院支持 ───────────────────────────────────
    # 行标题常见「岗位任务及支撑保障表」；部分模板无「支撑保障」字样但含「1. 岗位任务」正文。
    task_idx = next(
        (
            i
            for i, txt in enumerate(row_texts)
            if "岗位任务" in txt
            and (
                "支撑保障" in txt
                or "学院帮助" in txt
                or re.search(r"1[.．、]\s*岗位任务", txt)
            )
        ),
        -1,
    )
    if task_idx < 0:
        for i, txt in enumerate(row_texts):
            if "岗位任务" in txt and re.search(r"1[.．、]\s*岗位任务", txt):
                task_idx = i
                break
    if task_idx >= 0:
        # 读取自动编号段落：cell.text 不含 Word list 编号，需逐段读取并补全 numPr 前缀
        task_row = t.rows[task_idx]
        seen_cell_ids = set()
        cell_texts = []
        for cell in task_row.cells:
            cid = id(cell._element)
            if cid in seen_cell_ids:
                continue
            seen_cell_ids.add(cid)
            cell_texts.append(_cell_text_with_auto_num(cell, doc))
        block = "\n".join(ct.strip() for ct in cell_texts if ct.strip())
        jt, cs = _extract_job_tasks_college_support_block(block)
        data["job_tasks"] = jt
        data["college_support"] = cs

    # ── 相关待遇 ──────────────────────────────────────────────
    treat_idx = next((i for i, txt in enumerate(row_texts)
                      if "相关待遇" in txt), -1)
    if treat_idx >= 0:
        r_sal = consec_unique(t.rows[treat_idx])
        for cell_s in r_sal:
            if not data.get("salary"):
                # 提取「税前」后的完整描述，直到分隔符或科研/安家关键词
                m = re.search(r"税前\s*(.+?)(?=[|；。\n]|科研|安家|$)", cell_s)
                if m:
                    data["salary"] = m.group(1).strip().rstrip("；;，,。")
            if not data.get("research_fund"):
                m = re.search(r"科研启动(?:经费|费)[：:]\s*(.+?)(?=[；|]|安家费|$)", cell_s)
                if m:
                    data["research_fund"] = m.group(1).strip()
            if not data.get("relocation_fee"):
                m = re.search(r"安家费[：:]\s*(\d+(?:\.\d+)?)", cell_s)
                if m:
                    data["relocation_fee"] = m.group(1).strip()

    # ── 多表回退：部分简历将推荐人/岗位任务放在 table[1] 或更后的表格 ──────
    need_recs = not data.get("recommendations")
    need_tasks = not data.get("job_tasks")
    if need_recs or need_tasks:
        _COLLEGE_PAT_FB = r"(?:^|\n)[ \t]*(?:\d+[.．、]?[ \t]*)?学院(?:帮助|支持|协助|提供)[^\n]*"
        _REC_LABELS_FB = {"推荐人意见", "推|荐|人|意|见"}
        for extra_t in doc.tables[1:]:
            extra_rows = len(extra_t.rows)
            extra_row_texts = []
            for ridx in range(extra_rows):
                cells = consec_unique(extra_t.rows[ridx])
                extra_row_texts.append("|".join(str(x) for x in cells))

            if need_recs:
                recs = []
                # 格式一：行内有「推荐人意见」标签
                for ridx in range(extra_rows):
                    txt = extra_row_texts[ridx]
                    if "推荐人意见" not in txt and "推|荐|人|意|见" not in txt:
                        continue
                    cells = consec_unique(extra_t.rows[ridx])
                    rec_text = next((x for x in cells if re.match(r"\d+\.\D", x.strip())), "")
                    if not rec_text:
                        rec_text = next(
                            (x for x in cells if x.strip() and x.strip() not in _REC_LABELS_FB
                             and "推荐人意见" not in x and re.search(r"[，,]", x)),
                            "",
                        )
                    if rec_text:
                        recs.append(parse_recommendation(rec_text))
                # 格式二：独立表格，row0=「六、专家推荐意见」，row1=姓名/职务/单位，row3=正文
                if not recs and extra_rows >= 4:
                    hdr0 = extra_row_texts[0] if extra_rows > 0 else ""
                    if re.search(r"专家推荐意见|推荐意见", hdr0):
                        row1_cells = consec_unique(extra_t.rows[1]) if extra_rows > 1 else []
                        row3_cells = consec_unique(extra_t.rows[3]) if extra_rows > 3 else []
                        content = " ".join(
                            c.strip() for c in row3_cells if c and c.strip()
                        ).strip()
                        if content:
                            # 从 row1 提取姓名、单位
                            r1_vals = [c.strip() for c in row1_cells if c and c.strip()]
                            name_val, title_val = "", ""
                            for idx_v, v in enumerate(r1_vals):
                                if v in ("姓          名", "姓名"):
                                    name_val = r1_vals[idx_v + 1] if idx_v + 1 < len(r1_vals) else ""
                                if v in ("工作单位",):
                                    title_val = r1_vals[idx_v + 1] if idx_v + 1 < len(r1_vals) else ""
                            intro = f"{name_val}，{title_val}：" if name_val else ""
                            recs.append(parse_recommendation(intro + content if intro else content))
                if recs:
                    data["recommendations"] = recs
                    need_recs = False

            if need_tasks:
                task_idx_fb = next(
                    (i for i, txt in enumerate(extra_row_texts)
                     if "岗位任务" in txt and (
                         "支撑保障" in txt or "学院帮助" in txt
                         or re.search(r"1[.．、]\s*岗位任务", txt))),
                    -1,
                )
                if task_idx_fb < 0:
                    task_idx_fb = next(
                        (i for i, txt in enumerate(extra_row_texts)
                         if "岗位任务" in txt and re.search(r"1[.．、]\s*岗位任务", txt)),
                        -1,
                    )
                if task_idx_fb >= 0:
                    task_row_fb = extra_t.rows[task_idx_fb]
                    seen_fb = set()
                    ct_fb = []
                    for cell in task_row_fb.cells:
                        cid = id(cell._element)
                        if cid in seen_fb:
                            continue
                        seen_fb.add(cid)
                        ct_fb.append(_cell_text_with_auto_num(cell, doc))
                    block_fb = "\n".join(c.strip() for c in ct_fb if c.strip())
                    jt_fb, cs_fb = _extract_job_tasks_college_support_block(block_fb)
                    if jt_fb:
                        data["job_tasks"] = jt_fb
                        data["college_support"] = cs_fb
                        need_tasks = False

            if not need_recs and not need_tasks:
                break

    return data


def main():
    """
    从「正确」Word 批量抽取 JSON；作者等单元格的加粗/下划线由 read_cell_marked 编码。
    调优与正确版一致：用本脚本从正确 .docx 更新 json → 再运行 fill_template.py 生成 docx 对比。
    输入/输出目录默认为脚本所在目录下的「正确示例10份」「正确示例10份_json」。
    """
    base = Path(__file__).resolve().parent
    in_dir = base / "正确示例10份"
    out_dir = base / "正确示例10份_json"
    out_dir.mkdir(parents=True, exist_ok=True)

    for docx_path in sorted(in_dir.glob("*.docx")):
        if docx_path.name.startswith("~$"):
            continue
        try:
            data = extract_one(docx_path)
            out_file = out_dir / f"{docx_path.stem}.json"
            out_file.write_text(
                json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            print(f"done: {out_file.name}")
        except Exception as e:
            import traceback
            print(f"ERROR {docx_path.name}: {e}")
            traceback.print_exc()


if __name__ == "__main__":
    main()
